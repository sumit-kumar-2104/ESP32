"""Build the Word document: wireless wave-based perception and meta-perception."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "assets")
OUTPUT = os.path.join(os.path.dirname(HERE), "Wireless_Wave_Perception_and_MetaPerception.docx")

# Light, presentable palette
INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)     # headings
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)    # soft green accent
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_GREY = "F2F3F5"
HEADER_FILL = "DCE6EF"

doc = Document()

# ---- base styles -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), "8")
    e.set(qn("w:color"), color)
    pbdr.append(e)
    pPr.append(pbdr)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def para(text, italic=False, size=11, align=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def runs_para(parts):
    """parts: list of (text, bold, italic)."""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p


def plain_box(title, text):
    p = doc.add_paragraph()
    _shade(p, SHADE_GREEN)
    _border(p, color="AFC8B2")
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    return p


def note_box(text):
    p = doc.add_paragraph()
    _shade(p, SHADE_GREY)
    _border(p, color="C9CDD3")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


_eqn = {"n": 0}


def equation(img, label=True):
    _eqn["n"] += 1
    from PIL import Image
    path = os.path.join(ASSETS, img)
    with Image.open(path) as im:
        w_in = im.width / 200.0  # saved at 200 dpi
        h_in = im.height / 200.0
    max_w = 5.8
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if w_in > max_w:
        run.add_picture(path, width=Inches(max_w))
    else:
        # keep a readable minimum height for very short equations
        target_h = max(min(h_in, 0.55), 0.30)
        run.add_picture(path, height=Inches(target_h))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if label:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"( {_eqn['n']} )")
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE_LT
        cap.paragraph_format.space_after = Pt(8)
    return _eqn["n"]


_fig = {"n": 0}


def figure(img, caption, width=6.2):
    _fig["n"] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(width))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {_fig['n']}.  ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = SLATE
    r2 = cap.add_run(caption)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = INK
    cap.paragraph_format.space_after = Pt(12)


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        # also shade the cell
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = INK
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F6F8FA")
                tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ===========================================================================
# TITLE PAGE
# ===========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("From Light to Radio")
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(70)
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Mathematical Case for Wireless Wave-Based Perception")
r.font.size = Pt(16)
r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Extending optical wave-computation and diffractive models to wireless sensing, "
                 "with applications to multi-modal meta-perception")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = INK

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  Technical White Paper  \u2014")
r.font.size = Pt(11)
r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(24)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Prepared for the Multi-Modal Perception programme  \u00b7  April 2026")
r.font.size = Pt(10.5)
r.font.color.rgb = SLATE_LT

doc.add_page_break()

# ===========================================================================
# TABLE OF CONTENTS (field, updates on open)
# ===========================================================================
h1("Contents")
toc_p = doc.add_paragraph()
run = toc_p.add_run()
fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
t_run = OxmlElement("w:t"); t_run.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
run._r.append(fldChar); run._r.append(instr); run._r.append(fldChar2); run._r.append(t_run); run._r.append(fldChar3)
doc.add_page_break()

# ===========================================================================
# 1. EXECUTIVE SUMMARY
# ===========================================================================
h1("1.  Executive Summary")
para("Recent work has shown that light can do more than carry pictures: a properly engineered "
     "optical system can perform the computation of a neural network directly in the physical wave, "
     "spending almost no electronic energy. Diffractive optical neural networks and, most recently, "
     "optical generative models (Lin et al., Science 2018; Chen et al., Nature 2025) demonstrate that "
     "perception and even image synthesis can be embedded into wave propagation itself.")
para("This document makes the case \u2014 in plain language and with explicit mathematics \u2014 that "
     "the same idea transfers to wireless. Light and radio are the same physical phenomenon "
     "(electromagnetic waves) governed by the same equations; they differ only in wavelength. "
     "Because Maxwell\u2019s equations are scale-invariant, any wave-based computation or perception "
     "demonstrated optically has a faithful wireless counterpart, obtained by scaling the geometry. "
     "The single price paid is spatial resolution, which scales with wavelength.")
para("We then move from the single-sensor argument to the practical goal of the programme: "
     "meta-perception \u2014 combining cameras and wireless into one perception system. We give a fusion "
     "formulation, a catalogue of application directions with the data each one needs, a recipe for "
     "calculating those data requirements from first principles, and a focused comparison with event "
     "cameras that explains why an active wireless channel is the right tool for continuous, "
     "always-on monitoring.")
plain_box("In plain terms:",
          "Cameras see with light waves. Wi-Fi and radar \u201csee\u201d with radio waves. They are the same "
          "kind of wave, so the clever maths that lets light compute and perceive also works for radio "
          "\u2014 just at a coarser level of detail. The rest of this paper proves that and shows where to use it.")

# ===========================================================================
# 2. THE QUESTION AND THE CLAIM
# ===========================================================================
h1("2.  The Question and the Claim")
para("The motivating question is simple: optical models clearly work for perception and generation "
     "\u2014 can wireless models work the same way? We separate this into two claims, because the source "
     "material points at two distinct (and complementary) ideas.")
runs_para([("Claim A (sensing). ", True, False),
           ("Wireless signals can be used as a perception modality with the same deep-learning "
            "machinery used for images \u2014 classification, detection, pose, activity, identity \u2014 by "
            "treating the measured channel as the \u201cimage\u201d.", False, False)])
runs_para([("Claim B (computation). ", True, False),
           ("The wave itself can carry out part of the computation, exactly as in diffractive optical "
            "networks, using programmable metasurfaces at radio frequencies. This is the true analogue "
            "of the optical generative model.", False, False)])
para("Both are supported by published, peer-reviewed work. Sections 4\u20135 establish the mathematics "
     "behind both claims; the remainder of the document turns the result into an engineering plan.")

# ===========================================================================
# 3. PLAIN-LANGUAGE INTUITION
# ===========================================================================
h1("3.  Plain-Language Intuition: Waves Are Waves")
para("Imagine dropping a stone in a pond. The ripples spread, bounce off the edges, and where two "
     "ripples meet they add up or cancel out. If you place a few shaped obstacles in the pond, the "
     "pattern of ripples on the far side becomes a complicated, but completely predictable, function "
     "of the obstacles. If you are allowed to choose the shapes of those obstacles, you can sculpt the "
     "far-side pattern to be almost anything you want. That is, in one sentence, how a wave computes.")
para("A diffractive optical network is exactly this, using light instead of water: a few thin patterned "
     "surfaces bend the light so that the bright spot lands on \u201cdetector 7\u201d when the input is a "
     "handwritten seven. Nothing is calculated by a processor; the answer forms physically as the light "
     "crosses the surfaces.")
plain_box("In plain terms:",
          "Radio waves ripple, bounce and interfere in precisely the same way as light and water waves. "
          "So if patterned surfaces can make light \u201ccompute\u201d, patterned surfaces can make radio "
          "compute too. And because a person standing in a room changes how the radio ripples bounce, "
          "the radio field already carries a coarse picture of the room \u2014 we just have to learn to read it.")
figure("fig_spectrum.png",
       "Visible light and wireless bands are points on one continuous electromagnetic spectrum. They "
       "obey the identical wave equation; only the wavelength differs.", width=6.4)

# ===========================================================================
# 4. MATHEMATICAL FOUNDATIONS
# ===========================================================================
h1("4.  Mathematical Foundations")

h2("4.1  The one wave equation shared by optics and wireless")
para("Start from Maxwell\u2019s equations in a linear, source-free, isotropic medium. Each Cartesian "
     "component of the electric field U (and likewise the magnetic field) satisfies the time-harmonic "
     "Helmholtz equation:")
equation("eq_helmholtz.png")
para("Here k is the wavenumber, \u03bb the wavelength, \u03c9 the angular frequency and c the speed of "
     "light. This equation contains no notion of \u201coptical\u201d or \u201cradio\u201d. A visible photon "
     "(\u03bb \u2248 0.5 \u00b5m) and a Wi-Fi carrier (\u03bb \u2248 6 cm) are described by the same equation "
     "with a different numerical value of \u03bb. Everything that follows is a consequence of this single "
     "fact (Goodman 2017; Jackson 1999).")

h2("4.2  How a wave computes: diffraction as a linear operator")
para("Given the field on an input plane, the field on any later plane is obtained by the "
     "Rayleigh\u2013Sommerfeld diffraction integral \u2014 a convolution of the input field with a "
     "propagation kernel:")
equation("eq_rs.png")
equation("eq_rs_kernel.png")
para("Equivalently, in the spatial-frequency domain, propagation is a simple phase multiplication "
     "(the angular-spectrum method):")
equation("eq_angular.png")
para("The key structural facts are: (i) propagation is linear in the field; (ii) it is fully "
     "determined by \u03bb and the distance z; and (iii) every input point contributes to every output "
     "point within the numerical aperture \u2014 i.e. it is densely connected, like a fully-connected "
     "neural-network layer. These three properties hold for any wavelength.")

h2("4.3  The diffractive neural network")
para("A diffractive deep neural network (D\u00b2NN) interleaves free-space propagation with thin "
     "trainable surfaces. Each surface applies a complex transmission t (an amplitude a and a phase "
     "\u03c6) to the field. One layer is therefore a modulation followed by a propagation:")
equation("eq_layer.png")
para("After L layers the network reads out the intensity on a detector array, which supplies the "
     "nonlinearity:")
equation("eq_output.png")
para("The trainable parameters are the phase profiles \u03c6\u2097(x,y) of each surface. They are optimised "
     "by gradient descent against a task loss, exactly as digital network weights are (Lin et al. 2018). "
     "The optical generative model of Chen et al. (2025) is the generative version of this same "
     "construction: a shallow digital encoder produces phase \u201cseeds\u201d and a fixed diffractive "
     "decoder turns them into images, with essentially zero compute during synthesis.")

h2("4.4  The scale-invariance theorem (the proof that wireless inherits all of this)")
para("We now state precisely why a wireless system can reproduce any optical wave computation.")
note_box("Proposition (electromagnetic similitude). Let an optical diffractive system operate at "
         "wavelength \u03bb\u2092 with surfaces {t\u2097(x,y)} placed at distances {z\u2097}. Define the scale "
         "factor s = \u03bb_r / \u03bb\u2092. Construct a wireless system at wavelength \u03bb_r whose surfaces "
         "implement the same complex transmission sampled on coordinates stretched by s, placed at "
         "distances {s\u00b7z\u2097}. Then the two systems implement the identical input\u2013output map (up to the "
         "overall geometric scaling).")
para("Proof sketch. Apply the coordinate change x \u2192 s x, y \u2192 s y, z \u2192 s z and \u03bb \u2192 s \u03bb to "
     "the Helmholtz equation and to the Rayleigh\u2013Sommerfeld kernel. The wavenumber transforms as "
     "k \u2192 k/s, while every distance grows by s, so the dimensionless product k\u00b7r that controls all "
     "interference is left unchanged:")
equation("eq_scale.png")
para("Because the propagation operator depends on the geometry only through k\u00b7r, the propagation "
     "between the scaled surfaces is identical to the original. The modulation step is invariant by "
     "construction (we copied the same t\u2097). Composing L identical layers yields an identical overall "
     "transform. Hence the wireless network computes the same function as the optical one. \u220e")
para("This is the classical scale invariance of Maxwell\u2019s equations (Jackson 1999, \u00a7 on "
     "similitude). One caveat must be stated honestly: exact invariance also requires material "
     "conductivity to scale as \u03c3 \u2192 \u03c3/s. In practice we do not rely on scaling natural materials \u2014 "
     "we use programmable metasurfaces that are engineered to synthesise the required surface response "
     "{t\u2097} directly at the radio wavelength. This is precisely what reconfigurable-metasurface imagers "
     "and recognisers already do (Li et al., Nature Communications 2019; Li et al., Light: Science & "
     "Applications 2019; Hunt et al., Science 2013).")
figure("fig_diffractive.png",
       "The scale-invariance result in pictures: the same trainable diffractive network is realised "
       "optically (microns, centimetre-scale hardware) and at radio frequency (centimetres, "
       "metre-scale hardware). The topology is identical; only the physical scale changes.", width=6.6)

h2("4.5  What actually changes: the resolution limit")
para("Scale invariance preserves the computation but not the fineness of spatial detail. The smallest "
     "feature a wave system can resolve is set by the diffraction limit:")
equation("eq_res.png")
para("where NA is the numerical aperture. Because \u03bb is 10\u2074\u201310\u2075 times larger for radio than for "
     "visible light, a wireless system built at the same NA resolves features 10\u2074\u201310\u2075 times "
     "coarser. This is the fundamental reason wireless perception delivers presence, motion, pose and "
     "activity \u2014 not fine texture or readable faces. It is a feature, not only a limitation: coarse, "
     "non-identifying sensing is exactly what privacy-preserving, always-on monitoring wants.")
figure("fig_resolution.png",
       "Diffraction-limited resolution versus wavelength (fixed numerical aperture). The wireless bands "
       "sit far to the upper right: usable for coarse spatial perception, not for fine imaging.", width=5.6)

# ===========================================================================
# 5. WIRELESS PERCEPTION AS A MEASURED WAVE FIELD
# ===========================================================================
h1("5.  Wireless Perception as a Measured Wave Field")
para("Claim B above puts computation inside the hardware. Claim A is the route most deployable today: "
     "measure the wave field with ordinary radios and let a digital network read it. The two share the "
     "same physics; only the location of the \u201cnetwork\u201d differs.")

h2("5.1  The channel model")
para("A radio receiver measures the channel \u2014 the linear transformation the environment applies to "
     "the transmitted signal:")
equation("eq_channel.png")
para("For a wideband OFDM system (e.g. Wi-Fi), the receiver reports Channel State Information (CSI): "
     "one complex number per subcarrier, which is a sum over propagation paths p with complex gain "
     "a\u209a and delay \u03c4\u209a (Halperin et al. 2011; Ma et al. 2019):")
equation("eq_csi.png")
para("The CSI tensor \u2014 indexed by transmit antenna, receive antenna, subcarrier and time \u2014 is the "
     "wireless equivalent of an image\u2019s pixel grid. It is a direct sample of the same complex wave "
     "field that an optical sensor would record, only at radio wavelength.")

h2("5.2  Static and dynamic scattering")
para("Separating the channel into a stationary part and a moving part is what makes wireless useful "
     "for both presence and motion:")
equation("eq_csi_dyn.png")
para("The static term H_static encodes the fixed geometry of the room and any motionless objects or "
     "people; the moving term encodes Doppler shifts from displacement d_m(t). This decomposition is "
     "central to Section 9, where it explains the difference from event cameras.")

h2("5.3  Learning the inverse map")
para("Perception is the inverse problem: recover the scene state s (presence, pose, activity) from the "
     "measured field. A neural network f_\u03b8 is trained to approximate this inverse, exactly as a "
     "vision network maps pixels to labels. Published systems demonstrate the full range of tasks: "
     "through-wall human pose (Zhao et al., CVPR 2018; SIGCOMM 2018), fine-grained body perception "
     "(Wang et al., ICCV 2019), dense body surfaces from commodity Wi-Fi (Geng et al. 2023), gesture "
     "recognition (Pu et al., MobiCom 2013), and standardised benchmarks (Yang et al., Patterns 2023).")

# ===========================================================================
# 6. META-PERCEPTION
# ===========================================================================
h1("6.  From One Modality to Meta-Perception")
h2("6.1  Why fuse vision and wireless")
para("Vision is high-resolution but fails in darkness, glare and occlusion, is privacy-heavy and "
     "expensive to run continuously. Wireless is coarse and label-poor but cheap, private, "
     "penetrating and always-on. Their failure modes are largely complementary, which is the textbook "
     "condition under which sensor fusion yields more than the sum of parts.")

h2("6.2  A fusion formulation")
para("Treat the scene state s as a latent variable and each modality as a noisy measurement. Assuming "
     "conditional independence of the two sensors given the scene, the posterior factorises:")
equation("eq_fusion.png")
para("This single expression captures the whole design space. Late fusion evaluates the two "
     "likelihoods separately and multiplies; early fusion learns a joint likelihood; the uncertainty "
     "weighting falls out naturally because a confident modality has a sharply peaked likelihood and "
     "therefore dominates. It also justifies the wake-on-event policy: when the camera is off, its "
     "likelihood is uninformative (flat) and the posterior reduces to the wireless term.")

h2("6.3  Reference architecture")
figure("fig_fusion.png",
       "Meta-perception architecture. The wireless branch runs continuously at low cost; the decision "
       "engine wakes the camera only on confident events; late fusion combines both with "
       "uncertainty weighting.", width=6.4)

# ===========================================================================
# 7. APPLICATIONS AND DIRECTIONS
# ===========================================================================
h1("7.  Applications and Research Directions")
para("The combination of always-on wireless and on-demand vision opens a broad application space. "
     "The table catalogues the most promising directions, each with a concrete example, the sensing "
     "data it needs, and why fusion helps. The right-hand column is also a shortlist for choosing one "
     "direction to take forward.")
make_table(
    ["Direction", "Concrete example", "Data required", "Why wireless + vision"],
    [
        ["Elderly / home care", "Fall detection and 24/7 presence in a flat",
         "Wi-Fi CSI (antenna \u00d7 subcarrier \u00d7 time); sparse camera keyframes for labels",
         "Private, works in the dark; camera only confirms a confident fall event"],
        ["Security & search-rescue", "Through-wall presence and motion in smoke/darkness",
         "mmWave / UWB radar returns; IR or RGB when available",
         "RF penetrates walls and smoke; vision adds identity when line-of-sight returns"],
        ["Automotive / robotics", "Perception in fog, rain, glare",
         "Automotive radar point clouds + camera + (optionally) lidar",
         "Radar robust to weather; camera supplies semantics and texture"],
        ["Vital-sign monitoring", "Contactless respiration and heart-rate at the desk or bedside",
         "Fine-phase CSI or FMCW radar at sub-mm displacement sensitivity",
         "Continuous, contactless; camera contextualises posture and activity"],
        ["Smart buildings", "Occupancy and activity analytics for HVAC and lighting",
         "Coarse CSI presence features per zone; periodic camera audit",
         "Cheap always-on counting; camera calibrates and validates"],
        ["AR / VR & HCI", "Wearable-free body and gesture tracking",
         "Multi-antenna CSI or 60 GHz radar; camera for ground-truth skeletons",
         "RF tracks through occlusion; vision bootstraps and refines"],
    ],
    widths=[1.2, 1.7, 2.1, 1.9],
)
para("Selecting one direction. A practical down-selection scores each row on four axes: data "
     "availability (can we collect labelled CSI cheaply?), resolution sufficiency (does the diffraction "
     "limit of Section 4.5 permit the task?), regulatory and privacy fit, and hardware cost. Elderly "
     "care and smart-building occupancy typically score highest because they tolerate coarse resolution, "
     "benefit most from privacy, and can be labelled with a single temporary camera.")

# ===========================================================================
# 8. DATA REQUIREMENTS AND HOW TO CALCULATE THEM
# ===========================================================================
h1("8.  Data Requirements and How to Size Them")
para("A recurring question is: what data does a wireless perception system actually need, and how do "
     "we calculate it before building anything? The answer comes from four back-of-envelope formulas, "
     "each tied to a physical degree of freedom.")

h3("8.1  Spatial degrees of freedom (the wireless \u201cpixel count\u201d)")
para("The number of independent spatial channels between a transmit aperture of area A_T and a "
     "receive aperture of area A_R separated by distance d is the space\u2013bandwidth product, sometimes "
     "called the Shannon number:")
equation("eq_shannon.png")
para("This is the wireless analogue of an image\u2019s pixel count: it bounds how much spatial detail the "
     "measurement can hold. To perceive a scene whose description needs D degrees of freedom (for "
     "example, a J-joint skeleton needs 3J numbers), one requires N_dof \u2273 D, with margin for noise.")

h3("8.2  Range, velocity and angle resolution")
para("Three independent resolutions set what the system can distinguish:")
runs_para([("Range resolution ", True, False), ("from bandwidth B:", False, False)])
equation("eq_range.png")
runs_para([("Angular resolution ", True, False), ("from aperture size D:", False, False)])
equation("eq_angle.png")
para("Velocity (Doppler) resolution follows from the observation time T as \u0394v \u2248 \u03bb/(2T). Together "
     "these convert a perception specification into hardware requirements: \u201cresolve two people 30 cm "
     "apart in range\u201d demands B \u2265 c/(2\u00b70.30) \u2248 0.5 GHz, and so on.")
figure("fig_data_requirements.png",
       "Left: range resolution versus bandwidth \u2014 reading off the bandwidth a task demands. Right: "
       "independent spatial channels versus aperture size \u2014 reading off how much spatial detail an "
       "array can capture.", width=6.6)

h3("8.3  How many measurements: the compressed-sensing bound")
para("If the scene is sparse \u2014 S occupied voxels out of N \u2014 then far fewer measurements than N are "
     "needed. Compressed-sensing theory (Cand\u00e8s & Wakin 2008) gives the bound on the number of "
     "incoherent measurements M:")
equation("eq_cs.png")
para("For example, locating S = 3 people in a grid of N = 1000 candidate cells needs only on the order "
     "of a few tens of measurements, not a thousand \u2014 which is why a handful of antennas and "
     "subcarriers already supports coarse multi-person sensing.")

h3("8.4  Training-data budget")
para("Claim A is data-driven, so a labelled-data estimate matters. As a rule of thumb the number of "
     "labelled windows scales with model capacity and class count; published Wi-Fi benchmarks (SenseFi; "
     "Yang et al. 2023) reach usable activity-recognition accuracy with thousands to tens of thousands "
     "of labelled windows per environment, and transfer learning cuts the per-site cost substantially. "
     "A camera running only during a short enrolment phase can auto-label this data, which is why the "
     "fusion architecture also solves the wireless labelling problem.")

h3("8.5  Worked example")
note_box("Goal: 24/7 fall detection in a 5 m \u00d7 5 m room, distinguishing standing, sitting, walking and "
         "falling. Range: 30 cm is enough to separate posture transitions \u2192 B \u2248 0.5 GHz (achievable "
         "by aggregating Wi-Fi channels or a low-cost UWB radio). Spatial DoF: a 3 \u00d7 3 antenna pair at "
         "\u03bb = 6 cm over d = 4 m gives a modest but sufficient N_dof for whole-body posture. Measurements: "
         "with S \u2248 1 person in N \u2248 100 cells, M on the order of 10\u201320 incoherent CSI features suffice. "
         "Training: a few thousand labelled windows, auto-labelled by a camera during a one-day "
         "enrolment, then the camera sleeps. This is a buildable specification derived entirely from "
         "the formulas above.")

# ===========================================================================
# 9. EVENT CAMERAS VS ALWAYS-ON WIRELESS
# ===========================================================================
h1("9.  Event Cameras versus Always-On Wireless")
h2("9.1  What an event camera is")
para("An event (or neuromorphic) camera does not output frames. Each pixel fires asynchronously only "
     "when the brightness it sees changes by a set threshold, with microsecond latency, very high "
     "dynamic range and low power (Lichtsteiner et al. 2008; Gallego et al. 2022). It is an excellent "
     "motion and change detector.")

h2("9.2  The shared trait \u2014 and the trap")
para("Event cameras and Doppler-based wireless sensing share a tempting property: both respond "
     "strongly to change and stay quiet when nothing moves. For a purely motion-triggered design this "
     "looks efficient. But the programme\u2019s goal is continuous monitoring, including a person who is "
     "present but still \u2014 asleep, seated, or collapsed and motionless. A change-only sensor goes blind "
     "in exactly the situation that matters most.")
figure("fig_event_vs_wireless.png",
       "Top: an event camera emits data only while the scene changes and is silent on a motionless "
       "person. Bottom: an active wireless link is always illuminated, so a still person shifts the "
       "static channel (presence) and motion adds Doppler \u2014 both are observable.", width=6.6)

h2("9.3  The crucial difference: active illumination and static presence")
para("The decisive distinction is that an event camera is passive \u2014 it waits for ambient light to "
     "change \u2014 whereas a wireless sensor is active: the transmitter continuously illuminates the scene, "
     "so a fresh measurement of the full field is always available. Returning to Section 5.2, a "
     "motionless person still alters the static term H_static (their body reflects and absorbs the "
     "wave), so they remain detectable by comparing the present channel to a learned empty-room "
     "reference \u2014 no motion required. The Doppler term adds motion sensitivity on top. A wireless "
     "system therefore spans both regimes that an event camera cannot cover alone: static presence and "
     "dynamic motion.")

h2("9.4  Recommended design: differ by staying always-on")
para("We deliberately differ from the event-driven direction. Rather than trigger only on change, the "
     "wireless front-end maintains a continuous reference field and performs two operations at once: "
     "(i) background-referenced detection of static presence, and (ii) Doppler-based detection of "
     "motion. Cheap change-detection can still gate the heavier vision branch \u2014 capturing the energy "
     "benefit event cameras chase \u2014 but the wireless layer itself never goes blind. This is the "
     "property that makes wireless, not an event camera, the correct substrate for 24/7 perception.")

# ===========================================================================
# 10. LIMITATIONS
# ===========================================================================
h1("10.  Limitations and Honest Caveats")
para("Three caveats keep the claims rigorous. First, resolution: the diffraction limit (Section 4.5) "
     "means wireless perception is coarse; it will not read faces or fine texture and should be paired "
     "with vision when those are needed. Second, environment dependence: the radio channel is shaped by "
     "multipath and clutter, so models trained in one room transfer imperfectly; calibration and "
     "domain adaptation are active requirements, not optional polish. Third, hardware maturity: Claim B "
     "(in-wave computation by programmable metasurfaces) is demonstrated for imaging and recognition "
     "but is earlier-stage than the optical generative result, whereas Claim A (learned perception from "
     "measured CSI) is deployable today on commodity radios.")

# ===========================================================================
# 11. CONCLUSION
# ===========================================================================
h1("11.  Conclusion")
para("Optical models work for perception because perception is, at bottom, the reading and shaping of "
     "a wave field \u2014 and because Maxwell\u2019s equations let a trained piece of hardware do that reading "
     "and shaping physically. Wireless is the same wave at a longer wavelength. The scale-invariance "
     "argument of Section 4 shows the mathematics transfers exactly; the channel model of Section 5 "
     "shows it is already practical with ordinary radios; and Sections 6\u20139 turn the result into a "
     "concrete meta-perception system that monitors continuously, sizes its own data requirements, and "
     "succeeds precisely where vision and event cameras fail. The answer to the original question is "
     "therefore yes \u2014 wireless models can work like optical models for perception, with resolution as "
     "the one principled trade-off, and with privacy-preserving, always-on coverage as the reward.")

# ===========================================================================
# 12. REFERENCES
# ===========================================================================
h1("12.  References")
refs = [
    "Lin, X., Rivenson, Y., Yardimci, N. T., Veli, M., Luo, Y., Jarrahi, M., & Ozcan, A. (2018). "
    "All-optical machine learning using diffractive deep neural networks. Science, 361(6406), "
    "1004\u20131008. https://doi.org/10.1126/science.aat8084",
    "Chen, S., Li, Y., Wang, Y., Chen, H., & Ozcan, A. (2025). Optical generative models. Nature, "
    "644, 903\u2013911. https://doi.org/10.1038/s41586-025-09446-5",
    "Goodman, J. W. (2017). Introduction to Fourier Optics (4th ed.). W. H. Freeman.",
    "Jackson, J. D. (1999). Classical Electrodynamics (3rd ed.). Wiley. (Scale invariance / "
    "electromagnetic similitude.)",
    "Li, L., Ruan, H., Liu, C., Li, Y., Shuang, Y., Al\u00f9, A., Qiu, C.-W., & Cui, T. J. (2019). "
    "Machine-learning reprogrammable metasurface imager. Nature Communications, 10, 1082. "
    "https://doi.org/10.1038/s41467-019-09103-2",
    "Li, L., Shuang, Y., Ma, Q., Li, H., Zhao, H., Wei, M., Liu, C., Hao, C., Qiu, C.-W., & Cui, T. J. "
    "(2019). Intelligent metasurface imager and recognizer. Light: Science & Applications, 8, 97. "
    "https://doi.org/10.1038/s41377-019-0209-z",
    "Hunt, J., Driscoll, T., Mrozack, A., Lipworth, G., Reynolds, M., Brady, D., & Smith, D. R. "
    "(2013). Metamaterial apertures for computational imaging. Science, 339(6117), 310\u2013313. "
    "https://doi.org/10.1126/science.1230054",
    "Cui, T. J., Qi, M. Q., Wan, X., Zhao, J., & Cheng, Q. (2014). Coding metamaterials, digital "
    "metamaterials and programmable metamaterials. Light: Science & Applications, 3, e218. "
    "https://doi.org/10.1038/lsa.2014.99",
    "Zhao, M., Li, T., Abu Alsheikh, M., Tian, Y., Zhao, H., Torralba, A., & Katabi, D. (2018). "
    "Through-wall human pose estimation using radio signals. CVPR 2018, 7356\u20137365. "
    "https://doi.org/10.1109/CVPR.2018.00768",
    "Zhao, M., Tian, Y., Zhao, H., Abu Alsheikh, M., Li, T., Hristov, R., Kabelac, Z., Katabi, D., & "
    "Torralba, A. (2018). RF-based 3D skeletons. ACM SIGCOMM 2018, 267\u2013281. "
    "https://doi.org/10.1145/3230543.3230579",
    "Wang, F., Panev, S., Dai, Z., Han, J., & Huang, D. (2019). Person-in-WiFi: Fine-grained person "
    "perception using WiFi. ICCV 2019, 5451\u20135460. https://doi.org/10.1109/ICCV.2019.00540",
    "Geng, J., Huang, D., & De la Torre, F. (2023). DensePose from WiFi. arXiv:2301.00250.",
    "Yang, J., Chen, X., Zou, H., Wang, D., Xu, Q., & Xie, L. (2023). SenseFi: A library and benchmark "
    "on deep-learning-empowered WiFi human sensing. Patterns, 4(3), 100703. "
    "https://doi.org/10.1016/j.patter.2023.100703",
    "Ma, Y., Zhou, G., & Wang, S. (2019). WiFi sensing with channel state information: A survey. ACM "
    "Computing Surveys, 52(3), 46. https://doi.org/10.1145/3310194",
    "Halperin, D., Hu, W., Sheth, A., & Wetherall, D. (2011). Tool release: gathering 802.11n traces "
    "with channel state information. ACM SIGCOMM CCR, 41(1), 53. "
    "https://doi.org/10.1145/1925861.1925870",
    "Adib, F., & Katabi, D. (2013). See through walls with WiFi! ACM SIGCOMM 2013, 75\u201386. "
    "https://doi.org/10.1145/2486001.2486039",
    "Pu, Q., Gupta, S., Gollakota, S., & Patel, S. (2013). Whole-home gesture recognition using "
    "wireless signals (WiSee). ACM MobiCom 2013, 27\u201338. https://doi.org/10.1145/2500423.2500436",
    "Cand\u00e8s, E. J., & Wakin, M. B. (2008). An introduction to compressive sampling. IEEE Signal "
    "Processing Magazine, 25(2), 21\u201330. https://doi.org/10.1109/MSP.2007.914731",
    "Gallego, G., Delbr\u00fcck, T., Orchard, G., Bartolozzi, C., Taba, B., Censi, A., Leutenegger, S., "
    "Davison, A. J., Conradt, J., Daniilidis, K., & Scaramuzza, D. (2022). Event-based vision: A "
    "survey. IEEE TPAMI, 44(1), 154\u2013180. https://doi.org/10.1109/TPAMI.2020.3008413",
    "Lichtsteiner, P., Posch, C., & Delbr\u00fcck, T. (2008). A 128\u00d7128 120 dB 15 \u00b5s latency "
    "asynchronous temporal contrast vision sensor. IEEE J. Solid-State Circuits, 43(2), 566\u2013576. "
    "https://doi.org/10.1109/JSSC.2007.914337",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.space_after = Pt(6)
    rn = p.add_run(f"[{i}]  ")
    rn.bold = True
    rn.font.color.rgb = SLATE
    rn.font.size = Pt(10)
    r2 = p.add_run(ref)
    r2.font.size = Pt(10)

note_box("Note on sources: every reference above is a peer-reviewed paper, an archival conference "
         "proceeding, a standard textbook, or a public preprint with a verifiable DOI/arXiv identifier. "
         "The two foundational claims are anchored on Lin et al. (2018) and Chen et al. (2025) for "
         "optical wave computation, and on the Cui-group metasurface imagers and the MIT RF-sensing line "
         "for the wireless counterpart.")

doc.save(OUTPUT)
print("SAVED", OUTPUT)
