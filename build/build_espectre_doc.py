"""Build the ESPectre repository guide (Word document).

Explains what the ESPectre project is, what it is for, how it works, its
requirements, how to use and replicate it, and the exact steps - based on a
local copy of the repository placed in the workspace. Nothing was executed.
No emojis are used.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "assets_wireless")
OUTPUT = os.path.join(os.path.dirname(HERE), "ESPectre_Guide.docx")

INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
RUST = RGBColor(0xA8, 0x5A, 0x4A)
MONO = RGBColor(0x2B, 0x3A, 0x48)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_AMBER = "FBF2E7"
SHADE_GREY = "F2F3F5"
SHADE_CODE = "F4F6F8"
HEADER_FILL = "DCE6EF"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(size)); e.set(qn("w:space"), "8"); e.set(qn("w:color"), color)
    pbdr.append(e); pPr.append(pbdr)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def para(text, italic=False, size=11, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def runs_para(parts):
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        r = p.add_run(text); r.bold = bold; r.italic = italic
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def code_block(lines):
    for i, ln in enumerate(lines):
        p = doc.add_paragraph(); _shade(p, SHADE_CODE); _border(p, color="D3DAE0", where="left")
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9.5); r.font.color.rgb = MONO
        r._element.rPr.rFonts.set(qn("w:cs"), "Consolas")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def info_box(text, kind="green"):
    fill = {"green": SHADE_GREEN, "blue": SHADE_BLUE, "amber": SHADE_AMBER, "grey": SHADE_GREY}[kind]
    ec = {"green": "AFC8B2", "blue": "9DB6CC", "amber": "D9B98A", "grey": "C9CDD3"}[kind]
    lead = {"green": ("Note:  ", ACCENT), "blue": ("How it works:  ", SLATE),
            "amber": ("Important:  ", RUST), "grey": ("", INK)}[kind]
    p = doc.add_paragraph(); _shade(p, fill); _border(p, color=ec)
    if lead[0]:
        r = p.add_run(lead[0]); r.bold = True; r.font.color.rgb = lead[1]
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


_fig = {"n": 0}


def figure(img, caption, width=6.2):
    _fig["n"] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(width))
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {_fig['n']}.  "); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE
    r2 = cap.add_run(caption); r2.font.size = Pt(9.5); r2.font.color.rgb = INK
    cap.paragraph_format.space_after = Pt(12)


def make_table(headers, rows, widths=None, fs=9.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""; p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL); tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; p = cells[i].paragraphs[0]
            r = p.add_run(val); r.font.size = Pt(fs); r.font.color.rgb = INK
            if i == 0:
                r.bold = True; r.font.color.rgb = SLATE
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA"); tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ===========================================================================
# TITLE
# ===========================================================================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ESPectre: A Practical Guide"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(66); title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Wi-Fi (CSI) Motion Detection on ESP32 with Home Assistant")
r.font.size = Pt(15); r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("What it is, how it works, what you need, and the exact steps to set it up and replicate it")
r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = INK

line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  Repository Guide  \u2014"); r.font.size = Pt(11); r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(24)

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Project by Francesco Pace  \u00b7  GPLv3  \u00b7  github.com/francescopace/espectre\n"
                 "A local copy has been placed in this workspace under the espectre/ folder")
r.font.size = Pt(10); r.font.color.rgb = SLATE_LT
doc.add_page_break()

# ===========================================================================
# CONTENTS
# ===========================================================================
h1("Contents")
toc_p = doc.add_paragraph(); run = toc_p.add_run()
fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
t_run = OxmlElement("w:t"); t_run.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
for e in (fldChar, instr, fldChar2, t_run, fldChar3): run._r.append(e)

info_box("This guide is documentation only. As requested, the repository was cloned (obtained) locally and nothing "
         "was built, flashed, or run. The commands shown are the project's own instructions, reproduced here so the "
         "setup can be replicated later.", kind="amber")
doc.add_page_break()

# ===========================================================================
# 1. WHAT IS ESPECTRE
# ===========================================================================
h1("1.  What Is ESPectre?")
para("ESPectre is an open-source system that turns an inexpensive ESP32 Wi-Fi chip into a motion sensor \u2014 with no "
     "camera and no microphone. It listens to the fine-grained radio measurements that Wi-Fi hardware already "
     "produces (Channel State Information, or CSI) and detects when someone moves nearby, because a moving body "
     "disturbs the Wi-Fi waves travelling between the router and the sensor. It integrates natively into Home "
     "Assistant through ESPHome, so the motion signal appears as a normal smart-home sensor.")
runs_para([("In one line:  ", True, False),
           ("a roughly ten-euro ESP32 plus your existing Wi-Fi router becomes a privacy-preserving, through-wall "
            "presence and motion detector for your home.", False, False)])
info_box("When someone moves in a room, they disturb the Wi-Fi waves travelling between the router and the sensor \u2014 "
         "like moving your hand in front of a flashlight and watching the shadow change. The ESP32 listens to those "
         "changes and decides whether there is movement.", kind="blue")
para("Key facts about the project:")
make_table(
    ["Attribute", "Detail"],
    [
        ["Purpose", "Wi-Fi CSI-based motion / presence detection"],
        ["Hardware", "ESP32 with CSI support (C6, S3, C3, C5, original ESP32; S2 experimental)"],
        ["Integration", "ESPHome component with native Home Assistant auto-discovery"],
        ["Cost / setup", "About 10 euro per node; 10-15 minutes; YAML only, no programming required"],
        ["Detection", "MVS (moving variance) by default; an experimental on-device ML detector"],
        ["Privacy", "No cameras or microphones; works through walls; data stays local"],
        ["License", "GNU GPL v3.0"],
        ["Author / repo", "Francesco Pace; github.com/francescopace/espectre (approx. 8.8k stars)"],
    ],
    widths=[1.8, 4.4],
)

h2("1.1  Two platforms in one repository")
para("The project deliberately ships two complementary code bases:")
bullet("the ESPHome component in C++ (the components/espectre folder). Stable, production-ready, "
       "configured entirely in YAML, and meant for end users and smart-home enthusiasts.", "ESPectre (production): ")
bullet("a Python / MicroPython research-and-development platform (the micro-espectre folder). MQTT-based, with "
       "analysis and data-collection tools, meant for researchers who want to prototype new algorithms (people "
       "counting, activity recognition, localisation, gesture recognition).", "Micro-ESPectre (R&D): ")
info_box("Development flow: new algorithms are prototyped and validated in Micro-ESPectre (Python), then ported to "
         "the ESPectre C++ component once proven. For a research project, Micro-ESPectre is the part to study; for a "
         "working home sensor, the ESPHome component is the part to deploy.", kind="grey")

h2("1.2  What you can do with it")
bullet("home security: alert if someone enters while you are away.", None)
bullet("elderly care: monitor activity, detect prolonged inactivity or falls.", None)
bullet("smart automation and energy saving: switch lights, heating or devices by room occupancy.", None)
bullet("child or room monitoring: alert on room entry or exit at night.", None)
para("The R&D platform additionally targets people counting, activity recognition (walking, falling, sitting, "
     "sleeping), localisation and tracking, and gesture recognition.")

# ===========================================================================
# 2. HOW IT WORKS
# ===========================================================================
h1("2.  How It Works")
para("ESPectre runs a focused signal-processing pipeline on the ESP32. Each received Wi-Fi packet yields CSI (an "
     "amplitude and phase for each of the 64 subcarriers in HT20 mode). The pipeline reduces that to a single "
     "motion decision.")
figure("fig_espectre_pipeline.png",
       "The ESPectre processing pipeline. Raw CSI is stabilised (gain lock), reduced to the 12 most informative "
       "subcarriers (NBVI), summarised as a turbulence value, cleaned by filters, and turned into a motion decision "
       "by comparing its moving variance to an adaptive threshold.", width=6.6)
para("The stages, in order:")
numbered("locks the ESP32's automatic gain control and FFT scaling for about 3 seconds (300 packets) so CSI "
         "amplitudes stop drifting. On chips without this capability, a gain-invariant measure (coefficient of "
         "variation) is used instead.", "Gain lock: ")
numbered("automatically selects 12 non-consecutive, most informative subcarriers using the NBVI (Normalized "
         "Baseline Variability Index) algorithm, achieving high accuracy with zero manual configuration.", "Subcarrier selection (NBVI): ")
numbered("computes spatial turbulence \u2014 the standard deviation of the selected subcarrier amplitudes \u2014 as a "
         "single number per packet describing how disturbed the channel is.", "Turbulence: ")
numbered("a Hampel outlier filter (on by default) removes sudden interference spikes; an optional low-pass filter "
         "removes high-frequency noise.", "Filtering: ")
numbered("the variance of turbulence over a sliding window (default 100 packets) is compared to an adaptive "
         "threshold (the 95th percentile of the calibrated baseline times 1.1). Low variance means idle; high "
         "variance means motion. A hit filter debounces the IDLE/MOTION transitions.", "Moving variance and threshold: ")
info_box("Two detectors are available. MVS (Moving Variance Segmentation) is the default and is extremely light "
         "(about 150-440 microseconds per packet). An experimental ML detector replaces the final decision with a "
         "small neural network (an MLP with layers 9 -> 32 -> 16 -> 1) that needs no calibration. Reported motion "
         "detection quality is high on all supported chips (F1-score roughly 96-100 percent).", kind="blue")

# ===========================================================================
# 3. REPOSITORY STRUCTURE
# ===========================================================================
h1("3.  Repository Structure")
para("The local copy in the workspace (espectre/ folder) is laid out as follows:")
make_table(
    ["Path", "What it contains"],
    [
        ["components/espectre/", "The production ESPHome C++ component (detectors, CSI manager, NBVI, filters)"],
        ["examples/", "Ready-to-use YAML configs per board (c6, s3, c3, esp32, ...) plus dashboards"],
        ["micro-espectre/", "Python / MicroPython R&D platform: src, tools, notebooks, data, models, tests"],
        ["docs/", "Additional docs, including the browser game / USB streaming tool"],
        ["test/", "PlatformIO / Unity C++ tests and mocks"],
        ["images/", "Diagrams and photos used by the documentation"],
        ["README.md", "Project overview, quick start, FAQ"],
        ["SETUP.md", "Full installation and configuration guide"],
        ["TUNING.md", "Parameter tuning for optimal detection"],
        ["PERFORMANCE.md", "Benchmarks: confusion matrix, F1-score, resource usage"],
        ["micro-espectre/ALGORITHMS.md", "Scientific documentation of MVS, NBVI, filters, ML"],
        ["ROADMAP.md / CHANGELOG.md", "Plans and version history"],
        ["LICENSE", "GPLv3 license text"],
    ],
    widths=[2.5, 3.9],
)

# ===========================================================================
# 4. REQUIREMENTS
# ===========================================================================
h1("4.  What You Need (Requirements)")
h2("4.1  Hardware")
bullet("an ESP32 board with CSI support. Tested: ESP32-S3, ESP32-C6, ESP32-C5, ESP32-C3, and the original "
       "ESP32; ESP32-S2 is experimental. The C6 and S3 are recommended.", "ESP32 board: ")
bullet("a normal 2.4 GHz home Wi-Fi router (802.11 b/g/n/ax). The one you already own is fine.", "Wi-Fi router: ")
bullet("USB-C or Micro-USB, depending on the board, for flashing.", "USB cable: ")
bullet("optional external antenna (via IPEX connector) for better reception.", "Antenna: ")
info_box("ESPectre uses WiFi 4 (802.11 b/g/n) mode for a stable 64 subcarriers even on WiFi 6 chips, and on the "
         "ESP32-C5 it forces 2.4 GHz to keep CSI behaviour stable. The S3 has the most memory and is best for "
         "advanced/ML use; the C3 is the budget option.", kind="grey")

h2("4.2  Software")
bullet("free, runs on a Raspberry Pi, PC, NAS, or cloud. Optional but recommended.", "Home Assistant: ")
bullet("the firmware framework that hosts the ESPectre component (version 2026.5.0 or newer).", "ESPHome: ")
bullet("required only for the developer install path (version 3.12 recommended; 3.14 has known ESPHome issues).", "Python: ")
bullet("required only for the no-code web-flashing method.", "Google Chrome: ")

h2("4.3  Skills")
bullet("basic YAML editing for configuration.", None)
bullet("Home Assistant familiarity (optional but helpful).", None)
bullet("no programming and no router reconfiguration required for the standard install.", None)

# ===========================================================================
# 5. HOW TO USE / SET IT UP
# ===========================================================================
h1("5.  How to Use It: Setup Paths")
para("The project offers two installation methods. Method A is the fastest and needs no coding; Method B suits "
     "developers who want to customise. Both are reproduced from the project's SETUP.md.")

h2("5.1  Method A \u2014 Web flash (no coding)")
numbered("from the project's GitHub Releases, download the firmware .bin file that matches your chip (for example "
         "espectre-2.5.0-esp32c6.bin).", "Download firmware: ")
numbered("open ESPConnect in Chrome, connect the ESP32 by USB, select the port, choose the .bin, and click Flash.", "Flash: ")
numbered("after flashing, set your Wi-Fi via BLE (ESPHome / Home Assistant app), via USB at web.esphome.io, or via "
         "the \u201cESPectre Fallback\u201d captive-portal Wi-Fi.", "Configure Wi-Fi: ")
numbered("the device is discovered automatically by Home Assistant.", "Done: ")

h2("5.2  Method B \u2014 ESPHome CLI (for developers)")
para("Install ESPHome in a virtual environment, download the example config for your board, then build and flash:")
code_block([
    "# 1) Install ESPHome in a virtual environment",
    "python3 -m venv venv",
    "source venv/bin/activate      # Windows: venv\\Scripts\\activate",
    "pip install esphome",
    "",
    "# 2) Download the example config for your board (e.g. ESP32-C6),",
    "#    then build and flash it",
    "esphome run espectre-c6.yaml",
])
para("After flashing, configure Wi-Fi the same way as in Method A (BLE, USB, or captive portal). The example configs "
     "are pre-set to download the ESPectre component automatically from GitHub, so no manual component copying is "
     "needed.")

h2("5.3  Multiple rooms")
para("For whole-home coverage, flash one ESP32 per room. Each is auto-discovered by Home Assistant with a motion "
     "binary sensor, a movement-score sensor, and an adjustable threshold. This is exactly the per-room Wi-Fi layer "
     "used in the companion camera-reduction study.")
figure("fig_espectre_arch.png",
       "Multi-room deployment. One ESP32 per room reports into Home Assistant over the ESPHome native API, with "
       "automatic discovery; Home Assistant provides dashboards, automations and alerts.", width=6.2)

# ===========================================================================
# 6. REPLICATION FROM SOURCE
# ===========================================================================
h1("6.  How to Replicate It From Source (Development Setup)")
para("To build from the repository (for customisation, offline use, or contribution), the project documents this "
     "flow. It is reproduced here for reference; it was not executed as part of preparing this guide.")
numbered("obtain the repository (the standard command is below). Because Git was not available on this machine, the "
         "workspace copy was instead downloaded as the project's source archive and extracted to the espectre/ "
         "folder \u2014 the file contents are identical to a clone, without the .git history.", "Get the code: ")
code_block([
    "git clone https://github.com/francescopace/espectre.git",
    "cd espectre",
])
numbered("create the ESPHome environment:", "Install ESPHome: ")
code_block([
    "python3 -m venv venv",
    "source venv/bin/activate      # Windows: venv\\Scripts\\activate",
    "pip install esphome",
])
numbered("create a secrets file with your Wi-Fi credentials:", "Add secrets: ")
code_block([
    "# examples/secrets.yaml",
    "wifi_ssid: \"YourWiFiName\"",
    "wifi_password: \"YourWiFiPassword\"",
    "# Optional: lock to a specific access point",
    "# wifi_bssid: \"AA:BB:CC:DD:EE:FF\"",
])
numbered("build and flash a development configuration (these use the local component, DEBUG logging, and extra "
         "debug sensors):", "Build the dev config: ")
code_block([
    "esphome run examples/espectre-c6-dev.yaml    # or -s3-dev, -c3-dev, -esp32-dev",
])
para("Production vs development configs differ as follows:")
make_table(
    ["File", "Component source", "Wi-Fi", "Logger", "Debug sensors"],
    [
        ["espectre-c6.yaml", "GitHub", "Provisioning (BLE/USB/AP)", "INFO", "No"],
        ["espectre-c6-dev.yaml", "Local folder", "secrets.yaml", "DEBUG", "Yes"],
    ],
    widths=[1.9, 1.5, 1.7, 0.8, 0.9], fs=9,
)

h2("6.1  Key configuration parameters")
para("All parameters live under the espectre: section of the YAML. The most important ones:")
make_table(
    ["Parameter", "Default", "Meaning"],
    [
        ["detection_algorithm", "mvs", "Detector: mvs (variance) or ml (neural network)"],
        ["traffic_generator_rate", "100", "Packets/sec used to generate CSI (0-1000)"],
        ["segmentation_threshold", "auto", "Detection threshold: auto, min, or a number"],
        ["segmentation_window_size", "100", "Moving-variance window in packets (10-200)"],
        ["selected_subcarriers", "auto", "Fixed subcarriers, or auto-calibrate (NBVI)"],
        ["hampel_enabled", "true", "Outlier filter (recommended on)"],
        ["lowpass_enabled", "false", "Optional smoothing filter"],
        ["gain_lock", "auto", "AGC/FFT gain lock: auto, enabled, disabled"],
    ],
    widths=[2.4, 1.0, 3.0], fs=9,
)
info_box("Keep the room quiet and still for about 10 seconds after boot in MVS mode: the auto-calibration runs then, "
         "and movement during it hurts accuracy. The ML detector skips calibration. See TUNING.md for full ranges and "
         "troubleshooting.", kind="amber")

h2("6.2  The R&D platform and tests (for research)")
para("Micro-ESPectre (the micro-espectre folder) is the Python/MicroPython side for algorithm work. It exposes CSI "
     "APIs, an MQTT interface, a web monitor, analysis tools, and datasets under micro-espectre/data. The project's "
     "documented commands for running its validation and tests are below \u2014 listed for completeness only; they were "
     "not run here.")
code_block([
    "# C++ component tests (PlatformIO)",
    "cd test && pio test -f test_motion_detection -v",
    "",
    "# Python real-data validation",
    "cd micro-espectre && pytest tests/test_validation_real_data.py -v",
])
para("For building labelled datasets and training the ML detector, the repository documents the workflow in "
     "micro-espectre/ML_DATA_COLLECTION.md and the algorithm details in micro-espectre/ALGORITHMS.md.")

# ===========================================================================
# 7. PERFORMANCE EXPECTATIONS
# ===========================================================================
h1("7.  Performance You Can Expect")
para("From the project's PERFORMANCE.md (verified for v2.8.0), motion detection quality is high across chips:")
make_table(
    ["Chip", "Detector", "Recall", "Precision", "F1-score"],
    [
        ["ESP32-C6", "MVS + NBVI", "99.6%", "100%", "99.8%"],
        ["ESP32-C6", "ML", "100%", "100%", "100%"],
        ["ESP32-S3", "MVS + NBVI", "99.7%", "100%", "99.9%"],
        ["ESP32-C3", "MVS + NBVI", "96.3%", "100%", "98.1%"],
        ["ESP32-C3", "ML", "100%", "100%", "100%"],
    ],
    widths=[1.5, 1.6, 1.1, 1.2, 1.1], fs=9,
)
bullet("MVS detection takes roughly 150-440 microseconds per packet (about 1.5-4.4 percent CPU at 100 packets/sec); "
       "the ML path is heavier but still well within budget.", "Speed: ")
bullet("recommended sensor-to-router distance is 3-8 metres, at 1-1.5 metre height, avoiding metal obstacles and "
       "enclosed corners.", "Placement: ")

# ===========================================================================
# 8. PRIVACY, LICENSE, LOCAL COPY
# ===========================================================================
h1("8.  Privacy, License, and the Local Copy")
bullet("no cameras or microphones; the sensor sees only coarse channel disturbance, and data can stay entirely on "
       "the local hub. Note the project's own caution that Wi-Fi sensing can be misused for surveillance, so deploy "
       "it transparently and only in spaces you control.", "Privacy: ")
bullet("GNU GPL v3.0. You may use, study, modify and redistribute it, but modifications must be shared under the "
       "same license. Contributions require a DCO sign-off (git commit -s).", "License: ")
runs_para([("Local copy in this workspace:  ", True, False),
           ("the repository files are in the ", False, False),
           ("espectre/", False, True),
           (" folder of this MPL workspace. It was obtained by downloading the project's source archive (Git was not "
            "installed), so it contains all source files but no Git history. Nothing was built, flashed, or executed.",
            False, False)])

# ===========================================================================
# 9. HOW THIS FITS THE CAMERA-REDUCTION PLAN
# ===========================================================================
h1("9.  How ESPectre Fits the Camera-Reduction Plan")
para("In the companion feasibility study, ESPectre is the ready-made Wi-Fi-sensing node for each room. One ESP32 per "
     "room, discovered in Home Assistant, delivers reliable presence and motion immediately; the R&D platform is the "
     "path to add activity recognition, people counting, and localisation later, including via camera-supervised "
     "training. In short, ESPectre supplies the practical, low-cost per-room sensing that lets a single camera and a "
     "handful of ten-euro nodes stand in for three or four cameras at the level of the events that matter.")

closing = doc.add_paragraph(); closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("This guide summarises the ESPectre repository (README, SETUP, ALGORITHMS, PERFORMANCE and the "
                    "example configs) as of the locally stored copy. Commands are the project's own; none were "
                    "executed in preparing this document.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE_LT

doc.save(OUTPUT)
print("wrote", OUTPUT)
