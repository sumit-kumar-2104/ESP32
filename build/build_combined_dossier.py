"""Combine the seven MPL Word documents into one properly formatted dossier.

What it does
------------
* Merges the 7 source .docx files (keeping every figure, table and style) using
  docxcompose.
* Removes redundant material:
    - the 7 individual "Contents" / table-of-contents pages,
    - the per-document reference lists (they are merged into ONE deduplicated
      "Consolidated References" section at the end),
    - self-referential "companion document" citations (those documents are now
      part of this dossier).
* Adds a single title page, a short editor's note, one unified table of
  contents, "Part" dividers between sections and page-number footers.

Run:  python build/build_combined_dossier.py
Out:  MPL/Consolidated_Research_Dossier.docx
"""
import os
import re
import copy

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

HERE = os.path.dirname(__file__)
ROOT = os.path.dirname(HERE)
OUTPUT = os.path.join(ROOT, "Consolidated_Research_Dossier.docx")

# ---- palette (matches the source documents) --------------------------------
INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
SHADE_BLUE = "EAF1F7"

# ---- source documents, in the order they should appear ---------------------
# (filename, part-label, descriptive part title, collect_refs?)
PARTS = [
    ("Wireless_Perception_Reducing_Cameras.docx",
     "Part I", "Reducing Cameras with Wireless Perception", True),
    ("MetaAI_OverTheAir_Edge_AI.docx",
     "Part II", "MetaAI — Over-the-Air Edge AI", True),
    ("MetaPerception_Combining_MetaAI_with_WiFi_Sensing.docx",
     "Part III", "MetaPerception — Combining MetaAI with Wi-Fi Sensing", True),
    ("Camera_to_Wireless_Identity_Ideas.docx",
     "Part IV", "Camera-to-Wireless Identity — Idea Catalogue", True),
    ("Research_Ideas_LowCost_Wireless_Sensing.docx",
     "Part V", "Research Ideas — Low-Cost Wireless Sensing", True),
    ("ESPectre_Guide.docx",
     "Part VI", "ESPectre — Hardware and Tooling Guide", True),
    ("Weekly_Meeting_Update.docx",
     "Appendix A", "Weekly Meeting Update", False),
]


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def _remove(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def _is_heading(p, level=None):
    n = p.style.name if p.style else ""
    if level is None:
        return n.startswith("Heading")
    return n == f"Heading {level}"


def _is_toc(p):
    n = (p.style.name or "").lower()
    return n.startswith("toc") or "update field" in p.text.lower()


def strip_contents(doc):
    """Delete the 'Contents' heading and its TOC lines; keep real prose."""
    paras = list(doc.paragraphs)
    idx = None
    for i, p in enumerate(paras):
        if _is_heading(p, 1) and p.text.strip().lower() == "contents":
            idx = i
            break
    if idx is None:
        return
    _remove(paras[idx])
    for p in paras[idx + 1:]:
        if _is_heading(p):          # reached first real section -> stop
            break
        if _is_toc(p) or not p.text.strip():
            _remove(p)


REF_HEAD = re.compile(r"referen|further reading", re.I)


def strip_references(doc):
    """Delete a trailing References section; return its list items as text."""
    paras = list(doc.paragraphs)
    idx = None
    for i, p in enumerate(paras):
        if _is_heading(p, 1) and REF_HEAD.search(p.text or ""):
            idx = i
            break
    if idx is None:
        return []
    refs = []
    for p in paras[idx + 1:]:
        t = p.text.strip()
        if t and (p.style.name or "").startswith("List"):
            refs.append(t)
    for p in paras[idx:]:           # remove heading + everything after it
        _remove(p)
    return refs


def part_heading_texts(doc):
    """Top-level (Heading 1) section titles of a cleaned document."""
    out = []
    for p in doc.paragraphs:
        if _is_heading(p, 1) and p.text.strip():
            out.append(p.text.strip())
    return out


# ---------------------------------------------------------------------------
# base styles for the master document
# ---------------------------------------------------------------------------
def apply_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15
    for lvl, sz, col in [("Heading 1", 17, SLATE),
                         ("Heading 2", 13.5, SLATE),
                         ("Heading 3", 12, SLATE_LT)]:
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.color.rgb = col
        st.font.bold = True


def add_page_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldSimple")
    fld1.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld1)


def big_line(doc, text, size, color, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic
    return p


def add_rule(doc, color="9DB6CC"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# 1. load + clean every source document, gather outline + references
# ---------------------------------------------------------------------------
cleaned = []          # (part_label, part_title, Document, [heading1s])
all_refs = []

for fname, label, title, collect in PARTS:
    path = os.path.join(ROOT, fname)
    doc = Document(path)
    strip_contents(doc)
    refs = strip_references(doc)
    if collect:
        all_refs.extend(refs)
    cleaned.append((label, title, doc, part_heading_texts(doc)))


# dedupe references -----------------------------------------------------------
def norm(s):
    s = s.lower()
    s = re.sub(r"[^a-z0-9]", "", s)
    return s[:60]


seen = set()
merged_refs = []
for r in all_refs:
    if "companion" in r.lower() or "this document" in r.lower():
        continue                      # self-reference -> now part of the dossier
    k = norm(r)
    if k in seen:
        continue
    seen.add(k)
    merged_refs.append(r)


# ---------------------------------------------------------------------------
# 2. build the master front matter
# ---------------------------------------------------------------------------
master = Document()
apply_base_styles(master)
add_page_footer(master)

big_line(master, "Wireless Wave Perception & MetaPerception", 26, SLATE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
big_line(master, "A Consolidated Research Dossier", 16, SLATE_LT,
         bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_rule(master)
big_line(master,
         "Reducing cameras with Wi-Fi sensing  \u00b7  computing AI inside the "
         "wireless channel  \u00b7  camera-to-wireless identity  \u00b7  low-cost "
         "research directions  \u00b7  the ESPectre platform",
         11.5, INK, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
         italic=True)
big_line(master, "Multi-Modal Perception (MPL) programme",
         11, SLATE, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

# editor's note
big_line(master, "About this dossier", 13.5, SLATE, space_after=4)
note = master.add_paragraph()
note.add_run(
    "This single document consolidates seven working papers from the MPL "
    "programme into one arranged reference. The individual tables of contents "
    "and the separate, overlapping reference lists have been removed; every "
    "citation is merged once into the Consolidated References at the end. The "
    "papers are ordered so the argument builds from motivation, to enabling "
    "technology, to their combination, to concrete project ideas, and finally "
    "to the hardware used to prototype them."
)

# unified table of contents ---------------------------------------------------
page_break(master)
big_line(master, "Contents", 17, SLATE, space_after=6)
for label, title, doc, heads in cleaned:
    p = master.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}.  {title}")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = SLATE
    for h in heads:
        sub = master.add_paragraph()
        sub.paragraph_format.left_indent = Pt(18)
        sub.paragraph_format.space_after = Pt(0)
        rr = sub.add_run(h)
        rr.font.size = Pt(10.5)
        rr.font.color.rgb = INK
# reference entry in the TOC
p = master.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run("Consolidated References and Further Reading")
r.font.bold = True
r.font.size = Pt(12)
r.font.color.rgb = SLATE

# ---------------------------------------------------------------------------
# 3. compose: append each cleaned part behind a divider
# ---------------------------------------------------------------------------
composer = Composer(master)

for label, title, doc, heads in cleaned:
    # divider inside the master (goes to the current end of the body)
    page_break(master)
    big_line(master, label, 13, ACCENT, space_after=2)
    big_line(master, title, 22, SLATE, space_after=4)
    add_rule(master)
    composer.append(doc)

# ---------------------------------------------------------------------------
# 4. consolidated references
# ---------------------------------------------------------------------------
page_break(master)
big_line(master, "Consolidated References and Further Reading", 17, SLATE,
         space_after=4)
intro = master.add_paragraph()
intro.add_run(
    "Citations from the individual papers, merged and de-duplicated. "
    "Cross-references between the papers themselves have been dropped because "
    "they now live in this same dossier."
).italic = True
for r in merged_refs:
    master.add_paragraph(r, style="List Number")

# ---------------------------------------------------------------------------
master.save(OUTPUT)

print(f"Saved: {OUTPUT}")
print(f"Parts merged: {len(cleaned)}")
print(f"Merged references: {len(merged_refs)}")
