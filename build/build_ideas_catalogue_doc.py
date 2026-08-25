"""Build a Word document: a catalogue of concrete, low-cost, not-yet-done
research project ideas for a remote research intern, in the space where the
professor's direction sits (Wi-Fi CSI sensing, ESPectre, the reduce-cameras
idea, and the MetaAI over-the-air compute paper).

Each idea is written with What / Why (the gap) / How (minimal-cost method) /
Closest prior work / Equipment and cost / Deliverable and success metric /
Risk. All ideas are feasible with public datasets, simulation, one or two
ESP32 boards, a webcam, and household materials. No emojis.

Self-contained: generates its own figures with matplotlib, then builds .docx.
"""
import os
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle, Circle
from PIL import Image

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "assets_ideas_lowcost")
OUTPUT = os.path.join(os.path.dirname(HERE),
                      "Research_Ideas_LowCost_Wireless_Sensing.docx")
os.makedirs(ASSETS, exist_ok=True)

INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
PLUM = RGBColor(0x7A, 0x64, 0x9A)
ORANGE = RGBColor(0xB5, 0x72, 0x2F)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_PLUM = "F1EDF6"
SHADE_GREY = "F2F3F5"
SHADE_AMBER = "FAF0E6"
HEADER_FILL = "DCE6EF"

H_INK = "#33373d"; H_SLATE = "#4a6b8a"; H_SOFT = "#5b7fa6"
H_BLUE = "#9ecae1"; H_GREEN = "#a8d5a2"; H_ORANGE = "#f4b183"
H_PURPLE = "#c4b7e0"; H_PINK = "#e6a9b8"; H_GREY = "#d9dce1"; H_LGREY = "#eef1f4"

plt.rcParams.update({
    "font.family": "DejaVu Sans", "font.size": 11, "text.color": H_INK,
    "axes.edgecolor": "#b9bdc4", "axes.labelcolor": H_INK,
    "xtick.color": H_INK, "ytick.color": H_INK, "axes.titlecolor": H_INK,
    "figure.facecolor": "white", "axes.facecolor": "white",
})


def _save(fig, name):
    p = os.path.join(ASSETS, name)
    fig.savefig(p, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print("wrote", p)


def _box(ax, x, y, w, h, fc, ec=None, text="", fs=10, bold=False, tc=H_INK):
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.08",
                                fc=fc, ec=ec or fc, lw=1.4))
    if text:
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=fs, color=tc, fontweight="bold" if bold else "normal")


def _arrow(ax, p0, p1, color=H_SLATE, lw=2.0, style="-|>", ls="-"):
    ax.add_patch(FancyArrowPatch(p0, p1, arrowstyle=style, mutation_scale=15,
                                 color=color, lw=lw, linestyle=ls, shrinkA=2, shrinkB=2))


def fig_map():
    """Positioning map: cost (x) vs novelty (y), bubble size = feasibility for an intern."""
    fig, ax = plt.subplots(figsize=(9.2, 5.0))
    ideas = [
        ("1  OTA compute for\nthrough-wall CSI", 0.6, 4.4, 900, H_PURPLE),
        ("2  Robustness-training\ntransfer", 0.5, 3.4, 1000, H_GREEN),
        ("3  Poor-man's RIS\n(foil reflectors)", 1.7, 4.1, 780, H_ORANGE),
        ("4  Label-efficient\nself-supervised CSI", 0.7, 3.8, 820, H_BLUE),
        ("5  Camera-teaches-Wi-Fi\non one ESP32", 1.9, 3.6, 900, H_SLATE),
        ("6  CSI privacy audit\n+ suppression", 1.2, 4.0, 840, H_PINK),
    ]
    for lab, x, y, s, c in ideas:
        ax.scatter([x], [y], s=s, c=c, alpha=0.55, edgecolors="#5b6b7a", linewidths=1.3, zorder=3)
        ax.annotate(lab, (x, y), ha="center", va="center", fontsize=8.2, zorder=4)
    ax.set_xlim(0, 3.0); ax.set_ylim(2.8, 5.0)
    ax.set_xlabel("approx. hardware cost  (left = almost free)  \u2192", fontsize=10)
    ax.set_ylabel("degree of unexplored ground  \u2192", fontsize=10)
    ax.set_xticks([0.5, 1.5, 2.5]); ax.set_xticklabels(["~$0 (sim/data)", "~$10", "~$25"])
    ax.set_yticks([])
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_title("The six ideas, positioned by cost and novelty (bubble = intern feasibility)",
                 fontsize=10.5, pad=10)
    ax.text(1.5, 2.86, "All six sit in the low-cost, high-feasibility band \u2014 by design.",
            ha="center", fontsize=8.6, style="italic", color=H_SLATE)
    _save(fig, "fig_map.png")


def fig_toolkit():
    fig, ax = plt.subplots(figsize=(9.4, 4.6))
    ax.set_xlim(0, 12); ax.set_ylim(0, 6.0); ax.axis("off")
    ax.text(6.0, 5.65, "Everything a remote intern needs \u2014 nothing exotic",
            ha="center", fontsize=11.5, fontweight="bold", color=H_SLATE)
    _box(ax, 4.4, 2.55, 3.2, 1.0, H_LGREY, H_SLATE,
         "The minimal kit", 12, bold=True, tc=H_SLATE)
    ax.text(6.0, 2.2, "laptop  \u00b7  home router  \u00b7  1\u20132 ESP32 (~$8 each)\n"
                      "webcam  \u00b7  aluminium foil + cardboard  \u00b7  free datasets & Python",
            ha="center", fontsize=8.4, color=H_INK)
    items = [
        ("Public CSI\ndatasets", 0.7, 4.2, H_BLUE, "1, 2, 4, 6"),
        ("Simulation\n(Python)", 0.7, 0.9, H_PURPLE, "1, 3"),
        ("One ESP32\n+ router", 9.5, 4.2, H_GREEN, "2, 5, 6"),
        ("Webcam +\nfoil panels", 9.5, 0.9, H_ORANGE, "3, 5"),
    ]
    for lab, x, y, c, ideas in items:
        _box(ax, x, y, 1.9, 0.95, c, "#5b6b7a", lab, 8.6, bold=True)
        ax.text(x + 0.95, y - 0.28, f"enables ideas {ideas}", ha="center", fontsize=7.4,
                style="italic", color=H_SLATE)
        sx = x + 1.9 if x < 6 else x
        _arrow(ax, (sx, y + 0.48), (4.4 if x < 6 else 7.6, 3.05), color="#9aa6b2", lw=1.2, ls=(0, (3, 2)))
    _save(fig, "fig_toolkit.png")


def fig_pipeline():
    fig, ax = plt.subplots(figsize=(9.4, 3.2))
    ax.set_xlim(0, 12); ax.set_ylim(0, 3.6); ax.axis("off")
    y = 1.5
    steps = [
        ("Webcam +\nMediaPipe", H_ORANGE, "#b5722f"),
        ("Auto-label\n(pose/activity)", H_BLUE, H_SLATE),
        ("Sync with\nESP32 CSI", H_GREEN, "#4f7a53"),
        ("Train tiny\nCSI model", H_PURPLE, "#6b4f8a"),
        ("Remove camera,\nrun on CSI", H_GREY, H_SLATE),
    ]
    x = 0.3
    w = 2.05
    centers = []
    for lab, fc, ec in steps:
        _box(ax, x, y, w, 1.05, fc, ec, lab, 8.6, bold=True)
        centers.append(x + w)
        x += 2.3
    for cx in centers[:-1]:
        _arrow(ax, (cx, y + 0.52), (cx + 0.25, y + 0.52), color=H_SLATE)
    ax.text(6.0, 0.55, "A fully reproducible camera-supervised loop on ~$8 of hardware.",
            ha="center", fontsize=8.8, style="italic", color=H_SLATE)
    ax.text(6.0, 3.25, "Idea 5:  camera teaches Wi-Fi on a single ESP32",
            ha="center", fontsize=11, fontweight="bold", color=H_SLATE)
    _save(fig, "fig_pipeline.png")


def build_figures():
    fig_map()
    fig_toolkit()
    fig_pipeline()


build_figures()

# ===========================================================================
# DOCUMENT
# ===========================================================================
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = INK
pf = normal.paragraph_format; pf.space_after = Pt(8); pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), "8"); e.set(qn("w:color"), color)
    pbdr.append(e); pPr.append(pbdr)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def para(text, italic=False, size=11, align=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def field(label, text):
    """A labelled line: bold lead-in then body."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(label + "  ")
    r.bold = True; r.font.color.rgb = SLATE
    p.add_run(text)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def note_box(text, fill=SHADE_GREY, bc="C9CDD3", lead=None, lead_color=SLATE):
    p = doc.add_paragraph()
    _shade(p, fill); _border(p, color=bc)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.color.rgb = lead_color
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


_fig = {"n": 0}


def figure(img, caption, width=6.3):
    _fig["n"] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(width))
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {_fig['n']}.  "); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE
    r2 = cap.add_run(caption); r2.font.size = Pt(9.5); r2.font.color.rgb = INK
    cap.paragraph_format.space_after = Pt(12)


def make_table(headers, rows, widths=None, fontsize=9.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(fontsize + 0.5); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(val); r.font.size = Pt(fontsize); r.font.color.rgb = INK
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA")
                tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ---- TITLE ----------------------------------------------------------------
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Concrete, Low-Cost Research Ideas")
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(60); title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Six not-yet-done projects in Wi-Fi sensing and over-the-air computing, "
                "feasible for a remote research intern")
r.font.size = Pt(15); r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Each idea: what it is, why it matters, how to do it cheaply, the closest prior work, "
                 "and what you would deliver")
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = INK

line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  In the direction set by the MetaAI paper, the ESPectre repo, and the reduce-cameras idea  \u2014")
r.font.size = Pt(11); r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(22)

doc.add_page_break()

# ---- HOW TO READ ----------------------------------------------------------
h1("How to Read This, and How the Ideas Were Chosen")
para("The three-camera example the professor gave was exactly that, an example. This note collects six "
     "concrete project ideas in the same research direction (using ordinary radio to sense people, and "
     "pushing computation into the wireless channel), chosen against three hard constraints that fit a "
     "remote research intern:")
bullet("almost everything runs on a laptop, free public datasets, one or two ESP32 boards (about eight "
       "dollars each), a webcam you already own, and household materials.", bold_lead="Minimal funds and equipment: ")
bullet("each idea targets a real gap, something the literature has not yet done, so the work is publishable "
       "rather than a reproduction. Where close prior work exists, it is named honestly.", bold_lead="Not already done: ")
bullet("each has a clear method, a small deliverable, and a measurable success criterion, so progress can "
       "be shown at a weekly meeting.", bold_lead="Concrete and finishable: ")
figure("fig_map.png",
       "The six ideas positioned by hardware cost and by how much unexplored ground they cover; bubble size "
       "reflects how comfortably a remote intern can finish them. All sit in the low-cost, high-feasibility band.",
       width=6.4)
figure("fig_toolkit.png",
       "The complete equipment list and which ideas each item unlocks. Nothing exotic is required; several "
       "ideas need no hardware at all.", width=6.5)
note_box("A note on novelty claims: this field moves fast, so each idea lists the closest existing work and "
         "explains what is still open, rather than claiming nobody has touched the area. Do a fresh literature "
         "check before committing, but the gaps described here are, to the best of current knowledge, genuine.",
         fill=SHADE_AMBER, bc="D8B98A", lead="Read this first:  ", lead_color=ORANGE)

# ---------------------------------------------------------------------------
# IDEA TEMPLATE HELPER
# ---------------------------------------------------------------------------
def idea(num, name, tagline, what, why, how_items, prior, equip, deliver, risk, effort):
    h1(f"Idea {num}.  {name}")
    p = doc.add_paragraph()
    r = p.add_run(tagline); r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = SLATE_LT
    p.paragraph_format.space_after = Pt(8)
    field("What it is.", what)
    field("Why it matters (the gap).", why)
    ph = doc.add_paragraph(); rr = ph.add_run("How to do it, cheaply."); rr.bold = True; rr.font.color.rgb = SLATE
    ph.paragraph_format.space_after = Pt(3)
    for it in how_items:
        bullet(it)
    field("Closest prior work, and what is still open.", prior)
    field("Equipment and cost.", equip)
    field("Deliverable and success metric.", deliver)
    field("Main risk and mitigation.", risk)
    note_box(effort, fill=SHADE_BLUE, bc="9DB6CC", lead="Effort and level:  ")


# ---- IDEA 1 ---------------------------------------------------------------
idea(
    1,
    "Over-the-Air Compute for Through-Wall Human Sensing (Simulation)",
    "Take MetaAI's \u201ccompute in the channel\u201d idea and test it, for the first time, on through-wall "
    "human-activity radio data instead of line-of-sight images.",
    "MetaAI programs a metasurface to act as a neural-network layer so the wireless signal itself performs "
    "the classification. Their demonstration used line-of-sight image and gesture datasets. This project "
    "re-implements that computation model in software and evaluates it on public Wi-Fi CSI human-activity "
    "datasets (walking, sitting, falling), which are noisy, through-wall, and multipath-rich.",
    "The whole premise of the professor's direction is fusing sensing and computation. Nobody has yet asked "
    "whether MetaAI's linear-classifier-in-the-channel actually works on real human CSI, which is far harder "
    "than clean images. Answering that, positively or negatively, is a genuine contribution and directly "
    "extends both the paper and the reduce-cameras idea.",
    ["Download public CSI activity datasets (for example UT-HAR, NTU-Fi, or Widar3.0), all freely available.",
     "Write a small simulator of the MetaAI forward pass: send CSI features one per time slot, multiply each "
     "by a trained weight, accumulate, take the magnitude as the class score.",
     "Faithfully include the paper's hardware realities: 2-bit discrete weights, injected timing drift "
     "(Gamma-distributed shifts) and lowered SNR during training.",
     "Compare accuracy against an ordinary digital linear classifier and a small MLP baseline; report the gap.",
     "Ablate the number of meta-atoms, the bit-depth, and the timing/noise robustness, mirroring the paper."],
    "MetaAI (SIGCOMM 2025) itself, plus the many digital CSI activity-recognition papers. What is open: the "
    "intersection. No published work applies the metasurface-as-neural-network over-the-air model to "
    "through-wall human CSI, so both the benchmark and the released simulator would be new.",
    "None beyond a laptop. Pure simulation on free datasets. Effectively zero cost.",
    "An open-source simulator plus a benchmark table showing how close over-the-air linear inference gets to a "
    "digital classifier on CSI. Success = a clear, quantified answer with ablations (target: within a stated "
    "few percent of the digital baseline, or a clear explanation of why not).",
    "The concept may underperform on messy CSI. Mitigation: that is itself a publishable negative result with "
    "diagnosis, and the simulator remains a reusable community tool either way.",
    "Software only. Good first project. Strong paper potential because it bridges the two references directly.")

# ---- IDEA 2 ---------------------------------------------------------------
idea(
    2,
    "Do MetaAI's Training Tricks Fix Wi-Fi Sensing's Biggest Weakness?",
    "Borrow MetaAI's noise-aware and timing-robust training as plain data augmentation, and test whether they "
    "cure the well-known problem that Wi-Fi models break when the room changes.",
    "Wi-Fi sensing models are notoriously environment-specific: a model trained in room A often collapses in "
    "room B (domain shift). MetaAI introduced two training-time tricks, augmenting with lowered SNR and with "
    "random Gamma-distributed time shifts, to make its network robust. This project applies those exact "
    "augmentations to ordinary CSI activity classifiers and measures cross-room generalisation.",
    "Domain shift is the single biggest practical barrier to deploying Wi-Fi sensing, and it is exactly what "
    "the reduce-cameras idea flags as needing per-home calibration. If two cheap, purely-software "
    "augmentations measurably shrink the gap, that is immediately useful and cites the paper in a new light.",
    ["Use a public multi-environment CSI dataset (Widar3.0 is explicitly cross-domain), or collect a small "
     "two-room set with one ESP32.",
     "Train a baseline activity classifier; measure the accuracy drop from the training room to an unseen room.",
     "Add MetaAI-style SNR augmentation and Gamma-distributed temporal-shift augmentation during training only.",
     "Re-measure the cross-room drop; report how much of the gap each augmentation closes, alone and combined.",
     "Compare against standard augmentations (jitter, scaling) to isolate what is genuinely new."],
    "MetaAI (for the augmentations, but in a different setting) and the large cross-domain CSI literature "
    "(Widar, EI, domain-adaptation methods). What is open: nobody has repurposed MetaAI's specific over-the-air "
    "robustness augmentations as generic domain-shift remedies for through-wall CSI, nor benchmarked them "
    "against standard augmentations.",
    "Optional one ESP32 (about eight dollars) if collecting data; otherwise free with public datasets.",
    "A short empirical study with a clear before/after cross-room accuracy table. Success = a statistically "
    "meaningful reduction in the domain-shift gap, or a clean null result showing the tricks do not transfer.",
    "The tricks might not help outside their original setting. Mitigation: a rigorous negative result is still "
    "valuable, and the experiment is cheap and fast to run.",
    "Mostly software. Very achievable. High practical relevance, moderate novelty, low risk.")

# ---- IDEA 3 ---------------------------------------------------------------
idea(
    3,
    "The Poor-Man's Metasurface: Passive Foil Reflectors for Wi-Fi Coverage",
    "A programmable metasurface costs a lot. Test how far cheap static aluminium-foil reflector panels can push "
    "Wi-Fi sensing into dead rooms and around corners.",
    "A key limitation of single-link Wi-Fi sensing is that the signal may simply not reach a corner or a "
    "shadowed room. Expensive reconfigurable intelligent surfaces solve this actively. This project asks how "
    "much of that benefit is available for free using fixed, hand-made reflectors, panels of aluminium foil on "
    "cardboard, positioned to bounce the router's signal into the target area.",
    "It turns the abstract \u201cadd a metasurface\u201d recommendation into something an intern can actually "
    "test this week, and it produces practical placement guidance for the reduce-cameras deployment. Static "
    "passive reflectors for commodity-radio human sensing coverage are surprisingly under-studied compared to "
    "active surfaces.",
    ["Build several foil-on-cardboard panels of different sizes; they cost almost nothing.",
     "Set up one ESP32 CSI receiver in a room the router does not directly illuminate well.",
     "Measure motion-detection signal quality (CSI variance / detection reliability) with no reflector, then "
     "with reflectors at several positions and angles.",
     "Map which placements most improve the sensing signal, and quantify the gain over the no-reflector baseline.",
     "Optionally test a crude flat-panel corner reflector versus angled multi-panel arrangements."],
    "Active RIS-aided sensing (rich but expensive and simulation-heavy) and passive-reflector communication "
    "work. What is open: a systematic, measured study of cheap static reflectors specifically for commodity "
    "ESP32 CSI human-motion sensing coverage, with practical placement rules, is largely missing.",
    "One ESP32 (about eight dollars), aluminium foil, cardboard, tape. Under fifteen dollars total.",
    "A measurement report with a heat-map of reflector placements versus sensing-signal improvement, plus "
    "simple placement guidelines. Success = a repeatable, quantified coverage gain in a previously weak spot.",
    "Reflectors may give only marginal, position-sensitive gains. Mitigation: even a modest, well-characterised "
    "effect with clear guidance is a useful, honest, low-cost contribution.",
    "Hands-on hardware, but trivially cheap. Fun, visual, and a clear story for a meeting.")

# ---- IDEA 4 ---------------------------------------------------------------
idea(
    4,
    "Label-Efficient Wi-Fi Sensing with Self-Supervised Learning",
    "Cut the labelling burden: pretrain on unlabelled CSI, then fine-tune activity or people-counting with only "
    "a handful of labels.",
    "Every Wi-Fi sensing model needs labelled examples, and labelling is the main bottleneck (it is precisely "
    "why the reduce-cameras idea uses a camera as an automatic labeller). This project pretrains a CSI encoder "
    "with self-supervision, using temporal consistency and contrastive learning on unlabelled streams, then "
    "shows strong activity recognition or occupancy counting after fine-tuning on very few labels.",
    "If good accuracy is reachable with, say, one tenth of the usual labels, deployment in a new home becomes "
    "far cheaper, which is the whole point of a low-cost sensing system. Self-supervision for CSI, especially "
    "for people-counting and across domains, is still an open and active area.",
    ["Take a public CSI dataset and hide most labels to simulate a low-label regime.",
     "Pretrain an encoder with self-supervised objectives suited to CSI: contrastive pairs from augmentations, "
     "and predicting temporal order or masked segments.",
     "Fine-tune on 1, 5, and 10 percent of the labels; plot accuracy versus label budget against a "
     "trained-from-scratch baseline.",
     "Test whether the pretrained features also transfer to a second environment with few labels.",
     "Report the label budget at which self-supervision stops helping."],
    "A growing body of self-supervised CSI work exists (contrastive HAR, masked reconstruction). What is open: "
    "label-efficiency curves for people-counting specifically, and cross-domain transfer of self-supervised "
    "CSI features, remain under-explored and make a focused, defensible contribution.",
    "None beyond a laptop; free datasets. Zero cost. A GPU helps but Colab's free tier suffices.",
    "A label-efficiency study: accuracy-versus-labels curves showing how much labelling self-supervision saves. "
    "Success = matching the fully-supervised baseline with a small, clearly stated fraction of the labels.",
    "Self-supervised gains on small CSI datasets can be modest. Mitigation: even a clear characterisation of "
    "when it helps is useful; the study is cheap to iterate.",
    "Software only, more ML-heavy. Best if you are comfortable with PyTorch. Strong, current research angle.")

# ---- IDEA 5 ---------------------------------------------------------------
idea(
    5,
    "Camera-Teaches-Wi-Fi on a Single ESP32: A Reproducible Pipeline",
    "Build the actual camera-supervised loop from the reduce-cameras idea, end to end, on about eight dollars "
    "of hardware, and release it.",
    "The reduce-cameras idea proposes using a camera to auto-label CSI during a short enrolment, then removing "
    "the camera. This project implements that loop concretely on a single ESP32: a webcam runs free pose "
    "software (MediaPipe) to generate labels (present/absent, sitting/walking), synchronised with the ESP32 "
    "CSI stream, trains a small model, and then runs camera-free.",
    "Most camera-supervised Wi-Fi work used expensive specialist Wi-Fi cards (Intel 5300, Atheros). A clean, "
    "reproducible, open pipeline on the commodity ESP32, with a released dataset and code, would let anyone "
    "reproduce and build on the idea, which is a real contribution to the community and directly serves the "
    "professor's programme.",
    ["Flash one ESP32 with CSI-capable firmware (ESPectre or an ESP-IDF CSI example) and stream CSI to the laptop.",
     "Point a webcam at the same scene; run MediaPipe to auto-generate presence and simple activity labels.",
     "Timestamp-synchronise the two streams and build a paired dataset during a short enrolment.",
     "Train a small classifier from CSI to the camera labels; then evaluate with the camera switched off.",
     "Quantify how enrolment length and scene variety affect the camera-free accuracy."],
    "RF-Pose and DensePose-From-WiFi (camera-supervised, but heavy hardware and no reproducible ESP32 pipeline). "
    "What is open: an accessible, documented, open-source camera-to-ESP32-CSI distillation pipeline with a "
    "released dataset does not really exist yet.",
    "One ESP32 (about eight dollars) and a webcam you already own. Under ten dollars.",
    "A released dataset, code, and a report on camera-free accuracy versus enrolment effort. Success = a working "
    "camera-free classifier plus a reproducibility recipe others can follow.",
    "Time synchronisation and CSI noise on the ESP32 can be fiddly. Mitigation: start with coarse "
    "presence/absence labels, which are robust, before attempting finer activities.",
    "Hands-on, end-to-end systems work. Very tangible for weekly demos. Moderate build effort, high visibility.")

# ---- IDEA 6 ---------------------------------------------------------------
idea(
    6,
    "How Much Can Wi-Fi Really Identify You, and Can We Stop It Cheaply?",
    "Measure the privacy leak in commodity CSI, then test simple on-device tricks that keep activity accuracy "
    "while suppressing who-you-are.",
    "A 2026 result identified individuals from ordinary Wi-Fi with near-perfect accuracy, and MetaAI argues that "
    "computing in the channel keeps raw data private. This project empirically measures how much identity leaks "
    "from a cheap ESP32's CSI, then tests lightweight suppression, coarse quantisation, low-rank projection, or "
    "emitting only event labels, and reports the privacy-versus-utility trade-off.",
    "Privacy is the decisive factor for the sensitive rooms in the reduce-cameras idea (bedroom, bathroom). "
    "Turning the vague privacy claim into measured numbers, and offering a concrete cheap mitigation, is both "
    "socially useful and a clean research contribution.",
    ["Collect a small CSI set with one ESP32 and a few volunteers (or use a public identity-labelled CSI set).",
     "Train two probes: an activity classifier (utility) and a person-identifier (privacy risk); measure both.",
     "Apply simple, cheap transforms before classification: coarse quantisation, low-rank/PCA projection, "
     "temporal averaging, or emitting only the event label.",
     "Plot the trade-off: how much identity accuracy each transform removes versus how much activity accuracy "
     "it costs.",
     "Recommend an operating point that keeps activity usable while making identification unreliable."],
    "The KIT identification result and general CSI privacy discussion. What is open: an empirical "
    "privacy-utility trade-off curve on commodity ESP32 CSI, with concrete low-cost suppression mechanisms, is "
    "not established.",
    "One ESP32 (about eight dollars); optionally a public dataset. Under ten dollars.",
    "A privacy-utility trade-off study and a recommended cheap suppression setting. Success = a transform that "
    "substantially drops identification accuracy while keeping activity recognition usable.",
    "Collecting multi-person data remotely can be hard. Mitigation: begin with a public identity-labelled CSI "
    "dataset, then validate on a tiny self-collected set.",
    "Software-led with light data collection. Timely, responsible-AI angle. Moderate difficulty, high relevance.")

# ---- COMPARISON TABLE -----------------------------------------------------
h1("At a Glance: Comparing the Six Ideas")
para("A quick way to pick, depending on your comfort with hardware versus machine learning, and how much "
     "novelty versus safety you want.")
make_table(
    ["Idea", "Cost", "Hardware?", "Skill lean", "Novelty", "Risk"],
    [
        ["1  OTA compute for CSI", "~$0", "No (sim)", "Simulation", "High", "Medium"],
        ["2  Robustness-training transfer", "~$0\u2013$8", "Optional", "ML training", "Medium", "Low"],
        ["3  Poor-man's RIS reflectors", "<$15", "Yes", "Measurement", "Medium-High", "Medium"],
        ["4  Self-supervised CSI", "~$0", "No", "ML (deeper)", "Medium", "Medium"],
        ["5  Camera-teaches-Wi-Fi ESP32", "<$10", "Yes", "Systems", "Medium-High", "Medium"],
        ["6  CSI privacy audit", "<$10", "Light", "ML + analysis", "Medium-High", "Medium"],
    ],
    widths=[2.3, 0.9, 0.95, 1.15, 1.05, 0.85], fontsize=9.0)
figure("fig_pipeline.png",
       "Idea 5 in one line: the camera-supervised loop that turns about eight dollars of hardware into a "
       "camera-free Wi-Fi classifier, fully reproducible.", width=6.5)

# ---- RECOMMENDATION -------------------------------------------------------
h1("A Suggested Way to Start")
para("If you want the fastest path to a result you can show, and the lowest risk:")
bullet("it needs no hardware, uses free datasets, and produces a reusable tool plus a clear before/after "
       "number in a couple of weeks.", bold_lead="Begin with Idea 2 (robustness-training transfer): ")
bullet("in parallel, order one ESP32 (about eight dollars) and start Idea 5 (camera-teaches-Wi-Fi), which "
       "gives you real data and a tangible demo for meetings.", bold_lead="Then add hardware: ")
bullet("once you have data and a simulator, Idea 1 (over-the-air compute on CSI) becomes the strongest paper, "
       "because it directly unites the professor's two references.", bold_lead="Aim high: ")
note_box("All six ideas share the same cheap toolkit, so you are never locked in: start with the software ones, "
         "add the eight-dollar ESP32 when ready, and let the results steer which idea becomes the main project.",
         fill=SHADE_GREEN, bc="AFC8B2", lead="Bottom line:  ", lead_color=ACCENT)

# ---- REFERENCES -----------------------------------------------------------
h1("Pointers and Free Resources")
for rtext in [
    "MetaAI: C. Feng et al., Enabling Over-the-Air AI for Edge Computing via Metasurface-Driven Physical "
    "Neural Networks, ACM SIGCOMM 2025 (the over-the-air compute reference).",
    "ESPectre: open-source Wi-Fi CSI motion detection for Home Assistant, github.com/francescopace/espectre "
    "(the low-cost ESP32 sensing platform).",
    "Public CSI datasets: UT-HAR, NTU-Fi, Widar3.0 (cross-domain gestures), SignFi, and the DensePose-From-WiFi "
    "release \u2014 all free for research.",
    "ESP32 CSI: the ESP-IDF Wi-Fi CSI examples and the ESP32-CSI-Tool make streaming CSI from an eight-dollar "
    "board straightforward.",
    "MediaPipe (Google): free, on-device pose and person detection, ideal as the automatic camera labeller in "
    "Idea 5.",
    "Background on the direction: RF-Pose (CVPR 2018), Person-in-WiFi (ICCV 2019), DensePose-From-WiFi (2023), "
    "and the Widar line for cross-domain sensing.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(rtext); r.font.size = Pt(10)

closing = doc.add_paragraph(); closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("Prepared as an idea catalogue for a remote research internship. Cost figures are approximate "
                    "and reflect commodity ESP32 pricing and free/open datasets and tools. Novelty notes name the "
                    "closest known prior work; a fresh literature check is recommended before committing to any one "
                    "idea.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE_LT

if __name__ == "__main__":
    doc.save(OUTPUT)
    print("saved", OUTPUT)
