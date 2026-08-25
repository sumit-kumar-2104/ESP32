"""Build the single CONSOLIDATED study guide.

This is not a merge of the other .docx files. It is a freshly written, coherent
document that folds the content of all the scattered documents (wireless
perception survey, MetaAI explainer, MetaPerception combination study, idea
catalogues, ESP32 work) into one non-redundant study guide, organised around the
professor's two directions:
  (1) reproduce MetaAI and take it forward for MetaPerception / wireless tasks;
  (2) run ESP32 Wi-Fi perception.

Real citations only. No emojis.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
OUTPUT = os.path.join(os.path.dirname(HERE), "Consolidated_Study_Guide.docx")

INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
PLUM = RGBColor(0x7A, 0x64, 0x9A)
RUST = RGBColor(0xA8, 0x5A, 0x4A)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_PLUM = "F1EDF6"
SHADE_AMBER = "FBF2E7"
SHADE_GREY = "F2F3F5"
HEADER_FILL = "DCE6EF"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = INK
pf = normal.paragraph_format; pf.space_after = Pt(8); pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 18, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(size)); e.set(qn("w:space"), "8"); e.set(qn("w:color"), color)
    pbdr.append(e); pPr.append(pbdr)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def para(text, italic=False, size=11, space_after=8):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def field(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + "  "); r.bold = True; r.font.color.rgb = SLATE
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def layman_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREEN); _border(p, color="AFC8B2")
    r = p.add_run("In plain terms:  "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def note_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREY); _border(p, color="C9CDD3")
    r = p.add_run(text); r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


def caution_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_AMBER); _border(p, color="D9B98A")
    r = p.add_run("Reality check:  "); r.bold = True; r.font.color.rgb = RUST
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def key_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_BLUE); _border(p, color="9DB6CC")
    r = p.add_run("Key point:  "); r.bold = True; r.font.color.rgb = SLATE
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def make_table(headers, rows, widths=None, fs=9.5, first_bold=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""; p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL); tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; p = cells[i].paragraphs[0]
            r = p.add_run(val); r.font.size = Pt(fs); r.font.color.rgb = INK
            if i == 0 and first_bold:
                r.bold = True; r.font.color.rgb = SLATE
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA"); tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def toc():
    toc_p = doc.add_paragraph(); run = toc_p.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    t_run = OxmlElement("w:t"); t_run.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for e in (fldChar, instr, fldChar2, t_run, fldChar3): run._r.append(e)


# ===========================================================================
# TITLE
# ===========================================================================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Wireless Perception and In-Air Computing"); r.font.size = Pt(27); r.font.bold = True; r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(54); title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Consolidated Study Guide"); r.font.size = Pt(17); r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(16)

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("One place to understand the whole project: how people sense humans with radio, how the "
                 "MetaAI paper makes the wireless channel compute a neural network, how the two could meet "
                 "(MetaPerception), a concrete plan to reproduce and extend MetaAI, and the hands-on ESP32 "
                 "Wi-Fi sensing track. Written from scratch, no redundancy, real citations only.")
r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = INK

line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  Multi-Modal Perception (MPL) programme  \u2014"); r.font.size = Pt(11); r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(22)
doc.add_page_break()

h1("Contents"); toc(); doc.add_page_break()

# ===========================================================================
# 0. ORIENTATION
# ===========================================================================
h1("0.  Orientation: The Two Directions")
para("This guide serves the two directions set in the meetings. Everything is arranged so you can read "
     "top to bottom once and understand the whole picture, then dip back into any section.")
bullet("reproduce the MetaAI paper (making the wireless channel itself run a neural network) and take the "
       "idea forward for MetaPerception and wireless sensing tasks \u2014 what to build, what papers to use, "
       "what similar ideas exist, and the next steps.", "Direction 1 (the main focus now): ")
bullet("run ESP32 Wi-Fi perception on real hardware \u2014 sense people through Wi-Fi with cheap boards, "
       "collect data, and build toward camera-taught sensing.", "Direction 2 (hands-on): ")
key_box("The unifying thread of the whole project: a neural network is mostly multiply-and-add, and a "
        "radio channel already multiplies and adds signals. Wireless perception reads what the channel "
        "did to infer people; MetaAI programs what the channel does so the channel computes the network. "
        "MetaPerception is the idea of doing both on the same hardware.")
para("How the parts fit together:")
make_table(
    ["Part", "What it covers", "Serves"],
    [
        ["1", "Foundations of wireless perception (how radio senses people)", "Both"],
        ["2", "The wireless-perception literature, grouped by task", "Direction 2 + context"],
        ["3", "The MetaAI paper, explained fully but concisely", "Direction 1 (core)"],
        ["4", "Similar ideas and the real lineage MetaAI builds on", "Direction 1"],
        ["5", "MetaPerception: combining MetaAI with Wi-Fi sensing", "Direction 1"],
        ["6", "Reproduction plan, feasibility, hardware, next steps", "Direction 1 (action)"],
        ["7", "Project idea catalogue and how to choose", "Both"],
        ["8", "The ESP32 practical track and first results", "Direction 2 (action)"],
        ["9", "Consolidated references", "Both"],
    ],
    widths=[0.6, 4.3, 1.6],
)

# ===========================================================================
# 1. FOUNDATIONS
# ===========================================================================
h1("1.  Foundations of Wireless Perception")
para("Three facts make everything else easy to follow.")
bullet("radio waves from a router or radar travel through a room, reflect off walls, furniture and "
       "human bodies, and reach a receiver slightly changed. A moving body changes those reflections in "
       "a repeatable way.", "Radio reflects off people. ")
bullet("Wi-Fi chips can report Channel State Information (CSI): a detailed measure of how the signal's "
       "amplitude and phase changed across many frequencies. CSI is the raw material of Wi-Fi sensing "
       "(older work used coarse signal strength, RSSI, which carries much less information).",
       "Wi-Fi already measures the channel. ")
bullet("if a moving body changes the radio in a repeatable way, a neural network can be trained to read "
       "those changes and output presence, a gesture, a pose, an activity, or an identity.",
       "Machine learning reads the pattern. ")
layman_box("Picture the room as a still pond and the router as a hand tapping the water. A person moving "
           "makes ripples; the receiver watches the ripples and a neural network learns to say what "
           "caused them. No camera, works in the dark and through walls.")
key_box("A recurring trick ties this to the whole project: use a camera during training only, to "
        "automatically label what the radio is seeing, then remove the camera and let the radio work "
        "alone. This is called cross-modal supervision (camera teaches radio) and it is the backbone of "
        "the \u201creduce the cameras\u201d idea.")

# ===========================================================================
# 2. WIRELESS PERCEPTION LITERATURE
# ===========================================================================
h1("2.  The Wireless-Perception Literature, By Task")
para("These are the landmark works, arranged from easier tasks to harder ones. Each entry is deliberately "
     "compact; the companion survey document has a full page on each. Together they form a ladder from "
     "presence to identity.")

h2("2.1  Presence, motion and gesture (the easy, reliable base)")
field("Widar3.0 (MobiSys 2019).", "Recognises hand gestures from commodity Wi-Fi and, crucially, keeps "
      "working in rooms and orientations it never trained on. Its trick is the Body-coordinate Velocity "
      "Profile (BVP), a representation of how the body moves rather than how the signal looked in one "
      "room; motion is the same everywhere, so the model generalises. It also gave the field a widely "
      "reused public dataset. Read it to understand the central problem of the whole area: generalisation "
      "across environments.")
field("SenseFi (Patterns, 2023).", "Not a new trick but an open-source library and benchmark comparing "
      "many deep models (CNNs, RNNs, Transformers) across public CSI datasets for activity, gesture and "
      "identity. It is the fastest way to get baselines and hands-on code. Start any Wi-Fi sensing work "
      "here.")

h2("2.2  Pose and body shape (camera-taught radio)")
field("RF-Pose (CVPR 2018, MIT).", "Draws a 2D skeleton of a person, even through a wall, from radio "
      "alone. A camera watches the same people during training and produces skeletons; the radio network "
      "is trained to output the same skeletons (teacher-student). After training the camera is removed. "
      "This is the flagship proof that radio can do a rich vision task, and it popularised the "
      "camera-teaches-radio recipe. It used a custom FMCW radio, not commodity Wi-Fi.")
field("Person-in-WiFi (ICCV 2019).", "Does body segmentation and skeleton keypoints from just a few "
      "commodity Wi-Fi antennas, again supervised by a camera. It is the bridge from expensive radar to "
      "the cheap Wi-Fi you can actually deploy.")
field("DensePose From WiFi (arXiv 2023, CMU).", "The most detailed Wi-Fi body sensing so far: it maps "
      "Wi-Fi to a dense body surface (24 regions), the same output as camera-based DensePose. It is the "
      "clearest evidence of how much visual body information survives in ordinary Wi-Fi, motivated "
      "explicitly by privacy and low cost.")

h2("2.3  Identity (the hard, sensitive task)")
field("XModal-ID (MobiCom 2019, UCSB).", "Decides whether the person sensed by Wi-Fi behind a wall is "
      "the same person shown in a piece of video. It turns the video into the Wi-Fi signature that walk "
      "would produce and compares it with the measured Wi-Fi. This is the canonical demonstration that a "
      "camera identity can be handed to Wi-Fi through a wall \u2014 the scientific backbone of the "
      "camera-enrol, wireless-track idea. It works for a small candidate set, mainly while walking.")
field("RF-ReID (CVPR 2020, MIT).", "Re-identifies the same person over days using body shape and gait "
      "from radar, so it survives clothing changes that defeat camera re-identification. It builds a "
      "per-person signature with metric learning and discourages the model from latching onto the room. "
      "Strongest evidence that radio identity can be robust and long-term; uses a specialised FMCW radar.")
field("WhoFi (arXiv 2025).", "The freshest identity-from-Wi-Fi work: a Transformer encoder turns CSI "
      "into a biometric signature, trained with a contrastive loss, evaluated on the public NTU-Fi "
      "dataset with results competitive with prior methods. A recent preprint, so treat exact numbers as "
      "not yet peer-reviewed; useful as an up-to-date reference and for the privacy framing.")

h2("2.4  Handoff across sensors (scaling to corridors and buildings)")
field("Vi-Fi (IPSN 2022, Rutgers).", "Matches people seen by a camera with their phones' wireless (FTM) "
      "and motion (IMU) data, linking a camera track and a wireless track to the same person. It is the "
      "cleanest example of vision-to-wireless handoff at scale, with a public multimodal dataset. It ties "
      "identity to a carried phone, and modern MAC-address randomisation weakens passive phone "
      "identification, so it usually needs a consented app or ranging.")

note_box("Method families you meet across these papers: CNNs on CSI spectrograms; recurrent/Transformer "
         "encoders for time series; teacher-student cross-modal supervision; and metric/contrastive "
         "learning for identity. The recurring weakness is generalisation to new environments; the main "
         "cures are domain-independent features (like BVP), data from many rooms, and domain adaptation.")

h2("2.5  Datasets and capture tools you can actually use")
make_table(
    ["Resource", "Type", "What it gives you"],
    [
        ["Widar3.0", "Dataset", "Cross-domain gesture CSI across rooms/orientations; the generalisation test."],
        ["NTU-Fi", "Dataset", "CSI for identification and activity; used by WhoFi and SenseFi."],
        ["UT-HAR", "Dataset", "Early public CSI activity set; a common baseline."],
        ["Vi-Fi dataset", "Dataset", "Synced vision + phone wireless (FTM) + IMU for cross-modal association."],
        ["SenseFi", "Code/benchmark", "Ready models and loaders for the datasets above; your starting point."],
        ["Intel 5300 / Atheros CSI Tool", "Capture", "Classic laptop/PC CSI capture used by much prior work."],
        ["Nexmon CSI", "Capture", "CSI on Broadcom chips (some phones, Raspberry Pi)."],
        ["ESP32 CSI Toolkit / ESPectre", "Capture", "Cheapest CSI entry; matches your current hardware."],
    ],
    widths=[1.7, 1.1, 3.8],
)

# ===========================================================================
# 3. METAAI
# ===========================================================================
h1("3.  MetaAI: Computing the Network Inside the Channel")
para("This is the paper the professor asked you to focus on: Feng et al., \u201cEnabling Over-the-Air AI for "
     "Edge Computing via Metasurface-Driven Physical Neural Networks,\u201d ACM SIGCOMM 2025. This section "
     "explains it fully but without padding.")

h2("3.1  The core idea")
para("A neural-network layer is multiply-and-add: multiply each input by a learned weight, add the "
     "results, and read off a score per class. A wireless channel already multiplies a signal (it scales "
     "and rotates it) and adds signals that arrive together. So if you can set exactly how the channel "
     "multiplies the signal at each instant, the channel itself computes the network as the signal passes "
     "through it. A metasurface \u2014 a flat panel of hundreds of tiny, electronically controllable cells "
     "(meta-atoms) \u2014 is the device that lets you set it.")
layman_box("Normally a sensor mails raw data to a server and the server thinks. MetaAI installs a "
           "programmable smart mirror in the room; as the sensor's signal bounces off it on the way to the "
           "receiver, the mirror reshapes it so the receiver gets the answer (\u201cthis is a cat\u201d) instead "
           "of the raw picture. The mirror can be reprogrammed to hold whatever the model learned.")

h2("3.2  The two problems it had to solve")
bullet("wireless sends data one piece (symbol) at a time, but a network wants all inputs at once. "
       "Because a linear layer is a sum of products, it can be computed one term at a time and "
       "accumulated \u2014 the final total is identical. So the naturally sequential radio link is a fine way "
       "to compute a linear layer.", "Sequential vs. parallel: ")
bullet("the multiply must happen in the wave. The received signal equals the channel response times the "
       "transmitted signal; by programming the metasurface so its channel response equals the trained "
       "weight at each time slot, each input is multiplied by exactly the weight the network wants. The "
       "receiver accumulates the products; the magnitude of the complex sum is the class score and acts "
       "as the nonlinearity.", "Multiplication in the wave: ")
key_box("Because each input gets its own time slot, one metasurface can give each input an independent "
        "weight. That is why a single-layer, single-surface design does what older stacked physical "
        "networks needed many surfaces to approximate.")

h2("3.3  Making it work in a real room")
para("A clean equation assumes an ideal world; a real room adds echoes, timing drift and noise. MetaAI "
     "adds a specific, clever fix for each \u2014 and these fixes are the most reusable lessons for any "
     "physical-layer sensing project:")
bullet("room reflections would corrupt the result. Communication symbols average to zero over their "
       "period, and so do the echoes. MetaAI deliberately varies the surface within a symbol so only its "
       "own signal does not average to zero \u2014 the echoes cancel when summed, its signal survives.",
       "Multipath cancellation: ")
bullet("the transmitter and surface do not share a clock, so inputs can drift out of step with weights. "
       "CDFA (Coarse-grained Detection, Fine-grained Adjustment) uses a low-power energy detector to "
       "start the weight sequence roughly on time, then absorbs the remaining error by training on "
       "randomly time-shifted data. No shared clock needed.", "Timing (CDFA): ")
bullet("hardware and environmental noise are handled by simply training at a lower signal-to-noise "
       "ratio, so the model expects noise and copes at run time.", "Noise-aware training: ")
bullet("computing one class score needs one pass; several classes are computed in parallel across "
       "subcarriers (frequency slots) or receive antennas via one joint optimisation, trading a little "
       "accuracy for much lower latency.", "Parallelism: ")
bullet("several sensors (camera views, or accelerometer plus gyroscope) are processed in separate time "
       "windows through the same shared surface and combined by late fusion, which raises accuracy.",
       "Multi-sensor fusion: ")

h2("3.4  Hardware and results")
make_table(
    ["Component", "What the authors used"],
    [
        ["Metasurfaces", "Two prototypes: dual-band (2.4/5 GHz) and single-band (3.5 GHz)."],
        ["Meta-atom grid", "16 x 16 = 256 cells per surface."],
        ["Per-cell control", "PIN diodes giving 4 phase states (2-bit)."],
        ["Controller / radios", "STM32 microcontroller; USRP X310 software-defined radios (Tx and Rx)."],
        ["Test environments", "Corridor, laboratory, office; including non-line-of-sight and cross-room."],
    ],
    widths=[1.9, 4.4],
)
para("Reported results: about 82.8 percent average accuracy (up to about 89.8 percent) across six "
     "image/gesture datasets with a single linear layer; within roughly 7 percent of its own digital "
     "simulation; multi-sensor fusion adding up to about 27 percent; CDFA lifting accuracy from about 19 "
     "percent (no sync) to about 89 percent; and operation across 2.4, 3.5 and 5 GHz bands and several "
     "distances and angles.")
caution_box("MetaAI implements only linear networks (no deep nonlinear models yet). Larger models mean "
            "more transmissions and higher latency. Accuracy is capped by metasurface resolution (cell "
            "count and 2-bit depth). Movement forces recalibration. It is a lab prototype with custom "
            "metasurfaces and software-defined radios \u2014 not something an ESP32 can do today.")

# ===========================================================================
# 4. LINEAGE
# ===========================================================================
h1("4.  Similar Ideas and the Real Lineage")
para("You asked what work exists on this idea and what is similar. The honest picture: MetaAI (2025) is "
     "very recent, so there is not yet a large set of papers extending it specifically. But it sits on a "
     "deep, real lineage that is exactly what you should read to understand the field.")
field("Diffractive Deep Neural Networks (D2NN) \u2014 Lin et al., Science 2018.", "The founding demonstration "
      "that a wave passing through fixed patterned layers performs a neural network. Light through "
      "3D-printed plates classifies digits and fashion images, all optically. This is the intellectual "
      "root of physical neural networks, including MetaAI's microwave version. The layers are fixed once "
      "printed.")
field("Coherent nanophotonic circuits \u2014 Shen et al., Nature Photonics 2017.", "A photonic chip that does "
      "a network's matrix multiplications with light through interferometers \u2014 the reconfigurable, "
      "chip-scale sibling of D2NN. It shows the weights can be made tunable, not just fixed. Photonics, "
      "not radio, and small-scale.")
field("Programmable metasurface D2NN \u2014 Liu et al., Nature Electronics 2022 (Cui group).", "The closest "
      "precursor to MetaAI: a microwave metasurface that performs a diffractive network and can be "
      "reprogrammed on the fly. It proves the two ingredients MetaAI needs \u2014 microwave operation and live "
      "reprogramming of a surface into weights \u2014 already work. MetaAI's step is doing this over an "
      "ordinary communication link with a commodity transmitter.")
field("AirNN \u2014 Sanchez et al., 2022 (Northeastern).", "The most direct predecessor: it performs a "
      "convolution (the core CNN operation) over the air using reconfigurable intelligent surfaces, "
      "demonstrated experimentally with the full network validated in simulation. It is the bridge from "
      "\u201csurfaces can shape signals\u201d to \u201csurfaces can compute a network operation.\u201d")
field("Over-the-air computation (AirComp) \u2014 e.g. Yang et al., IEEE TWC 2020.", "Uses the natural "
      "adding-up of overlapping radio signals to compute a sum directly in the air, used to speed up "
      "federated-learning aggregation. It is the \u201caddition for free\u201d half of the story; MetaAI adds "
      "programmable multiplication so the channel can do a full layer, not just a sum.")
key_box("The accurate framing for the professor: MetaAI is a fresh capstone on a mature line of "
        "optical / metasurface / over-the-air computing. The opportunity is not to chase a crowded "
        "follow-up field, but to connect this computation idea to the wireless-perception tasks in "
        "Part 2 \u2014 which is exactly what MetaPerception (Part 5) proposes.")

# ===========================================================================
# 5. METAPERCEPTION
# ===========================================================================
h1("5.  MetaPerception: Combining MetaAI with Wi-Fi Sensing")
para("MetaPerception is the project's own thesis: the same physics that carries Wi-Fi sensing also "
     "carries the neural-network math, so a metasurface could improve sensing in three ways at once \u2014 "
     "richer perception, cheaper devices, and better privacy. This section states the concept honestly, "
     "including where it is proven and where it is still speculative.")

h2("5.1  Why the two fit together")
para("Wi-Fi sensing reads what the channel did to a signal to infer people. MetaAI programs what the "
     "channel does so the channel computes a network. Both rest on the same linear physics. So a "
     "programmable surface in the room could, in principle, both shape the illumination for better "
     "sensing and carry part of the inference \u2014 letting a tiny node just transmit while the channel does "
     "the first layer of the sensing model.")

h2("5.2  A staged roadmap (each stage is useful on its own)")
make_table(
    ["Phase", "What you deploy", "Status"],
    [
        ["0. ESPectre now", "One camera plus ESP32 CSI nodes; presence, motion, which-room.", "Working today"],
        ["1. Camera-taught Wi-Fi", "Cross-modal distillation adds activity and fall classes.", "Proven in prior work"],
        ["2. Metasurface coverage", "A surface steers signal into dead rooms and raises degrees of freedom.", "Plausible, needs RF gear"],
        ["3. Over-the-air inference", "The channel computes the classifier; nodes only transmit.", "Research frontier"],
    ],
    widths=[1.6, 3.4, 1.5],
)
caution_box("Be honest in the meeting: Phases 0 and 1 are grounded in existing, working systems. Phases 2 "
            "and 3 are research directions that need a real metasurface and software-defined radios, and "
            "combining sensing with over-the-air inference on one surface has not been demonstrated yet. "
            "That gap is the opportunity, not a solved result.")

# ===========================================================================
# 6. REPRODUCTION PLAN
# ===========================================================================
h1("6.  Reproducing MetaAI and Taking It Forward: The Plan")
para("This is the action part for Direction 1. It answers: how would we reproduce the work, what does it "
     "need, and what are the next steps \u2014 at three levels of ambition, honestly matched to the hardware "
     "you actually have.")

h2("6.1  What full reproduction needs")
make_table(
    ["Component", "Needed", "Accessibility for you"],
    [
        ["Reconfigurable metasurface", "16x16, 2-bit, PIN-diode controlled", "Hard: custom-fabricated; not off-the-shelf"],
        ["Software-defined radios", "USRP X310 (Tx and Rx)", "Expensive; sometimes available in RF labs"],
        ["Controller", "STM32 + shift registers", "Easy once the panel exists"],
        ["Compute", "GPU/PC for offline training", "Easy"],
        ["Expertise", "RF/antenna, SDR, metasurface control", "This is an RF-systems project, not just ML"],
    ],
    widths=[1.9, 2.1, 2.3],
)
caution_box("Reproducing MetaAI end-to-end is a serious hardware-lab effort. It is realistic only if the "
            "lab has, or will fund, a reconfigurable metasurface and SDRs. It is not an ESP32-scale build.")

h2("6.2  Three realistic paths")
h3("Path A \u2014 Faithful simulator (start immediately, no special hardware)")
numbered("reimplement the MetaAI pipeline in software: model the channel-as-weights, the 2-bit "
         "metasurface quantisation, and the impairments (timing drift, multipath, noise).", "")
numbered("reproduce the paper's key curves \u2014 accuracy versus timing error, versus number of meta-atoms, "
         "versus SNR \u2014 and confirm the fixes (CDFA, noise-aware training) behave as reported.", "")
numbered("write it up as a reproduction study. This alone is a strong, honest contribution and teaches "
         "the whole idea deeply.", "")
para("Needs only a laptop and GPU. This is the recommended first step for Direction 1.")

h3("Path B \u2014 Scaled hardware demo (only if the lab has RF gear)")
para("With a commercial reconfigurable intelligent surface (RIS) development kit and one or two SDRs, "
     "reproduce a single over-the-air linear classifier on a small dataset, and study one specific "
     "robustness issue (for example, how timing error degrades accuracy and how the training-time fix "
     "helps). A focused, publishable contribution without inventing hardware.")

h3("Path C \u2014 Computation-meets-perception study (the novel angle)")
para("In simulation, study whether an over-the-air linear front-end could pre-process CSI sensing "
     "features before a tiny node transmits \u2014 i.e., combine the MetaAI computation idea with the Wi-Fi "
     "sensing task (MetaPerception, Phase 3). Framed honestly as an exploratory feasibility study, this "
     "is the distinctive research narrative that ties the whole project together.")

h2("6.3  Papers to use, by purpose")
make_table(
    ["Purpose", "Papers to read/use"],
    [
        ["Understand the target", "MetaAI (SIGCOMM 2025)."],
        ["Understand the lineage", "AirNN (2022); D2NN (2018); Shen et al. (2017); Liu et al. metasurface D2NN (2022); AirComp (Yang et al. 2020)."],
        ["Ground the sensing side", "Widar3.0; RF-Pose; Person-in-WiFi; DensePose-from-WiFi; XModal-ID; RF-ReID; WhoFi."],
        ["Get code and baselines", "SenseFi library and benchmark; ESPectre / ESP32 CSI Toolkit."],
        ["Datasets", "Widar3.0, NTU-Fi, UT-HAR, Vi-Fi."],
    ],
    widths=[1.8, 4.5],
)

h2("6.4  Recommended sequence")
numbered("keep building ESP32 Wi-Fi sensing (Part 8) for real data and demos \u2014 you are already doing "
         "this.", "")
numbered("in parallel, build the MetaAI simulator (Path A) to master the idea cheaply.", "")
numbered("let the computation-meets-perception study (Path C) be the ambitious research thread.", "")
numbered("pursue the hardware demo (Path B) only when and if the lab commits a metasurface and SDRs.", "")

# ===========================================================================
# 7. IDEA CATALOGUE
# ===========================================================================
h1("7.  Project Idea Catalogue and How to Choose")
para("A condensed, de-duplicated shortlist drawn from the fuller catalogues. Each is a concrete option "
     "you could take to the professor; the scorecard columns are relative guidance, not measurements.")
make_table(
    ["#", "Idea", "Effort", "Cost", "Novelty", "Best used as"],
    [
        ["1", "MetaAI faithful simulator + reproduction study", "Med", "Low", "Med", "First Direction-1 result"],
        ["2", "ESP32 Wi-Fi gait/zone sensing (household)", "Low-Med", "Low", "Med", "Real data engine (Direction 2)"],
        ["3", "Camera-taught Wi-Fi activity/fall (distillation)", "Med", "Low-Med", "Med", "Capability add-on"],
        ["4", "Over-the-air front-end for CSI (MetaPerception)", "Med-High", "Low (sim)", "High", "Ambitious thesis narrative"],
        ["5", "Scaled RIS + SDR linear classifier", "High", "Med-High", "High", "If lab funds RF gear"],
        ["6", "Vision-to-wireless identity handoff (corridors)", "Med", "Low-Med", "Med-High", "Scale-up / networks"],
        ["7", "Privacy-preserving consented re-ID", "Med", "Low", "High", "Distinctive framing"],
    ],
    widths=[0.4, 2.7, 0.8, 0.8, 0.8, 1.8],
)
para("Recommendation: pair Idea 1 (learn and reproduce MetaAI) with Idea 2 (real ESP32 data), and let "
     "Idea 4 be the research story that unites them. Ideas 5 and 6 are natural extensions once the "
     "basics are solid.")
note_box("Ethics and privacy are not optional here. Identity-from-Wi-Fi (WhoFi and related) shows ordinary "
         "Wi-Fi can identify people; any identity work should enrol only consenting people, keep "
         "signatures on-device, reject unknowns, and document the trade-offs. Making privacy a design "
         "contribution is itself a strong thesis angle.")

# ===========================================================================
# 8. ESP32 TRACK
# ===========================================================================
h1("8.  The ESP32 Practical Track (Direction 2)")
para("This is the hands-on work that produces your own data and demos, and grounds every abstract idea "
     "above in something real and cheap.")

h2("8.1  What the hardware does")
para("Two ESP32-WROOM-32 boards run the open-source ESPectre firmware, which reads Wi-Fi CSI and turns "
     "it into a Movement Score (0\u201310) and a Motion Detected on/off flag, served on a small live web "
     "dashboard. Each board connects to the 2.4 GHz router like a phone, pings it about 100 times a "
     "second to keep packets flowing, and measures the CSI of those packets \u2014 so a single board is both "
     "the listener and the traffic source. Two boards cover two areas (for example two rooms, or two "
     "corners of a lobby).")
layman_box("You do not wire the two boards together. Each one independently watches how the Wi-Fi around "
           "it wobbles as people move, and reports a live \u201chow much motion\u201d number in a web page.")

h2("8.2  The zone-localization experiment")
para("The research script divides the space into a grid of zones (a 4x3 grid was used) and, while a "
     "person walks around, records for each moment: the zone the person is in, both boards' movement "
     "scores, and the person's pixel position. The zone label comes automatically from a camera with a "
     "person detector (YOLO) \u2014 the same camera-teaches-radio idea from Part 1, used here just to label "
     "the data. One run captured about 4,660 samples over roughly three and a half minutes.")
make_table(
    ["What was set up", "Detail"],
    [
        ["Boards", "ESP32-1 at 192.168.0.10, ESP32-2 at 192.168.0.11 (\u201cRoom 2\u201d)."],
        ["Grid", "4 columns x 3 rows of zones over the lobby."],
        ["Labeling", "Camera + YOLO person detector maps pixel position to a zone."],
        ["Logged per sample", "Zone, esp1/esp2 movement + motion, person x/y, timestamp."],
        ["One run", "~4,660 samples, ~3.5 minutes, free-walk mode."],
    ],
    widths=[1.8, 4.5],
)

h2("8.3  What the first results show (honestly)")
bullet("both boards clearly respond to a moving person: the movement score rises when someone is near "
       "and falls when the area is empty, and the motion flag toggles accordingly.", "It works: ")
bullet("ESP32-2 gave a nicely graded response (movement roughly 4\u20136 across occupied zones), which is "
       "what you want for localisation. ESP32-1 tended to saturate near the maximum (about 9.9) in most "
       "occupied zones, so it acted more like a coarse presence detector than a graded sensor \u2014 likely a "
       "sensitivity/threshold or placement effect to tune.", "One board saturated: ")
bullet("the walked data mostly covered the middle row of zones, so coverage of the full grid was "
       "uneven \u2014 a data-collection issue to fix with a more systematic (guided) walk next time.",
       "Coverage was uneven: ")
key_box("Correct status: this is a working data-collection pipeline and a clear proof that two cheap "
        "ESP32 boards sense human motion and can be camera-labelled by zone. It is not yet a trained "
        "localisation model \u2014 that is the next step, and the data to train it is now being collected.")

h2("8.4  Next steps on the ESP32 track")
numbered("tune ESP32-1 (threshold, placement, calibration) so both boards give graded, comparable "
         "scores instead of saturating.", "")
numbered("collect a balanced dataset with a guided walk that visits every zone equally, across a few "
         "days and small furniture changes to test generalisation.", "")
numbered("train a simple model (CSI features to zone) and show a live predicted-presence heatmap; "
         "compare to the camera labels.", "")
numbered("add a second capability by camera-taught distillation (for example activity or fall), reusing "
         "the same collection setup.", "")

# ===========================================================================
# 9. REFERENCES
# ===========================================================================
h1("9.  Consolidated References")
refs = [
    "C. Feng, S. Liang, C. Li, G. Zhao, B. Jing, Y. Xie, X. Chen. Enabling Over-the-Air AI for Edge "
    "Computing via Metasurface-Driven Physical Neural Networks. ACM SIGCOMM 2025. DOI: 10.1145/3718958.3750474.",
    "S. Garcia Sanchez, G. Reus Muns, C. Bocanegra, Y. Li, U. Muncuk, Y. Naderi, Y. Wang, S. Ioannidis, "
    "K. R. Chowdhury. AirNN: Neural Networks with Over-the-Air Convolution via Reconfigurable Intelligent "
    "Surfaces. arXiv:2202.03399, 2022 (later IEEE/ACM Transactions on Networking).",
    "X. Lin, Y. Rivenson, N. T. Yardimci, M. Veli, Y. Luo, M. Jarrahi, A. Ozcan. All-optical machine "
    "learning using diffractive deep neural networks. Science 361, 1004-1008 (2018).",
    "Y. Shen, N. C. Harris, S. Skirlo, M. Prabhu, T. Baehr-Jones, M. Hochberg, X. Sun, S. Zhao, "
    "H. Larochelle, D. Englund, M. Soljacic. Deep learning with coherent nanophotonic circuits. Nature "
    "Photonics 11, 441-446 (2017).",
    "C. Liu, Q. Ma, Z. J. Luo, Q. R. Hong, Q. Xiao, H. C. Zhang, L. Miao, W. M. Yu, Q. Cheng, L. Li, "
    "T. J. Cui. A programmable diffractive deep neural network based on a digital-coding metasurface "
    "array. Nature Electronics 5, 113-122 (2022).",
    "K. Yang, T. Jiang, Y. Shi, Z. Ding. Federated Learning via Over-the-Air Computation. IEEE "
    "Transactions on Wireless Communications, 2020.",
    "Y. Zheng, Y. Zhang, K. Qian, G. Zhang, Y. Liu, C. Wu, Z. Yang. Zero-Effort Cross-Domain Gesture "
    "Recognition with Wi-Fi. ACM MobiSys 2019.",
    "M. Zhao, T. Li, M. Abu Alsheikh, Y. Tian, H. Zhao, A. Torralba, D. Katabi. Through-Wall Human Pose "
    "Estimation Using Radio Signals. IEEE/CVF CVPR 2018.",
    "F. Wang, S. Zhou, S. Panev, J. Han, D. Huang. Person-in-WiFi: Fine-grained Person Perception using "
    "WiFi. IEEE/CVF ICCV 2019.",
    "J. Geng, D. Huang, F. De la Torre. DensePose From WiFi. arXiv:2301.00250, 2023.",
    "B. Korany, C. R. Karanam, H. Cai, Y. Mostofi. XModal-ID: Using WiFi for Through-Wall Person "
    "Identification from Candidate Video Footage. ACM MobiCom 2019.",
    "L. Fan, T. Li, R. Fang, R. Hristov, Y. Yuan, D. Katabi. Learning Longterm Representations for Person "
    "Re-Identification Using Radio Signals. IEEE/CVF CVPR 2020.",
    "D. Avola, E. Emam, D. Montagnini, D. Pannone, A. Ranaldi. WhoFi: Deep Person Re-Identification via "
    "Wi-Fi Channel Signal Encoding. arXiv:2507.12869, 2025.",
    "H. Liu, A. Alali, M. Ibrahim, B. B. Cao, N. Meegan, H. Li, M. Gruteser, et al. Vi-Fi: Associating "
    "Moving Subjects across Vision and Wireless Sensors. ACM/IEEE IPSN 2022.",
    "J. Yang, X. Chen, H. Zou, D. Wang, Q. Xu, L. Xie. SenseFi: A library and benchmark on "
    "deep-learning-empowered WiFi human sensing. Patterns (Cell Press), 2023.",
    "ESPectre (open-source ESP32 Wi-Fi CSI sensing firmware) and the ESP32 CSI Toolkit.",
]
for i, rtext in enumerate(refs, 1):
    p = doc.add_paragraph()
    rn = p.add_run(f"[{i}]  "); rn.bold = True; rn.font.color.rgb = SLATE; rn.font.size = Pt(9.5)
    r2 = p.add_run(rtext); r2.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(4)

closing = doc.add_paragraph(); closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("This guide consolidates the project's scattered notes into one coherent narrative. "
                    "Every paper is real and cited so each claim can be checked at source; where results "
                    "are new or uncertain, that is stated plainly rather than overstated.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE_LT

doc.save(OUTPUT)
print("wrote", OUTPUT)
