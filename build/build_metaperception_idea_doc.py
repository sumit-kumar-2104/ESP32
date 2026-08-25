"""Build the Word document that researches how to COMBINE the MetaAI paper
with the "Camera-Enrolled, Wireless-Tracked Identity" idea (ESPectre Wi-Fi
sensing to reduce cameras).

Source material (all in the MPL folder):
  - Camera_to_Wireless_Identity_Ideas.docx  (the idea catalogue)
  - Wireless_Perception_Reducing_Cameras.docx (the feasibility study)
  - ESPectre_Guide.docx  (the reference Wi-Fi sensing platform)
  - 3718958.3750474.pdf / MetaAI_OverTheAir_Edge_AI.docx  (the MetaAI paper,
        "Enabling Over-the-Air AI for Edge Computing via Metasurface-Driven
        Physical Neural Networks", SIGCOMM 2025, Feng et al.)

This document answers the user's question: how can the MetaAI paper be combined
with our idea, what concretely can be done, and can we use something similar to
get better results.  It is a research + design study, written in plain language
with a scientific layer, and contains no emojis.

The script is self-contained: it generates its own figures and rendered
equations with matplotlib (no external LaTeX needed), then assembles the .docx.
"""
import os
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle, Circle
from PIL import Image

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "assets_metaperception")
OUTPUT = os.path.join(os.path.dirname(HERE),
                      "MetaPerception_Combining_MetaAI_with_WiFi_Sensing.docx")
os.makedirs(ASSETS, exist_ok=True)

# ---- shared light palette --------------------------------------------------
INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
PLUM = RGBColor(0x7A, 0x64, 0x9A)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_PLUM = "F1EDF6"
SHADE_GREY = "F2F3F5"
HEADER_FILL = "DCE6EF"

# matplotlib hex mirror of the palette
H_INK = "#33373d"
H_SLATE = "#4a6b8a"
H_SOFT = "#5b7fa6"
H_BLUE = "#9ecae1"
H_GREEN = "#a8d5a2"
H_ORANGE = "#f4b183"
H_PURPLE = "#c4b7e0"
H_PINK = "#e6a9b8"
H_GREY = "#d9dce1"
H_LGREY = "#eef1f4"

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 11,
    "text.color": H_INK,
    "axes.edgecolor": "#b9bdc4",
    "axes.labelcolor": H_INK,
    "xtick.color": H_INK,
    "ytick.color": H_INK,
    "axes.titlecolor": H_INK,
    "figure.facecolor": "white",
    "axes.facecolor": "white",
})


# ===========================================================================
# FIGURE GENERATION
# ===========================================================================
def _save(fig, name):
    p = os.path.join(ASSETS, name)
    fig.savefig(p, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print("wrote", p)


def render_eq(tex, name, fontsize=20):
    fig = plt.figure(figsize=(0.1, 0.1))
    t = fig.text(0, 0, tex, fontsize=fontsize, color=H_INK)
    fig.canvas.draw()
    bb = t.get_window_extent()
    w, h = bb.width / fig.dpi, bb.height / fig.dpi
    fig.set_size_inches(w + 0.2, h + 0.2)
    t.set_position((0.1 / (w + 0.2), 0.1 / (h + 0.2)))
    _save(fig, name)


def _box(ax, x, y, w, h, fc, ec=None, text="", fs=10, bold=False, tc=H_INK, round=True):
    style = "round,pad=0.02,rounding_size=0.08" if round else "square,pad=0.02"
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle=style,
                                fc=fc, ec=ec or fc, lw=1.4, mutation_aspect=1))
    if text:
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=fs, color=tc, fontweight="bold" if bold else "normal",
                wrap=True)


def _arrow(ax, p0, p1, color=H_SLATE, lw=2.0, style="-|>", ls="-"):
    ax.add_patch(FancyArrowPatch(p0, p1, arrowstyle=style, mutation_scale=16,
                                 color=color, lw=lw, linestyle=ls,
                                 shrinkA=2, shrinkB=2))


def fig_overview():
    fig, ax = plt.subplots(figsize=(9.2, 4.2))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")

    # Left pillar: the idea
    _box(ax, 0.2, 1.2, 3.1, 3.2, H_LGREY, H_SLATE)
    ax.text(1.75, 4.05, "Our idea", ha="center", fontsize=12, fontweight="bold", color=H_SLATE)
    ax.text(1.75, 3.65, "Camera-taught Wi-Fi", ha="center", fontsize=9.5, color=H_INK)
    for i, t in enumerate([
        "One camera kept for detail",
        "ESP32 CSI nodes per room",
        "See through walls (RF)",
        "Camera supervises Wi-Fi",
        "Events, not imagery"]):
        ax.text(1.75, 3.25 - i * 0.42, t, ha="center", fontsize=8.6, color=H_INK)

    # Right pillar: MetaAI
    _box(ax, 6.7, 1.2, 3.1, 3.2, H_LGREY, H_PLUM if False else H_SLATE)
    ax.text(8.25, 4.05, "MetaAI (SIGCOMM'25)", ha="center", fontsize=12, fontweight="bold", color="#6b4f8a")
    ax.text(8.25, 3.65, "Compute in the channel", ha="center", fontsize=9.5, color=H_INK)
    for i, t in enumerate([
        "Metasurface = network weights",
        "Multiply over the air",
        "Device only transmits",
        "Private: server sees results",
        "Multi-sensor time-share"]):
        ax.text(8.25, 3.25 - i * 0.42, t, ha="center", fontsize=8.6, color=H_INK)

    # Middle: fusion
    _box(ax, 3.7, 1.8, 2.6, 2.0, "#eef4ee", H_GREEN, round=True)
    ax.text(5.0, 3.35, "MetaPerception", ha="center", fontsize=12.5, fontweight="bold", color="#4f7a53")
    ax.text(5.0, 2.85, "Shared linearity of", ha="center", fontsize=8.8, color=H_INK)
    ax.text(5.0, 2.55, "radio + neural nets", ha="center", fontsize=8.8, color=H_INK)
    ax.text(5.0, 2.15, "Richer, cheaper,", ha="center", fontsize=8.8, color=H_INK)
    ax.text(5.0, 1.9, "more private sensing", ha="center", fontsize=8.8, color=H_INK)

    _arrow(ax, (3.35, 2.8), (3.7, 2.8), color=H_SLATE)
    _arrow(ax, (6.65, 2.8), (6.3, 2.8), color="#6b4f8a")
    ax.text(5.0, 0.7, "The same physics that carries Wi-Fi sensing also carries the neural-network math.",
            ha="center", fontsize=9, style="italic", color=H_SLATE)
    _save(fig, "fig_overview.png")


def fig_architecture():
    fig, ax = plt.subplots(figsize=(9.4, 5.0))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5.4); ax.axis("off")

    # apartment outline: 3 rooms
    ax.add_patch(Rectangle((0.4, 0.6), 9.2, 4.2, fill=False, ec=H_INK, lw=1.6))
    ax.plot([3.5, 3.5], [0.6, 4.8], color=H_INK, lw=1.2)
    ax.plot([6.6, 6.6], [0.6, 4.8], color=H_INK, lw=1.2)
    ax.text(1.95, 4.55, "Bedroom (camera-free)", ha="center", fontsize=9, color=H_SLATE)
    ax.text(5.05, 4.55, "Hall / kitchen (camera-free)", ha="center", fontsize=9, color=H_SLATE)
    ax.text(8.1, 4.55, "Living room (1 camera)", ha="center", fontsize=9, color=H_SLATE)

    # router + metasurface near centre
    _box(ax, 4.55, 2.05, 0.9, 0.55, H_BLUE, H_SLATE, "Router", 8.2, bold=True)
    # metasurface panel
    ax.add_patch(Rectangle((4.5, 2.75), 1.0, 0.28, fc=H_PURPLE, ec="#6b4f8a", lw=1.4))
    for gx in np.linspace(4.56, 5.44, 9):
        ax.plot([gx, gx], [2.77, 3.01], color="#6b4f8a", lw=0.7)
    ax.text(5.0, 3.2, "Programmable metasurface", ha="center", fontsize=8.4, color="#6b4f8a", fontweight="bold")

    # ESP32 nodes
    node_pos = [(2.6, 1.5), (5.0, 1.3), (8.0, 1.5)]
    for i, (nx, ny) in enumerate(node_pos):
        _box(ax, nx - 0.42, ny - 0.28, 0.84, 0.5, H_GREEN, "#4f7a53", "ESP32", 8, bold=True)
    # camera
    ax.add_patch(Circle((8.6, 3.7), 0.22, fc=H_ORANGE, ec="#b5722f", lw=1.4))
    ax.text(8.6, 3.35, "camera", ha="center", fontsize=8, color="#b5722f")

    # steered/illuminating beams from metasurface to each room
    for (nx, ny) in node_pos:
        _arrow(ax, (5.0, 2.75), (nx, ny + 0.25), color="#6b4f8a", lw=1.6, style="-|>", ls=(0, (4, 2)))

    # hub
    _box(ax, 0.65, 0.75, 1.15, 0.5, H_GREY, H_SLATE, "Hub / HA", 8, bold=True)
    for (nx, ny) in node_pos:
        _arrow(ax, (nx, ny - 0.28), (1.25, 1.0), color=H_SLATE, lw=1.0, style="-|>", ls=(0, (2, 2)))
    _arrow(ax, (8.6, 3.5), (1.4, 1.25), color="#b5722f", lw=1.0, style="-|>", ls=(0, (2, 2)))

    ax.text(5.0, 0.25,
            "Metasurface steers illumination into camera-free rooms and encodes the sensing/inference weights; "
            "ESP32 nodes report, the hub fuses.",
            ha="center", fontsize=8.4, style="italic", color=H_SLATE)
    _save(fig, "fig_architecture.png")


def fig_ota_pipeline():
    fig, ax = plt.subplots(figsize=(9.4, 3.4))
    ax.set_xlim(0, 12); ax.set_ylim(0, 4); ax.axis("off")
    y = 1.7
    _box(ax, 0.2, y, 1.9, 1.1, H_GREEN, "#4f7a53", "Sensor / CSI\nx1 x2 ... xU", 8.6, bold=True)
    _box(ax, 2.7, y, 2.2, 1.1, H_BLUE, H_SLATE, "Transmit inputs\none per time slot", 8.6, bold=True)
    _box(ax, 5.5, y, 2.3, 1.1, H_PURPLE, "#6b4f8a", "Metasurface\nmultiplies by weight\nH(t_i) = w_(r,i)", 8.4, bold=True)
    _box(ax, 8.4, y, 1.8, 1.1, H_BLUE, H_SLATE, "Receiver\naccumulates", 8.6, bold=True)
    _box(ax, 10.5, y, 1.3, 1.1, H_ORANGE, "#b5722f", "Class\nscore", 9, bold=True)
    for x0, x1 in [(2.1, 2.7), (4.9, 5.5), (7.8, 8.4), (10.2, 10.5)]:
        _arrow(ax, (x0, y + 0.55), (x1, y + 0.55), color=H_SLATE)
    ax.text(6.0, 0.7, "Multiplication happens in the air; only the running sum is done in software.",
            ha="center", fontsize=8.8, style="italic", color=H_SLATE)
    ax.text(6.0, 3.35, "Over-the-air inference for a Wi-Fi sensing node",
            ha="center", fontsize=11, fontweight="bold", color=H_SLATE)
    _save(fig, "fig_ota_pipeline.png")


def fig_uplift():
    fig, ax = plt.subplots(figsize=(8.6, 3.6))
    labels = ["Multi-sensor\nfusion", "Noise-aware\ntraining", "Timing fix\n(CDFA)", "Multipath\ncancellation"]
    base = [63, 80.5, 19, 60]
    withm = [90, 87.9, 89, 82]
    x = np.arange(len(labels)); w = 0.36
    ax.bar(x - w / 2, base, w, label="without remedy", color=H_GREY, edgecolor="#b9bdc4")
    ax.bar(x + w / 2, withm, w, label="with MetaAI remedy", color=H_GREEN, edgecolor="#4f7a53")
    for xi, (b, m) in enumerate(zip(base, withm)):
        ax.text(xi - w / 2, b + 1.5, f"{b:g}", ha="center", fontsize=8, color=H_INK)
        ax.text(xi + w / 2, m + 1.5, f"{m:g}", ha="center", fontsize=8, color="#4f7a53", fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=8.6)
    ax.set_ylabel("accuracy (%)"); ax.set_ylim(0, 105)
    ax.legend(frameon=False, fontsize=8.5, loc="upper center", ncol=2, bbox_to_anchor=(0.5, 1.16))
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_title("What MetaAI's engineering buys, as a template for our Wi-Fi nodes",
                 fontsize=10, pad=22)
    _save(fig, "fig_uplift.png")


def fig_roadmap():
    fig, ax = plt.subplots(figsize=(9.4, 3.2))
    ax.set_xlim(0, 12); ax.set_ylim(0, 3.4); ax.axis("off")
    phases = [
        ("Phase 0", "ESPectre now", "One camera + ESP32 CSI nodes.\nPresence, motion, which-room.", H_GREEN, "#4f7a53"),
        ("Phase 1", "Camera-taught Wi-Fi", "Cross-modal distillation adds\nactivity and fall classes.", H_BLUE, H_SLATE),
        ("Phase 2", "Metasurface coverage", "RIS/metasurface steers signal\ninto dead rooms; raises DoF.", H_PURPLE, "#6b4f8a"),
        ("Phase 3", "Over-the-air inference", "Channel computes the classifier;\nnodes just transmit.", H_ORANGE, "#b5722f"),
    ]
    w = 2.7
    for i, (ph, title, body, fc, ec) in enumerate(phases):
        x = 0.2 + i * 2.95
        _box(ax, x, 0.7, w, 1.9, "#ffffff", ec)
        ax.text(x + w / 2, 2.35, ph, ha="center", fontsize=9.5, fontweight="bold", color=ec)
        ax.text(x + w / 2, 2.02, title, ha="center", fontsize=9, color=H_INK, fontweight="bold")
        ax.text(x + w / 2, 1.35, body, ha="center", fontsize=7.9, color=H_INK)
        if i < 3:
            _arrow(ax, (x + w, 1.65), (x + w + 0.25, 1.65), color=H_SLATE)
    ax.text(6.0, 0.25, "Each phase is deployable on its own; value compounds toward the frontier.",
            ha="center", fontsize=8.6, style="italic", color=H_SLATE)
    _save(fig, "fig_roadmap.png")


def fig_capability():
    fig, ax = plt.subplots(figsize=(8.8, 3.8))
    tasks = ["Presence", "Count", "Coarse\nlocation", "Activity", "Fall", "Coarse\npose", "Identity"]
    wifi_now = [0.95, 0.7, 0.85, 0.6, 0.65, 0.3, 0.05]
    with_meta = [0.97, 0.85, 0.92, 0.82, 0.85, 0.6, 0.1]
    x = np.arange(len(tasks)); w = 0.36
    ax.bar(x - w / 2, wifi_now, w, label="Wi-Fi today (single link)", color=H_GREY, edgecolor="#b9bdc4")
    ax.bar(x + w / 2, with_meta, w, label="+ metasurface / OTA fusion", color=H_PURPLE, edgecolor="#6b4f8a")
    ax.set_xticks(x); ax.set_xticklabels(tasks, fontsize=8.4)
    ax.set_ylabel("expected capability"); ax.set_ylim(0, 1.08)
    ax.legend(frameon=False, fontsize=8.5, loc="upper right")
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_title("Where a metasurface layer should help most", fontsize=10, pad=8)
    _save(fig, "fig_capability.png")


def build_figures():
    fig_overview()
    fig_architecture()
    fig_ota_pipeline()
    fig_uplift()
    fig_roadmap()
    fig_capability()
    # equations
    render_eq(r"$y = H\,x + n$", "eq_channel.png")
    render_eq(r"$H(f_k)=\sum_{p=1}^{P} a_p\,e^{-j 2\pi f_k \tau_p}$", "eq_csi.png")
    render_eq(r"$y_r=\left|\;\sum_{i=0}^{U} H_r(t_i)\cdot x_i\;\right|$", "eq_core.png")
    render_eq(r"$H_{mts}=\alpha\sum_{m=1}^{M} e^{\,j\phi_m^{p}}\,e^{\,j\phi_m}\;\;\Rightarrow\;\;\mathrm{DoF}\uparrow$",
              "eq_dof.png")
    render_eq(r"$\theta^{\star}=\arg\min_{\theta}\sum_{(x_{csi},\,y_{cam})\in O}\mathcal{L}\left(f_\theta(x_{csi}),\,y_{cam}\right)$",
              "eq_distill.png")
    render_eq(r"$p(s\,|\,e_{cam},e_{rf})\;\propto\;p(e_{cam}\,|\,s)\,p(e_{rf}\,|\,s)\,p(s)$",
              "eq_fusion.png")
    render_eq(r"$\Phi=\arg\min_{\phi_m}\left|\,H_{mts}-(H_{des}-H_e)\,\right|$", "eq_multipath.png")


# Generate all figures and equations up front, before the document body is
# assembled (the body embeds these image files as it goes).
build_figures()

# ===========================================================================
# DOCUMENT SETUP
# ===========================================================================
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), "8")
    e.set(qn("w:color"), color)
    pbdr.append(e)
    pPr.append(pbdr)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def para(text, italic=False, size=11, align=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def layman_box(text):
    p = doc.add_paragraph()
    _shade(p, SHADE_GREEN)
    _border(p, color="AFC8B2")
    r = p.add_run("In plain terms:  ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    return p


def sci_box(text):
    p = doc.add_paragraph()
    _shade(p, SHADE_BLUE)
    _border(p, color="9DB6CC")
    r = p.add_run("For the researcher:  ")
    r.bold = True
    r.font.color.rgb = SLATE
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    return p


def idea_box(text):
    p = doc.add_paragraph()
    _shade(p, SHADE_PLUM)
    _border(p, color="B9A9D0")
    r = p.add_run("Why this is new:  ")
    r.bold = True
    r.font.color.rgb = PLUM
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    return p


def note_box(text):
    p = doc.add_paragraph()
    _shade(p, SHADE_GREY)
    _border(p, color="C9CDD3")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


_eqn = {"n": 0}


def equation(img, label=True):
    _eqn["n"] += 1
    path = os.path.join(ASSETS, img)
    with Image.open(path) as im:
        w_in = im.width / 200.0
        h_in = im.height / 200.0
    max_w = 5.9
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if w_in > max_w:
        run.add_picture(path, width=Inches(max_w))
    else:
        target_h = max(min(h_in, 0.6), 0.28)
        run.add_picture(path, height=Inches(target_h))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    if label:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"( {_eqn['n']} )")
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE_LT
        cap.paragraph_format.space_after = Pt(8)
    return _eqn["n"]


_fig = {"n": 0}


def figure(img, caption, width=6.2):
    _fig["n"] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(width))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {_fig['n']}.  ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = SLATE
    r2 = cap.add_run(caption)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = INK
    cap.paragraph_format.space_after = Pt(12)


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = INK
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F6F8FA")
                tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ===========================================================================
# TITLE PAGE
# ===========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MetaPerception")
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(66)
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Combining the MetaAI Over-the-Air Compute Paper with Camera-Taught Wi-Fi Sensing")
r.font.size = Pt(15)
r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("How the metasurface neural network of SIGCOMM 2025 can make our "
                 "\u201creduce-cameras-with-Wi-Fi\u201d idea richer, cheaper, and more private")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = INK

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  A Research and Design Study  \u2014")
r.font.size = Pt(11)
r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(24)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Builds on: the ESPectre Wi-Fi sensing platform, the "
                 "\u201cCamera-Enrolled, Wireless-Tracked Identity\u201d idea catalogue, and "
                 "Feng et al., \u201cEnabling Over-the-Air AI for Edge Computing via "
                 "Metasurface-Driven Physical Neural Networks\u201d (SIGCOMM 2025).")
r.font.size = Pt(10)
r.font.color.rgb = SLATE_LT

doc.add_page_break()

# ===========================================================================
# CONTENTS
# ===========================================================================
h1("Contents")
toc = [
    "1.  The Question, Answered in One Page",
    "2.  Recap of the Two Pieces We Are Joining",
    "3.  Why They Fit: One Shared Idea (Linearity)",
    "4.  A Catalogue of Concrete Ways to Combine Them",
    "5.  A Worked Example: The Three-Room Apartment, Upgraded",
    "6.  Can We Get Better Results? What the Numbers Suggest",
    "7.  The Mathematics That Ties Them Together",
    "8.  Requirements: Mathematical, Realistic, Practical",
    "9.  A Phased Roadmap",
    "10.  Limitations and Honest Expectations",
    "11.  Novelty: What Would Be Publishable",
    "12.  References and Further Reading",
]
for t in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
doc.add_page_break()

# ===========================================================================
# 1. EXECUTIVE SUMMARY
# ===========================================================================
h1("1.  The Question, Answered in One Page")
para("The task is simple to state: we have an idea for reducing the number of home cameras by "
     "letting cheap Wi-Fi sensors see the rooms a single camera cannot, and we have just studied "
     "the MetaAI paper, which shows how to perform artificial-intelligence computation inside the "
     "wireless channel using a programmable metasurface. The question is whether the paper can be "
     "folded into our idea, what concretely can be built, and whether it can give better results.")
para("The short answer is yes, and for a reason that is deeper than convenience: both systems rest "
     "on the very same mathematics. Wi-Fi sensing reads a person out of the linear wireless channel; "
     "MetaAI programs that same linear channel to act as the weights of a neural network. Because the "
     "physics is shared, MetaAI is not a competing project but a natural upgrade path for the "
     "Wi-Fi half of our design.")
figure("fig_overview.png",
       "The two pieces and their meeting point. Our camera-taught Wi-Fi idea (left) and the MetaAI "
       "over-the-air compute paper (right) share the linear wireless channel, so they combine into a "
       "single richer system (centre).", width=6.6)
para("Concretely, this study identifies five ways to combine them, ranging from deployable-soon to "
     "research-frontier:")
bullet("use a metasurface (a reconfigurable intelligent surface) to steer Wi-Fi illumination into "
       "the camera-free rooms and raise the effective resolution of a single cheap link;", bold_lead="Coverage and resolution: ")
bullet("push the Wi-Fi node's classifier into the channel itself, so the ESP32 only transmits and "
       "the metasurface plus receiver compute the answer, saving energy and hardware;", bold_lead="Over-the-air inference: ")
bullet("keep the camera-as-teacher scheme, but let the metasurface implement the learned "
       "student model physically, closing the loop from training to deployment;", bold_lead="Physical distillation: ")
bullet("time-share one metasurface across the camera and every Wi-Fi node, fusing modalities the "
       "way MetaAI fuses multiple sensors for up to a 27 percent accuracy gain;", bold_lead="Shared-surface fusion: ")
bullet("exploit that the server receives only inference results, never raw CSI or images, which is "
       "exactly the privacy property our bedroom and bathroom rooms demand.", bold_lead="Privacy by design: ")
note_box("Bottom line: MetaAI does not replace our idea; it supercharges its weakest half. Wi-Fi "
         "sensing gives us presence, motion, coarse location and activity today. A metasurface layer "
         "promises the same at higher fidelity, lower node energy, and stronger privacy, and it turns "
         "the programme's long-term frontier note into a concrete engineering plan.")

# ===========================================================================
# 2. RECAP
# ===========================================================================
h1("2.  Recap of the Two Pieces We Are Joining")

h2("2.1  Our idea: camera-enrolled, wireless-tracked sensing")
para("The idea keeps one camera where visual detail truly matters (a face at the door, verifying an "
     "alarm) and covers every other room with a single cheap Wi-Fi link: the home router as "
     "transmitter and a small ESP32 running the open-source ESPectre firmware as receiver. Each node "
     "reads Channel State Information (CSI) and reports presence, motion, and, with a model, activity "
     "and falls. Radio passes through walls, so one link per room replaces several cameras for "
     "event-level tasks. A short enrolment period, during which the camera and a Wi-Fi node see the "
     "same scene, lets the camera automatically label CSI and teach the Wi-Fi model, after which the "
     "camera can leave.")
bullet("Strengths: through-wall, private, cheap (about ten euro per node), no per-room wiring.")
bullet("Weaknesses: a single commodity link is low-resolution; it does presence and motion well, "
       "activity and pose poorly; models are environment-specific and must be calibrated per home.")

h2("2.2  MetaAI: computing inside the wireless channel")
para("MetaAI observes that a linear neural-network layer is just a matrix multiplying a vector, and "
     "that the wireless channel already multiplies whatever passes through it. So it programs a "
     "metasurface, a flat array of hundreds of tiny elements each adding a controllable phase shift, "
     "to make the channel equal to the network's weights. The sensor transmits its inputs one per "
     "time slot; the metasurface multiplies each by the right weight over the air; the receiver adds "
     "the results into a running total; the magnitude of that total is the class score.")
bullet("Strengths: the sensing device only transmits and never runs the AI (saving energy and cost); "
       "one surface implements a whole network and serves many sensors by time-division; the server "
       "sees results not raw data (private by design); it runs on ordinary commodity radios.")
bullet("Limitations: it currently implements only linear networks; larger models cost more "
       "transmission time (latency); accuracy is capped by the metasurface resolution; and motion "
       "forces the weight mapping to be recomputed.")
note_box("Reported results to anchor expectations: 82.8 percent average accuracy across six datasets "
         "with a single linear layer; multi-sensor fusion adds up to 27 percent; the timing scheme "
         "lifts accuracy from about 19 to 89 percent; noise-aware training lifts the 80th-percentile "
         "accuracy from about 80.5 to 87.9 percent; and it works across 2.4, 3.5 and 5 GHz, several "
         "modulations, non-line-of-sight corners, and across rooms.")

# ===========================================================================
# 3. WHY THEY FIT
# ===========================================================================
h1("3.  Why They Fit: One Shared Idea (Linearity)")
para("The reason these two projects belong together is not a marketing slogan; it is a single "
     "mathematical fact. What a Wi-Fi sensor measures and what a neural-network layer computes are "
     "both linear operations on the wireless channel.")
para("A receiver never sees an image. It measures how the environment transformed the transmitted "
     "signal:")
equation("eq_channel.png")
para("Everything about the room is carried inside the channel H. Wi-Fi sensing reads a person out of "
     "H because a moving body is a moving reflector that writes patterns into the per-frequency "
     "channel, the CSI:")
equation("eq_csi.png")
para("MetaAI writes into that same H. It sets the metasurface so the channel becomes the weights of "
     "a classifier, then lets the signal carry out the multiply-and-accumulate as it propagates:")
equation("eq_core.png")
layman_box("Our Wi-Fi sensor listens to what the channel does to a signal in order to learn about the "
           "room. MetaAI does the reverse: it tells the channel what to do so the signal comes out "
           "already classified. Same channel, read in one project and written in the other. That is "
           "why one can host the other.")
sci_box("Formally, both operations live in the linear map y = H x. Sensing is the inference problem "
        "\u201cgiven y and x, infer the scene inside H\u201d; over-the-air compute is the synthesis "
        "problem \u201cchoose H so that y is the desired linear functional of x\u201d. A metasurface "
        "gives us direct, reconfigurable control of H, which is exactly the missing actuator that "
        "turns a passive sensing link into an active, programmable one.")

# ===========================================================================
# 4. CATALOGUE
# ===========================================================================
h1("4.  A Catalogue of Concrete Ways to Combine Them")
para("Five options, ordered from the most immediately practical to the most ambitious. They are "
     "additive: each can be adopted on its own, and later ones assume the earlier ones.")

h2("4.1  Option A \u2014 Metasurface coverage and resolution boost")
para("The single biggest weakness of our Wi-Fi half is that one cheap link is low-resolution and may "
     "not illuminate a whole room, especially around corners and behind furniture. A reconfigurable "
     "intelligent surface (the same hardware family as the MetaAI metasurface) placed on a wall can "
     "steer and focus the router's signal into the exact rooms and blind spots we care about, and can "
     "synthesise a much larger effective aperture than the ESP32's single antenna.")
equation("eq_dof.png")
para("By adding hundreds of controllable reflection paths, the surface raises the number of "
     "independent measurements a single link can make, which is the fundamental quantity that limits "
     "how much a Wi-Fi link can resolve. In our idea's own terms, this is the cheapest way to climb "
     "the capability ladder from presence toward coarse location and pose without adding access "
     "points.")
bullet("steer illumination into non-line-of-sight rooms and corners the direct link misses;", bold_lead="What it buys: ")
bullet("raise angular and spatial resolution so a single link reports more than presence;")
bullet("null out clutter from a static room to make the human-motion signal cleaner.")
idea_box("MetaAI validates that these surfaces work with ordinary commodity radios and that far-field "
         "control needs only an angle, not the receiver's exact position. That removes the main "
         "deployment objection to putting a reconfigurable surface in a home.")

h2("4.2  Option B \u2014 Over-the-air inference for the Wi-Fi nodes")
para("Today the ESPectre node runs a small machine-learning detector on the ESP32 itself. MetaAI's "
     "central move lets us delete that on-device model: transmit the CSI features one per time slot, "
     "let the metasurface multiply each by the classifier's weights over the air, and let the "
     "receiver accumulate the class score. The node becomes a pure transmitter.")
figure("fig_ota_pipeline.png",
       "Over-the-air inference for a Wi-Fi sensing node. The sensor transmits its CSI features; the "
       "metasurface multiplies each by a trained weight in the air; the receiver accumulates the "
       "result into a class score. No AI runs on the node.", width=6.6)
bullet("longer battery life and cheaper nodes, because the heavy computation leaves the device;", bold_lead="What it buys: ")
bullet("the classifier can be updated by reprogramming the surface, not reflashing every node;")
bullet("one surface can serve many nodes by time-division, amortising its cost across the home.")
sci_box("The Wi-Fi activity or fall classifier we train in Section 4.3 of the idea is, in its "
        "simplest form, a linear layer over CSI features. That is precisely the model class MetaAI "
        "can realise physically. A non-linear classifier would need either a digital non-linearity at "
        "the receiver or the paper's absolute-value nonlinearity, which already gave 82.8 percent "
        "average accuracy with a single layer.")

h2("4.3  Option C \u2014 Physical knowledge distillation (camera still teaches)")
para("Our idea's key trick is camera-supervised learning: during enrolment the camera generates "
     "pseudo-labels that train the Wi-Fi model. MetaAI supplies the missing second half, a way to "
     "deploy that trained model in hardware. Train the student model in software with continuous "
     "weights exactly as the paper recommends, then snap those weights onto the metasurface's "
     "discrete states and let the channel run the student at inference time.")
equation("eq_distill.png")
para("This closes the loop cleanly: the camera teaches during enrolment, the software model captures "
     "the mapping, and the metasurface becomes the deployed student that runs for free in the "
     "channel afterwards.")
layman_box("The camera is the tutor and the Wi-Fi node is the pupil, just as before. What MetaAI adds "
           "is that once the pupil has learned, the lesson is written directly into a physical surface "
           "on the wall, so the answer appears in the signal itself with no computer thinking about it.")

h2("4.4  Option D \u2014 One shared surface fuses camera and Wi-Fi")
para("MetaAI fuses multiple sensors through a single metasurface by giving each sensor its own time "
     "window and summing the partial scores. Our system already has multiple sensors: one camera and "
     "several Wi-Fi nodes. Routing them all through one shared surface lets us do principled late "
     "fusion in hardware, which is the same posterior combination our idea already specifies.")
equation("eq_fusion.png")
figure("fig_capability.png",
       "Where a metasurface layer should help most. A single Wi-Fi link today (grey) is strong for "
       "presence and location but weak for activity, pose and identity. A metasurface and over-the-air "
       "fusion layer (purple) is expected to lift the middle of the ladder most; identity still needs "
       "the camera.", width=6.4)
bullet("the paper reports up to 27 percent accuracy gain from multi-sensor fusion;", bold_lead="What it buys: ")
bullet("complementary sensors reinforce the right answer (camera for identity, Wi-Fi for coverage);")
bullet("one surface, time-shared, avoids separate hardware per sensor.")

h2("4.5  Option E \u2014 Privacy-preserving edge sensing")
para("The most sensitive rooms in our design (bedroom, bathroom) are exactly where cameras are "
     "unacceptable and where even raw CSI could, in principle, be abused. MetaAI's architecture "
     "means the server receives only the inference result, for example the single fact "
     "\u201csomeone fell\u201d, and never the raw signal. This is not an add-on; it is inherent to "
     "computing in the channel.")
bullet("raw CSI and images never leave the sensing point in reconstructable form;", bold_lead="What it buys: ")
bullet("mitigates the very privacy risk our idea flags, that Wi-Fi can identify individuals;")
bullet("aligns with keeping data on the local hub, which our deployment already requires.")
idea_box("A 2026 result identified individuals from ordinary Wi-Fi with near-perfect accuracy. Doing "
         "the classification in the channel, and emitting only a coarse event, is a concrete "
         "engineering answer to that privacy hazard rather than a policy promise.")

# ===========================================================================
# 5. WORKED EXAMPLE
# ===========================================================================
h1("5.  A Worked Example: The Three-Room Apartment, Upgraded")
para("Take the same three-room apartment from the idea, one camera in the living room and one ESP32 "
     "per other room, and add a single wall-mounted metasurface near the router. Nothing else about "
     "the layout changes.")
figure("fig_architecture.png",
       "The upgraded reference architecture. A programmable metasurface near the router steers "
       "illumination into the two camera-free rooms and encodes the sensing and inference weights; the "
       "ESP32 nodes report and the hub fuses the camera and Wi-Fi evidence into events.", width=6.7)
para("The metasurface plays three roles at once, cycled in time slots:")
numbered("steer and focus the router's signal into the bedroom and hall so a single link covers each "
         "room fully, including corners the direct path misses;", bold_lead="Illuminate: ")
numbered("during each node's slot, set the surface to the trained classifier weights so the returning "
         "signal already carries the activity or fall score;", bold_lead="Compute: ")
numbered("in a dedicated slot, combine the camera's per-frame evidence with each Wi-Fi node's score "
         "into a single fused event for the hub.", bold_lead="Fuse: ")
para("Deployment steps, extending the idea's own recipe:")
numbered("stand up the Wi-Fi layer with three ESPectre nodes and one camera, exactly as before;")
numbered("mount one metasurface panel near the router; calibrate its steering angles per room using "
         "standard beam scanning (only the angle is needed, not exact positions);")
numbered("run the short camera-supervised enrolment to train activity and fall classifiers;")
numbered("map each trained classifier onto the surface's discrete phase states;")
numbered("define fusion events in Home Assistant, now sourced from in-channel scores rather than "
         "on-device inference.")

# ===========================================================================
# 6. BETTER RESULTS
# ===========================================================================
h1("6.  Can We Get Better Results? What the Numbers Suggest")
para("\u201cBetter\u201d must be defined carefully, because a metasurface does not turn Wi-Fi into a "
     "camera. It cannot deliver faces, colour, or readable detail; identity still needs the one "
     "camera. Where it should help is in three measurable directions: coverage, fidelity of the "
     "event classes, and node energy. The paper's own ablations are the best available evidence for "
     "how much each engineering remedy is worth, and they transfer directly to our node classifiers.")
figure("fig_uplift.png",
       "What MetaAI's engineering buys, as a template for our Wi-Fi nodes. Each remedy the paper "
       "introduces (fusion, noise-aware training, timing correction, multipath cancellation) yields a "
       "large, separately measured accuracy gain that our activity and fall classifiers can inherit.", width=6.4)
para("Reading the transfer conservatively:")
bullet("more independent measurements per link should push single-link tasks up the capability "
       "ladder, most visibly for activity and coarse pose, which are the idea's known weak spots.", bold_lead="Coverage and resolution (Option A): ")
bullet("noise-aware and timing-robust training, which the paper shows are worth tens of points, are "
       "pure software techniques we can apply to the ESPectre models immediately, with or without a "
       "surface.", bold_lead="Robustness: ")
bullet("up to 27 percent from combining the camera and Wi-Fi evidence in one place, which our late "
       "fusion already aims at but can now do in hardware.", bold_lead="Fusion (Option D): ")
bullet("moving inference off the ESP32 lowers node power and cost, improving the practical "
       "\u201cten-euro node\u201d economics rather than raw accuracy.", bold_lead="Energy (Option B): ")
note_box("Honest calibration: the paper's numbers were measured on image and gesture datasets, not on "
         "our through-wall CSI activity task, so treat them as an upper-bound template, not a promise. "
         "The robustness techniques transfer with high confidence because they are training-side; the "
         "metasurface accuracy gains must be validated in our own domain.")

# ===========================================================================
# 7. MATHEMATICS
# ===========================================================================
h1("7.  The Mathematics That Ties Them Together")
para("This section collects the equations that justify the combination, with every symbol explained. "
     "They are deliberately the same objects that appear in both source documents, which is the point.")

h2("7.1  The shared linear channel")
equation("eq_channel.png")
para("What it says. The received signal y is the channel H acting on the transmitted signal x, plus "
     "noise n. Sensing infers the room from H; over-the-air compute chooses H. Symbols: y received "
     "signal, H channel, x transmitted signal, n noise.")

h2("7.2  What the Wi-Fi node reads: CSI as a sum over paths")
equation("eq_csi.png")
para("What it says. The channel at frequency f_k is a sum over the P paths the signal travelled; each "
     "path contributes an amplitude a_p and a phase set by its delay tau_p. A moving person changes "
     "one path's delay and amplitude, which is the sensing signal. This is unchanged from our idea.")

h2("7.3  What MetaAI writes: the channel as network weights")
equation("eq_core.png")
para("What it says. The class score y_r is the magnitude of the accumulated sum, over inputs i, of "
     "the channel response H_r(t_i) multiplied by input x_i. Setting the metasurface so H_r(t_i) "
     "equals the trained weight makes the propagating signal compute a neural-network layer. Symbols: "
     "y_r score for class r, H_r(t_i) channel at slot i for class r, x_i i-th input, U number of "
     "inputs.")

h2("7.4  The new actuator: a metasurface raises the degrees of freedom")
equation("eq_dof.png")
para("What it says. The metasurface channel is an amplitude times a sum over its M meta-atoms of two "
     "phase rotations, the natural propagation phase and the deliberate phase the atom adds. Because "
     "we control each atom's phase, we add hundreds of tunable paths, which raises the independent "
     "measurements (degrees of freedom) a single link can make, the quantity that caps resolution in "
     "our idea. Symbols: M number of meta-atoms, phi_m the controllable phase of atom m.")

h2("7.5  Deploying the student: multipath-aware configuration")
equation("eq_multipath.png")
para("What it says. Choose the atom phases that make the realised channel match the desired trained "
     "weight after subtracting the known static room response H_e. This is how a software-trained "
     "classifier becomes a physical surface in a real, cluttered room. Symbols: H_mts realised "
     "channel, H_des desired weight, H_e static environment response, phi_m atom phases.")

h2("7.6  Combining the sensors: fusion")
equation("eq_fusion.png")
para("What it says. The probability of a scene state s given both camera and radio evidence is "
     "proportional to each sensor's likelihood times a prior. The more confident sensor dominates "
     "automatically, camera for identity, Wi-Fi for through-wall coverage. This is the same fusion "
     "rule in the idea, now realisable on one shared surface.")

h2("7.7  Training the student from the camera")
equation("eq_distill.png")
para("What it says. Choose model parameters theta that minimise, over the enrolment overlap O, the "
     "loss between the model's prediction from CSI and the camera-derived label. After training, only "
     "CSI is needed, and the weights are mapped to the surface. Symbols: theta model parameters, "
     "f_theta the CSI-to-label model, x_csi CSI input, y_cam camera pseudo-label, O overlap window.")

# ===========================================================================
# 8. REQUIREMENTS
# ===========================================================================
h1("8.  Requirements: Mathematical, Realistic, Practical")

h2("8.1  Mathematical requirements")
bullet("only what the physics allows is recoverable. Presence, motion, coarse location, activity and "
       "breathing are identifiable from CSI; identity and appearance are not, so the camera stays.", bold_lead="Identifiability: ")
bullet("the metasurface must add enough independent, controllable paths (the paper found returns "
       "saturate near 256 meta-atoms) to lift resolution meaningfully.", bold_lead="Degrees of freedom: ")
bullet("the classifier realised in the channel is linear (plus a magnitude nonlinearity); anything "
       "deeper needs a digital nonlinearity at the receiver.", bold_lead="Linearity: ")
bullet("transmitter and surface do not share a clock, so timing drift must be handled; the paper's "
       "training-time correction holds accuracy to about a 4-microsecond drift.", bold_lead="Synchronisation: ")

h2("8.2  Realistic requirements")
note_box("Reality check: like all Wi-Fi sensing, the models are environment-specific and drift after "
         "big furniture changes, so expect per-home calibration. A metasurface adds its own "
         "calibration (steering angles, static-clutter subtraction) and its accuracy gains must be "
         "proven on our through-wall CSI task, not assumed from the paper's image datasets.")
bullet("mobility of the target changes the paths and forces the weight mapping to be recomputed; fast "
       "motion is a race between the person's speed and recalibration.")
bullet("multi-person scenes remain hard from a single link; more surface paths or a brief camera "
       "assist help.")
bullet("a metasurface adds cost and mounting effort, so its benefit must beat simply adding a second "
       "ESP32 link.")

h2("8.3  Practical requirements")
bullet("place the router-to-node line across each room, and mount the metasurface where it can "
       "illuminate the target rooms, ideally with line of sight to the router.", bold_lead="Geometry: ")
bullet("run beam scanning to find each room's steering angle; only the angle is needed, not exact "
       "positions.", bold_lead="Calibrate: ")
bullet("run the short camera-supervised enrolment, train with continuous weights, then map to the "
       "surface's discrete states.", bold_lead="Enrol and map: ")
bullet("collect in-channel scores in Home Assistant and define the same fusion events as before.", bold_lead="Fuse: ")

# ===========================================================================
# 9. ROADMAP
# ===========================================================================
h1("9.  A Phased Roadmap")
para("The combination is best built incrementally, so that each phase delivers value and de-risks the "
     "next. Nothing requires waiting for the frontier before the system is useful.")
figure("fig_roadmap.png",
       "A phased roadmap. Start with today's ESPectre coverage, add camera-taught Wi-Fi, then a "
       "metasurface for coverage and resolution, and finally over-the-air inference. Each phase is "
       "independently deployable.", width=6.7)
bullet("one camera plus ESPectre CSI nodes; reliable presence, motion and which-room today.", bold_lead="Phase 0 (now): ")
bullet("cross-modal distillation from the camera adds activity and fall classes; apply MetaAI's "
       "noise-aware and timing-robust training tricks in software for immediate robustness gains.", bold_lead="Phase 1: ")
bullet("add one metasurface to steer illumination into dead rooms and raise single-link resolution; "
       "validate the capability uplift on our own CSI task.", bold_lead="Phase 2: ")
bullet("move the node classifiers into the channel; the ESP32 becomes a pure transmitter, and the "
       "server sees only events, completing the privacy and energy story.", bold_lead="Phase 3: ")

# ===========================================================================
# 10. LIMITATIONS
# ===========================================================================
h1("10.  Limitations and Honest Expectations")
bullet("a metasurface does not give faces, appearance, or readable detail; keep the one camera for "
       "identity and alarm verification.", bold_lead="Not a camera: ")
bullet("the in-channel classifier is linear; rich, deep models still need digital computation "
       "somewhere.", bold_lead="Linear ceiling: ")
bullet("the paper's accuracy gains were measured on image and gesture data; our through-wall CSI "
       "activity task must be validated separately.", bold_lead="Domain transfer unproven: ")
bullet("larger models cost more transmission time, and moving people force frequent recalibration.", bold_lead="Latency and mobility: ")
bullet("a metasurface adds cost; it must beat the simpler alternative of adding another cheap link.", bold_lead="Cost: ")
note_box("The safe reading is that MetaAI's software-side techniques (noise-aware training, timing "
         "robustness, principled fusion) can be adopted now with high confidence, while the metasurface "
         "hardware is a genuine but unproven frontier for through-wall home sensing, and should be "
         "treated as a research phase with its own validation.")

# ===========================================================================
# 11. NOVELTY
# ===========================================================================
h1("11.  Novelty: What Would Be Publishable")
para("The combination is not just an engineering convenience; it opens research questions that "
     "neither source document answers.")
bullet("MetaAI was demonstrated on line-of-sight image and gesture classification. Applying a "
       "metasurface neural network to through-wall human-activity CSI is, to our knowledge, "
       "unexplored and directly extends both works.", bold_lead="Over-the-air compute for through-wall sensing: ")
bullet("using the enrolment camera to generate labels and then compiling the learned student onto a "
       "metasurface unifies the idea's teacher-student scheme with MetaAI's train-then-discretise "
       "workflow into one pipeline.", bold_lead="Camera-supervised physical distillation: ")
bullet("one surface that alternates between illuminating a room, computing a classifier, and fusing "
       "sensors within one deployment is a richer use than the paper's single-role demonstrations.", bold_lead="Joint sense-compute-fuse surface: ")
bullet("classifying in the channel and emitting only coarse events is a concrete, measurable answer "
       "to the Wi-Fi-identification privacy hazard, a contribution of independent interest.", bold_lead="Privacy-by-construction sensing: ")
idea_box("Each of these is a defensible research contribution that sits squarely between the two "
         "documents already in the programme, which is exactly the gap this study was asked to find.")

# ===========================================================================
# 12. REFERENCES
# ===========================================================================
h1("12.  References and Further Reading")
refs = [
    "C. Feng, S. Liang, C. Li, G. Zhao, B. Jing, Y. Xie, X. Chen. Enabling Over-the-Air AI for Edge "
    "Computing via Metasurface-Driven Physical Neural Networks (MetaAI). ACM SIGCOMM, 2025.",
    "Camera-Enrolled, Wireless-Tracked Identity \u2014 A Catalogue of Concrete Project Options "
    "(companion idea document).",
    "Wireless Perception: Reducing Cameras with Biometric-to-Wireless Handoff (companion feasibility "
    "study).",
    "F. Pace. ESPectre \u2014 Wi-Fi CSI motion detection for Home Assistant (open source). "
    "github.com/francescopace/espectre (companion guide).",
    "IEEE 802.11bf Task Group \u2014 WLAN Sensing standardisation (2020- ).",
    "D. Halperin, W. Hu, A. Sheth, D. Wetherall. Tool release: gathering 802.11n traces with channel "
    "state information. ACM SIGCOMM CCR, 2011.",
    "F. Adib, D. Katabi. See Through Walls with WiFi. ACM SIGCOMM, 2013.",
    "M. Zhao et al. Through-Wall Human Pose Estimation Using Radio Signals (RF-Pose). CVPR, 2018.",
    "F. Wang et al. Person-in-WiFi: Fine-grained Person Perception using WiFi. ICCV, 2019.",
    "J. Geng, D. Huang, F. De la Torre. DensePose From WiFi. arXiv:2301.00250, 2023.",
    "Y. Zheng et al. Widar3.0: Zero-Effort Cross-Domain Gesture Recognition with Wi-Fi. IEEE TPAMI / "
    "MobiSys.",
    "M. Di Renzo et al. Reconfigurable Intelligent Surfaces: principles and opportunities. IEEE "
    "surveys (2020- ).",
    "Karlsruhe Institute of Technology. Person identification from commodity Wi-Fi (2026) \u2014 "
    "reported near-perfect accuracy; a privacy caution.",
]
for rtext in refs:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(rtext)
    r.font.size = Pt(10)

closing = doc.add_paragraph()
closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("This note is a research and design study prepared to answer how the MetaAI paper "
                    "can be combined with the reduce-cameras Wi-Fi idea. Figures and rendered equations "
                    "were generated for this document. Performance figures attributed to MetaAI are "
                    "drawn from the SIGCOMM 2025 paper and should be validated in our own through-wall "
                    "CSI domain before being relied upon.")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = SLATE_LT

# ===========================================================================
# BUILD
# ===========================================================================
if __name__ == "__main__":
    doc.save(OUTPUT)
    print("saved", OUTPUT)
