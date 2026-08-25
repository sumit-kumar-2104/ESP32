"""Build Document 2 of the literature survey:
"Computing in the Wireless Channel: The MetaAI Idea, Its Ancestors, and What
We Can Actually Do".

Focus: the SIGCOMM 2025 MetaAI paper (metasurface-driven physical neural
network / over-the-air AI), the real lineage of work it builds on, an honest
feasibility and hardware assessment, and concrete next steps. One page per
paper, plain-language + technical, real citations only. No emojis.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
OUTPUT = os.path.join(os.path.dirname(HERE), "Survey_2_Computing_in_the_Channel_MetaAI.docx")

INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
PLUM = RGBColor(0x7A, 0x64, 0x9A)
RUST = RGBColor(0xA8, 0x5A, 0x4A)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_PLUM = "F1EDF6"
SHADE_AMBER = "FBF2E7"
SHADE_GREY = "F2F3F5"
HEADER_FILL = "DCE6EF"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = INK
pf = normal.paragraph_format; pf.space_after = Pt(8); pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(size)); e.set(qn("w:space"), "8"); e.set(qn("w:color"), color)
    pbdr.append(e); pPr.append(pbdr)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)


def para(text, italic=False, size=11, space_after=8):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def field(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + "  "); r.bold = True; r.font.color.rgb = SLATE
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def layman_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREEN); _border(p, color="AFC8B2")
    r = p.add_run("In plain terms:  "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def note_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREY); _border(p, color="C9CDD3")
    r = p.add_run(text); r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


def caution_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_AMBER); _border(p, color="D9B98A")
    r = p.add_run("Honest caveat:  "); r.bold = True; r.font.color.rgb = RUST
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def make_table(headers, rows, widths=None, fs=9.5, first_bold=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""; p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL); tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; p = cells[i].paragraphs[0]
            r = p.add_run(val); r.font.size = Pt(fs); r.font.color.rgb = INK
            if i == 0 and first_bold:
                r.bold = True; r.font.color.rgb = SLATE
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA"); tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


_paper = {"n": 0}


def paper(title):
    _paper["n"] += 1
    doc.add_heading(f"Paper {_paper['n']}.  {title}", level=2)


# ===========================================================================
# TITLE PAGE
# ===========================================================================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Computing in the Wireless Channel"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(60); title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("The MetaAI Idea, Its Real Ancestors, and What We Can Actually Do")
r.font.size = Pt(15); r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("A focused study of \u201cEnabling Over-the-Air AI for Edge Computing via Metasurface-Driven "
                 "Physical Neural Networks\u201d (SIGCOMM 2025) \u2014 the line of work it grew from, an honest "
                 "feasibility and hardware assessment, and concrete directions. Real citations only.")
r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = INK

line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  Document 2 of 2  \u00b7  Companion to \u201cWireless Perception\u201d  \u2014")
r.font.size = Pt(11); r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(24)

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Multi-Modal Perception (MPL) programme  \u00b7  Prepared as a reading and decision guide")
r.font.size = Pt(10); r.font.color.rgb = SLATE_LT
doc.add_page_break()

# ===========================================================================
# HOW TO READ
# ===========================================================================
h1("How to Read This Document")
para("This is the second of two survey documents. Document 1 covered wireless perception (seeing people "
     "with radio). This one covers a different and more radical idea: making the wireless signal do the "
     "AI computation on its way to the receiver, so that transmitting the data and computing on it become "
     "one step. The MetaAI paper is the centrepiece; the other papers are its genuine ancestors and "
     "cousins, so you can see where the idea came from and what is proven versus speculative.")
para("Layout is the same per paper: what it is, plain-language explanation, method/data/result, why it "
     "matters, and an honest caveat, with a real citation. After the papers comes the part you asked for "
     "specifically: how feasible the MetaAI idea is, what hardware it needs, what has happened around it, "
     "and what you could realistically do.")
note_box("Scope honesty. MetaAI is very new (2025). There is not yet a large body of follow-up work "
         "citing it, so this document does not pretend there is. Instead it maps the established lineage "
         "the paper itself builds on (optical and metasurface physical neural networks, over-the-air "
         "computation), which is the accurate way to understand \u201cwhat work has been done on this idea.\u201d")

# ===========================================================================
# 1. THE ONE-PARAGRAPH IDEA
# ===========================================================================
h1("1.  The Idea in One Page")
para("A neural-network layer is, at heart, multiply-and-add: multiply each input by a learned weight and "
     "add the results. A wireless channel already multiplies a signal (it scales and rotates it) and adds "
     "signals when they arrive together. So if you could set exactly how the channel multiplies the "
     "signal at each instant, the channel itself would compute the neural network as the signal passes "
     "through it. A metasurface \u2014 a flat panel of hundreds of tiny, electronically controllable cells \u2014 "
     "is the device that lets you set it.")
layman_box("Normally a sensor mails its raw data to a server and the server thinks about it. MetaAI "
           "installs a programmable \u201csmart mirror\u201d in the room; as the sensor's signal bounces off the "
           "mirror on the way to the receiver, the mirror reshapes it so the receiver gets the answer "
           "(\u201cthis is a cat\u201d) instead of the raw picture. The mirror can be reprogrammed to hold "
           "whatever the AI model learned.")
para("Two things make this hard, and MetaAI's contribution is solving both on ordinary wireless hardware: "
     "(1) wireless sends data one piece at a time, but a network wants all inputs at once \u2014 solved by "
     "noticing a linear layer can be computed one term at a time and accumulated; and (2) the multiply "
     "must happen in the wave itself \u2014 solved by programming the metasurface so its channel equals the "
     "trained weight at each instant.")

# ===========================================================================
# 2. THE PAPERS
# ===========================================================================
h1("2.  The Papers, One Per Page")

# ---- Paper 1: MetaAI (the target) ----
paper("MetaAI \u2014 Over-the-Air AI via a Metasurface-Driven Physical Neural Network (SIGCOMM 2025) [THE TARGET PAPER]")
field("What it is.", "The paper your professor pointed you to. It turns the wireless channel into a "
      "single-layer neural network by programming one metasurface, so a small device only has to transmit "
      "while the network is computed over the air and the receiver reads out the classification result.")
layman_box("It is the first design to run a complete, end-to-end neural network inside an ordinary "
           "wireless link using just one reprogrammable panel and a normal transmitter. Earlier systems "
           "needed exotic setups; MetaAI works with commodity-style radios.")
field("How it works.", "The device streams the input one symbol per time slot. The metasurface is set so "
      "that at each time slot its channel response equals the trained weight for that input. The receiver "
      "accumulates the products, so a matrix-vector multiply (a linear layer) is computed physically as "
      "the signal travels; the magnitude of the complex sum gives the class score and acts as the "
      "nonlinearity. It adds practical fixes: multipath cancellation (vary the surface within a symbol so "
      "room echoes average out), timing robustness (CDFA: coarse energy detection plus training-time "
      "error injection), noise-aware training, frequency/antenna parallelism, and multi-sensor fusion by "
      "time-sharing the same surface.")
field("Data and result.", "A physical prototype (two metasurfaces: dual-band 2.4/5 GHz and single-band "
      "3.5 GHz; 16x16 = 256 meta-atoms; 2-bit phase control; STM32 controller; USRP X310 radios) was "
      "tested on six image/gesture datasets, reaching about 82.8 percent average accuracy (up to about "
      "89.8 percent), within roughly 7 percent of its own digital simulation, with multi-sensor fusion "
      "adding up to about 27 percent, across corridor, lab and office environments including "
      "non-line-of-sight.")
field("Why it matters.", "It unifies communication and computation into one physical step. For always-on, "
      "low-power, privacy-sensitive edge sensing it is a genuinely new direction: the device never runs "
      "the AI, and the server sees results, not raw data.")
caution_box("It currently implements only linear networks (no deep nonlinear models like Transformers). "
            "Larger models mean more transmissions and higher latency. Accuracy is capped by metasurface "
            "resolution (atom count and 2-bit depth). Movement forces recalibration. It is a lab prototype "
            "with specialised metasurfaces and software-defined radios \u2014 not something an ESP32 can do "
            "today.")
field("Citation.", "C. Feng, S. Liang, C. Li, G. Zhao, B. Jing, Y. Xie, X. Chen. \u201cEnabling Over-the-Air "
      "AI for Edge Computing via Metasurface-Driven Physical Neural Networks.\u201d ACM SIGCOMM 2025. "
      "DOI: 10.1145/3718958.3750474.")

# ---- Paper 2: AirNN ----
paper("AirNN \u2014 Over-the-Air Convolution via Reconfigurable Intelligent Surfaces (the closest predecessor)")
field("What it is.", "The most direct ancestor of MetaAI: it performs a convolution (the core operation of "
      "a CNN) over the air using reconfigurable intelligent surfaces (RIS, another name for programmable "
      "metasurfaces), and demonstrates it for a neural-network inference task.")
layman_box("AirNN showed for the first time that you can make radio reflections carry out a convolution \u2014 "
           "one of the two key operations in deep learning \u2014 by shaping the environment with metasurfaces. "
           "MetaAI later generalised this from a single convolution to a complete end-to-end network on an "
           "ordinary link.")
field("How it works.", "It treats each neuron's weights as a channel impulse response that corresponds to "
      "a realisable filter, then engineers each such response with an RIS; the reflected signals combine "
      "at the receiver to produce the convolution output. It is a proof-of-concept: the over-the-air "
      "convolution is demonstrated experimentally, and the full CNN accuracy is then validated in "
      "simulation for a modulation-classification task.")
field("Data and result.", "Experimental demonstration of over-the-air convolution plus simulation of the "
      "resulting CNN. It establishes feasibility rather than a deployable product.")
field("Why it matters.", "It is the scientific bridge from \u201cRIS can shape signals\u201d to \u201cRIS can compute "
      "a neural-network operation.\u201d Reading it makes MetaAI's leap much clearer, and it is the paper "
      "MetaAI positions itself against.")
caution_box("It needed a more complex, multi-surface / relay-style arrangement and validated the full "
            "network mostly in simulation. It is a stepping stone, not an end system.")
field("Citation.", "S. Garcia Sanchez, G. Reus Muns, C. Bocanegra, Y. Li, U. Muncuk, Y. Naderi, Y. Wang, "
      "S. Ioannidis, K. R. Chowdhury. \u201cAirNN: Neural Networks with Over-the-Air Convolution via "
      "Reconfigurable Intelligent Surfaces.\u201d arXiv:2202.03399, 2022 (later IEEE/ACM Transactions on "
      "Networking).")

# ---- Paper 3: D2NN ----
paper("Diffractive Deep Neural Networks (D2NN) \u2014 All-Optical Machine Learning (the founding idea)")
field("What it is.", "The foundational demonstration that a physical wave passing through fixed patterned "
      "layers can perform a full neural network's computation \u2014 in this case with light through 3D-printed "
      "plates.")
layman_box("This 2018 Science paper is where \u201clet the physics do the maths\u201d became real. Light shines "
           "through a stack of specially shaped transparent layers; the way the light bends and interferes "
           "as it passes is exactly a neural-network calculation, and the answer appears as a bright spot "
           "at the output. No electricity is used for the computation itself \u2014 the wave does it at the "
           "speed of light.")
field("How it works.", "Each layer is an array of tiny features that shift the phase/amplitude of the "
      "passing wave; diffraction spreads and recombines the light so that stacked layers implement "
      "successive linear transforms. The layer patterns are trained in a computer, then physically "
      "fabricated (here for terahertz waves via 3D printing).")
field("Data and result.", "Demonstrated on classic benchmarks such as handwritten digits (MNIST) and "
      "fashion images, achieving accuracies in the low-to-mid 90s percent for digits, all optically after "
      "fabrication.")
field("Why it matters.", "It is the intellectual root of physical neural networks, including MetaAI's "
      "microwave version. It proves the core claim that wave propagation equals neural computation.")
caution_box("The original D2NN layers are fixed once printed (not reprogrammable), it uses optics rather "
            "than radio, and it implements linear layers (plus a measurement nonlinearity). It is a "
            "principle demonstrator, not a communication system.")
field("Citation.", "X. Lin, Y. Rivenson, N. T. Yardimci, M. Veli, Y. Luo, M. Jarrahi, A. Ozcan. "
      "\u201cAll-optical machine learning using diffractive deep neural networks.\u201d Science 361, 1004\u20131008 "
      "(2018).")

# ---- Paper 4: Coherent nanophotonic circuits ----
paper("Deep Learning with Coherent Nanophotonic Circuits (matrix multiply with light on a chip)")
field("What it is.", "A photonic chip that performs the matrix multiplications of a neural network using "
      "beams of light routed through interferometers \u2014 another founding demonstration of computing with "
      "waves, but on-chip and reconfigurable.")
layman_box("Where D2NN used free-space light, this uses light inside tiny channels on a silicon chip. By "
           "tuning little optical components, the chip multiplies vectors by matrices with light instead "
           "of transistors, promising very fast, low-energy AI math.")
field("How it works.", "A mesh of Mach-Zehnder interferometers implements the linear (matrix) part of a "
      "network; the settings realise the trained weights. A nonlinearity is applied separately. The "
      "authors ran a small classification task through the photonic hardware.")
field("Data and result.", "Demonstrated on a vowel-recognition task, showing a real optical chip can "
      "carry out neural-network matrix operations with accuracy approaching the digital baseline.")
field("Why it matters.", "It is the reconfigurable, chip-scale sibling of D2NN and part of the same family "
      "of \u201canalog physical computation\u201d that MetaAI joins in the radio domain. It shows the weights can "
      "be made tunable, not just fixed.")
caution_box("It is photonics, not wireless, and small-scale; scaling and integrating the nonlinearity "
            "remain hard. It informs the concept behind MetaAI rather than being directly reusable for a "
            "Wi-Fi project.")
field("Citation.", "Y. Shen, N. C. Harris, S. Skirlo, M. Prabhu, T. Baehr-Jones, M. Hochberg, X. Sun, "
      "S. Zhao, H. Larochelle, D. Englund, M. Soljacic. \u201cDeep learning with coherent nanophotonic "
      "circuits.\u201d Nature Photonics 11, 441\u2013446 (2017).")

# ---- Paper 5: Programmable metasurface D2NN ----
paper("Programmable Diffractive Deep Neural Network on a Coding-Metasurface Array (reprogrammable, in microwave)")
field("What it is.", "A microwave metasurface that performs a diffractive neural network and, importantly, "
      "can be reprogrammed on the fly \u2014 the closest precursor to MetaAI in both frequency band (microwave) "
      "and reconfigurability.")
layman_box("This takes the D2NN idea out of fixed 3D-printed optics and into an electronically "
           "reprogrammable microwave panel. You can change what network it computes by sending it new "
           "control codes, without rebuilding anything. That reprogrammability is exactly what MetaAI "
           "relies on.")
field("How it works.", "A digital-coding metasurface array (cells switch between discrete states) is "
      "arranged so that microwaves passing/reflecting through it implement a diffractive network; a "
      "controller loads different weight patterns in real time, enabling tasks to be switched or retrained "
      "onto the same hardware.")
field("Data and result.", "Demonstrated real-time programmable classification tasks on the metasurface "
      "hardware in the microwave regime (the authors report recognition tasks such as handwritten digits "
      "and coded patterns).")
field("Why it matters.", "It proves the two ingredients MetaAI needs \u2014 microwave operation and live "
      "reprogramming of a metasurface into network weights \u2014 already work. MetaAI's novelty is doing this "
      "over a normal communication link and unifying it with data transmission.")
caution_box("It is still a specialised lab metasurface system focused on the computation itself, not on "
            "riding an ordinary wireless link with a commodity transmitter (which is MetaAI's step).")
field("Citation.", "C. Liu, Q. Ma, Z. J. Luo, Q. R. Hong, Q. Xiao, H. C. Zhang, L. Miao, W. M. Yu, "
      "Q. Cheng, L. Li, T. J. Cui. \u201cA programmable diffractive deep neural network based on a "
      "digital-coding metasurface array.\u201d Nature Electronics 5, 113\u2013122 (2022).")

# ---- Paper 6: Over-the-air computation (AirComp) ----
paper("Over-the-Air Computation (AirComp) for Federated Learning (the \u201caddition for free\u201d line)")
field("What it is.", "A body of work that uses the natural adding-up of overlapping radio signals to "
      "compute a sum (an average) directly in the air \u2014 used to speed up federated learning aggregation. "
      "It is the \u201caddition\u201d half of the over-the-air computing story that MetaAI completes with "
      "\u201cmultiplication.\u201d")
layman_box("If many devices transmit at the same time, their signals literally add together on the way to "
           "the receiver. AirComp exploits this so that, instead of collecting each device's numbers and "
           "summing them, the sum just appears at the receiver. It gives you addition for free; what was "
           "missing was cheap multiplication \u2014 which is where metasurfaces come in.")
field("How it works.", "Devices pre-scale their transmissions so that the superimposed waveform at the "
      "receiver equals the desired weighted sum (for example, the average of local model updates in "
      "federated learning), turning many uploads into one aggregated measurement.")
field("Data and result.", "A large literature shows this speeds up and scales distributed model "
      "aggregation, trading some numerical precision for large communication savings (theory plus "
      "simulations, and some testbeds).")
field("Why it matters.", "It explains why \u201caddition in the channel\u201d was already understood, and frames "
      "MetaAI's real contribution: adding programmable multiplication (via the metasurface) so the channel "
      "can do a full neural-network layer, not just a sum.")
caution_box("Classic AirComp does aggregation (sums), not general neural-network inference, and needs "
            "careful power control and synchronisation. It is background context, not a sensing system.")
field("Citation.", "Representative: K. Yang, T. Jiang, Y. Shi, Z. Ding. \u201cFederated Learning via "
      "Over-the-Air Computation.\u201d IEEE Transactions on Wireless Communications, 2020. (One entry point "
      "into a broad AirComp literature.)")

# ===========================================================================
# 3. WHAT HAS BEEN DONE AROUND THE IDEA
# ===========================================================================
h1("3.  What Has Actually Been Done On This Idea (An Honest Map)")
para("You asked specifically what work exists on the MetaAI idea and what has happened since it was "
     "published. Here is the honest picture.")
bullet("MetaAI (2025) is very recent. As of now there is not a large set of peer-reviewed papers that "
       "extend MetaAI specifically. Anyone claiming a big follow-up literature would be exaggerating.",
       "It is new. ")
bullet("the idea sits on a well-established lineage: diffractive optical neural networks (D2NN, 2018), "
       "photonic matrix-multiply chips (2017), reprogrammable microwave metasurface D2NNs (2022), "
       "over-the-air convolution with RIS (AirNN, 2022), and over-the-air computation for aggregation "
       "(AirComp). That lineage is real and is what you should read to understand the field.",
       "But the lineage is deep. ")
bullet("the same communities work on RIS/metasurfaces for communication and for wireless sensing "
       "(imaging, localisation). Combining computation, communication and sensing on one surface is an "
       "active, open research theme \u2014 which is exactly where your MPL angle could contribute.",
       "The surrounding field is active. ")
note_box("So the accurate framing for your professor is: MetaAI is a fresh capstone on a mature line of "
         "physical/optical/over-the-air computing. The opportunity is not to chase a crowded follow-up "
         "field, but to connect this computation idea to the wireless-perception tasks in Document 1.")

# ===========================================================================
# 4. FEASIBILITY AND HARDWARE
# ===========================================================================
h1("4.  Feasibility and Hardware: What Would It Actually Take?")
para("A blunt, honest assessment of building on the MetaAI idea, at three levels of ambition.")

h2("4.1  What MetaAI itself needed (the full thing)")
make_table(
    ["Component", "What the paper used", "Rough accessibility for you"],
    [
        ["Metasurface", "Custom 16x16 (256-cell) 2-bit reconfigurable panels, dual-band + 3.5 GHz",
         "Hard: custom-fabricated PCB metasurface with PIN diodes; not off-the-shelf"],
        ["Radios", "USRP X310 software-defined radios (Tx and Rx)",
         "Expensive: several thousand USD each; sometimes available in RF labs"],
        ["Controller", "STM32 microcontroller driving the cells via shift registers",
         "Easy and cheap once the panel exists"],
        ["Compute", "GPU/PC to train the complex-valued network offline",
         "Easy: a normal training PC"],
        ["Expertise", "RF/antenna design, SDR programming, metasurface control",
         "Significant: this is an RF-systems project, not just ML"],
    ],
    widths=[1.3, 2.7, 2.3],
)
caution_box("Reproducing MetaAI end-to-end is a serious hardware-lab effort (custom metasurface plus SDRs "
            "plus RF expertise). It is realistic only if your lab already has, or will fund, a "
            "reconfigurable metasurface and software-defined radios. It is not an ESP32-scale build.")

h2("4.2  A middle path (metasurface, but simpler)")
para("If the lab has or can buy a commercial reconfigurable intelligent surface (RIS) development kit and "
     "one or two SDRs, a realistic scaled-down project is to reproduce a single over-the-air linear "
     "classifier (as in AirNN/MetaAI) on a small dataset, and study one specific robustness issue \u2014 for "
     "example how timing error or multipath degrades accuracy and how the training-time fixes help. This "
     "is a focused, publishable contribution without inventing new hardware.")

h2("4.3  A low-cost path you can start now (simulation + ESP32 sensing)")
para("Most realistic for a remote intern with an ESP32 and a laptop:")
numbered("simulate the MetaAI pipeline in software \u2014 model the channel-as-weights, the 2-bit metasurface, "
         "timing/multipath/noise \u2014 and reproduce the accuracy-versus-impairment curves. This needs no "
         "special hardware and teaches the whole idea deeply.", "")
numbered("in parallel, keep building the real Wi-Fi sensing from Document 1 on the ESP32 (presence, "
         "gesture, activity). This is where your actual data and demos come from.", "")
numbered("the research bridge: study, in simulation, whether an over-the-air linear front-end could "
         "pre-process CSI sensing features before the tiny device transmits \u2014 i.e., combine the "
         "computation idea with the perception task. Frame it honestly as an exploratory feasibility "
         "study.", "")
layman_box("Short version: fully rebuilding MetaAI needs a metasurface and lab radios you probably do not "
           "have yet. But you can learn and contribute right now by simulating its maths faithfully and by "
           "growing your ESP32 sensing work \u2014 then connect the two ideas on paper and in simulation, which "
           "is itself a legitimate research direction.")

# ===========================================================================
# 5. WHAT WE CAN DO
# ===========================================================================
h1("5.  Concrete, Honest Directions You Could Propose")
make_table(
    ["Direction", "What you would do", "Hardware needed", "Honest verdict"],
    [
        ["A. Faithful MetaAI simulator",
         "Reimplement the channel-as-weights pipeline; reproduce impairment curves.",
         "Laptop + GPU", "Very feasible; great learning + a solid write-up"],
        ["B. Scaled over-the-air classifier",
         "One linear classifier over a real RIS + SDR on a small dataset.",
         "RIS kit + 1-2 SDRs", "Feasible only if lab has the RF gear"],
        ["C. ESP32 Wi-Fi sensing (Doc 1 ladder)",
         "Presence/gesture/activity, camera-taught, cross-domain tested.",
         "ESP32 + webcam", "Most feasible; your real data engine"],
        ["D. Computation-meets-perception study",
         "Simulate an over-the-air front-end for CSI sensing features.",
         "Laptop", "Feasible as exploratory research; novel angle"],
        ["E. Robustness/generalisation focus",
         "Systematically study domain shift and impairments across A-D.",
         "Varies", "High value; reviewers love honest generalisation studies"],
    ],
    widths=[1.5, 2.3, 1.1, 1.7],
)
para("Recommended sequencing: start with C (you are already doing it) for real results, do A in parallel "
     "to master the MetaAI idea cheaply, and let D/E be the ambitious research narrative that ties both "
     "survey documents together. Only pursue B if and when the lab commits RF hardware.")

# ===========================================================================
# 6. SUMMARY
# ===========================================================================
h1("6.  One-Page Summary for Your Meeting")
bullet("MetaAI makes the wireless channel compute a neural network by programming a metasurface into the "
       "network's weights; the device only transmits and the receiver reads the answer.", "The idea: ")
bullet("it is real, peer-reviewed (SIGCOMM 2025), and works on a lab prototype, but only for linear "
       "networks and with custom metasurfaces plus software-defined radios.", "The status: ")
bullet("D2NN (2018), photonic circuits (2017), reprogrammable microwave metasurface D2NN (2022), AirNN "
       "(2022) and AirComp form the genuine lineage; MetaAI is a recent capstone with little follow-up "
       "yet.", "The lineage: ")
bullet("full reproduction needs a metasurface and SDRs (hard); but a faithful simulation plus your ESP32 "
       "sensing is a realistic, honest starting point that still contributes.", "The feasibility: ")
bullet("combine the computation idea (this document) with the perception tasks (Document 1): start with "
       "ESP32 sensing for real data, simulate MetaAI to master it, and study whether the two can meet.",
       "The move: ")

closing = doc.add_paragraph(); closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("End of Document 2. Every paper listed is real and cited so it can be verified. Where "
                    "MetaAI is new and follow-up work is thin, this document says so plainly rather than "
                    "inventing a literature. Feasibility is assessed honestly against the hardware you "
                    "actually have.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE_LT

doc.save(OUTPUT)
print("wrote", OUTPUT)
