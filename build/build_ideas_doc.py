"""Build the ideas-catalogue Word document:
"Camera-Enrolled, Wireless-Tracked Identity - A Catalogue of Concrete Project
Options" so the user and their professor can finalise one.

Central idea: use a few cameras to enrol facial/biometric signatures per person,
build a pattern, then use wireless perception to track/monitor those identities in
camera-free areas. Includes a catalogue of concrete ideas, each with guided
execution, requirements, feasibility, challenges and prior work; a decision matrix;
a recommendation; and an ethics section. No emojis are used.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "assets_ideas")
OUTPUT = os.path.join(os.path.dirname(HERE), "Camera_to_Wireless_Identity_Ideas.docx")

INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
PLUM = RGBColor(0x7A, 0x64, 0x9A)
RUST = RGBColor(0xA8, 0x5A, 0x4A)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_PLUM = "F1EDF6"
SHADE_AMBER = "FBF2E7"
SHADE_GREY = "F2F3F5"
HEADER_FILL = "DCE6EF"
TAG_FILL = "E7EEF4"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = INK
pf = normal.paragraph_format; pf.space_after = Pt(8); pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(size)); e.set(qn("w:space"), "8"); e.set(qn("w:color"), color)
    pbdr.append(e); pPr.append(pbdr)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def para(text, italic=False, size=11, space_after=8):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def runs_para(parts, space_after=8):
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        r = p.add_run(text); r.bold = bold; r.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def field(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + "  "); r.bold = True; r.font.color.rgb = SLATE
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def layman_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREEN); _border(p, color="AFC8B2")
    r = p.add_run("In plain terms:  "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def sci_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_BLUE); _border(p, color="9DB6CC")
    r = p.add_run("Technical detail:  "); r.bold = True; r.font.color.rgb = SLATE
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def ethics_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_AMBER); _border(p, color="D9B98A")
    r = p.add_run("Ethics and privacy:  "); r.bold = True; r.font.color.rgb = RUST
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def note_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREY); _border(p, color="C9CDD3")
    r = p.add_run(text); r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


_eqn = {"n": 0}


def equation(img, label=True):
    _eqn["n"] += 1
    from PIL import Image
    path = os.path.join(ASSETS, img)
    with Image.open(path) as im:
        w_in = im.width / 200.0; h_in = im.height / 200.0
    max_w = 5.9
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if w_in > max_w:
        run.add_picture(path, width=Inches(max_w))
    else:
        run.add_picture(path, height=Inches(max(min(h_in, 0.55), 0.26)))
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    if label:
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"( {_eqn['n']} )"); r.font.size = Pt(9); r.font.color.rgb = SLATE_LT
        cap.paragraph_format.space_after = Pt(8)
    return _eqn["n"]


_fig = {"n": 0}


def figure(img, caption, width=6.2):
    _fig["n"] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(width))
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {_fig['n']}.  "); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE
    r2 = cap.add_run(caption); r2.font.size = Pt(9.5); r2.font.color.rgb = INK
    cap.paragraph_format.space_after = Pt(12)


def variable_table(rows):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(["Symbol", "What it means"]):
        hdr[i].text = ""; p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL); tcPr.append(shd)
    for ri, (sym, mean) in enumerate(rows):
        cells = t.add_row().cells
        cells[0].text = ""; p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(sym); r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = SLATE
        cells[1].text = ""; p1 = cells[1].paragraphs[0]
        r1 = p1.add_run(mean); r1.font.size = Pt(9.5); r1.font.color.rgb = INK
        if ri % 2 == 1:
            for c in cells:
                tcPr = c._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA"); tcPr.append(shd)
    t.columns[0].width = Inches(1.3); t.columns[1].width = Inches(4.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def make_table(headers, rows, widths=None, fs=9.5, first_bold=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""; p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL); tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; p = cells[i].paragraphs[0]
            r = p.add_run(val); r.font.size = Pt(fs); r.font.color.rgb = INK
            if i == 0 and first_bold:
                r.bold = True; r.font.color.rgb = SLATE
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA"); tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def scorecard(effort, hardware, data, novelty, robustness, best_for):
    make_table(
        ["Build effort", "Hardware cost", "Data / labeling", "Novelty", "Expected robustness", "Best used as"],
        [[effort, hardware, data, novelty, robustness, best_for]],
        widths=[1.05, 1.05, 1.15, 0.9, 1.25, 1.6], fs=8.8, first_bold=False,
    )


_idea = {"n": 0}


def idea(title):
    _idea["n"] += 1
    doc.add_heading(f"Idea {_idea['n']}:  {title}", level=2)


# ===========================================================================
# TITLE
# ===========================================================================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Camera-Enrolled, Wireless-Tracked Identity"); r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(56); title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Catalogue of Concrete Project Options for Reducing Cameras with Biometric-to-Wireless Handoff")
r.font.size = Pt(14.5); r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Use a few cameras to learn each person's biometric signature, then let wireless perception "
                 "recognise and follow that person through rooms with no cameras. Eight concrete ideas, each with "
                 "guided execution, requirements, feasibility and prior work \u2014 so one can be finalised.")
r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = INK

line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  Ideas Catalogue and Decision Aid  \u2014"); r.font.size = Pt(11); r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(24)

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Multi-Modal Perception (MPL) programme  \u00b7  Grounded in XModal-ID, RF-ReID, WhoFi, Vi-Fi and related work")
r.font.size = Pt(10); r.font.color.rgb = SLATE_LT
doc.add_page_break()

# ===========================================================================
# CONTENTS
# ===========================================================================
h1("Contents")
toc_p = doc.add_paragraph(); run = toc_p.add_run()
fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
t_run = OxmlElement("w:t"); t_run.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
for e in (fldChar, instr, fldChar2, t_run, fldChar3): run._r.append(e)
doc.add_page_break()

# ===========================================================================
# 1. THE IDEA AND SHORT ANSWER
# ===========================================================================
h1("1.  The Idea, and the Short Answer")
para("The proposal is to invert the usual heavy-camera setup. Instead of putting a camera in every room, put a "
     "small number of cameras only at a few key points \u2014 an entrance, a hallway \u2014 and use them once, to learn "
     "who each person is. From that footage, build a compact biometric signature per person (face, but also the way "
     "they walk and their body shape). Then hand the tracking over to wireless perception: because radio passes "
     "through walls and reflects off the body, a wireless sensor can recognise that same signature and follow the "
     "person through rooms that have no camera at all.")
runs_para([("The question:  ", True, False),
           ("is it possible to recognise a person with a few cameras, turn that into a per-person pattern, and then "
            "monitor and track that identity wirelessly where there are no cameras?", False, False)])
runs_para([("The short answer:  ", True, False),
           ("yes, in principle, and it has been demonstrated in research \u2014 with important limits. Matching a person "
            "in video to a wireless measurement through a wall was shown by XModal-ID (2019). Identifying people from "
            "Wi-Fi channel data was shown by WhoFi (2025) at about 95 percent on a benchmark, and from radar by "
            "RF-ReID (2020) using body shape that survives clothing changes. The limits: it works well for a small, "
            "known set of people (a household, a team), needs an enrolment step, and is environment-specific. It is a "
            "strong research and thesis direction rather than a plug-and-play product.", False, False)])
figure("fig_handoff.png",
       "The core concept. One camera at an enrolment gate learns each known person's face, gait and body shape and "
       "turns them into a signature. Wireless sensors in camera-free rooms then match the live wireless signature to "
       "that enrolled identity, tracking the same person without any camera in those rooms.", width=6.6)
layman_box("It is like giving each family member a name tag that only the Wi-Fi can read. The camera writes the tag "
           "once at the door by looking at the face and walk; afterwards the Wi-Fi in each room reads the invisible "
           "tag and knows who is there \u2014 no camera needed inside the private rooms.")

# ===========================================================================
# 2. HOW TO READ THIS CATALOGUE
# ===========================================================================
h1("2.  How to Read This Catalogue")
para("Sections 3 gives the shared technical foundation. Section 4 presents eight concrete ideas as cards. Each card "
     "has the same structure so they can be compared directly:")
bullet("what the idea is, in one or two sentences.", "Concept: ")
bullet("the mechanism and the closest proven system.", "How it works: ")
bullet("a numbered, do-this-then-that recipe.", "Execution steps: ")
bullet("hardware, software, data and skills needed.", "Requirements: ")
bullet("a six-dimension scorecard (effort, cost, data need, novelty, robustness, best use).", "Feasibility: ")
bullet("the main risks and how to reduce them.", "Challenges and mitigations: ")
bullet("the peer-reviewed work it builds on.", "Closest prior work: ")
note_box("Ratings in the scorecards are relative guidance (Low / Medium / High) to help prioritise, not precise "
         "measurements. Section 5 collects them into one decision matrix, and Section 6 gives a recommended path.")

# ===========================================================================
# 3. SHARED TECHNICAL FOUNDATION
# ===========================================================================
h1("3.  Shared Technical Foundation")
para("Almost every idea below rests on the same three-phase pattern and the same handful of equations. Understanding "
     "them once makes the whole catalogue easy to read.")
figure("fig_phases.png",
       "The three phases common to these systems: enrol (camera and wireless observe the same person together), "
       "train (learn matched signatures), deploy (wireless recognises the enrolled identity alone).", width=6.6)

h2("3.1  Signatures are embeddings")
para("A signature is a short list of numbers (a vector) produced by a neural network from raw sensor data. Two "
     "measurements of the same person should produce nearby vectors; different people should produce distant ones.")
equation("eq_embed.png")
runs_para([("What it says.  ", True, False),
           ("An encoder f with parameters theta maps a raw input x (a video clip or a wireless measurement) to a "
            "d-dimensional signature vector z.", False, False)])
variable_table([
    ("z", "the signature (embedding) vector for one observation."),
    ("f_theta", "the encoder network; theta are its learned weights."),
    ("x", "the raw input (camera clip, or Wi-Fi / radar measurement)."),
    ("d", "the length of the signature (for example 128 numbers)."),
])

h2("3.2  Recognition is nearest-neighbour matching")
para("At deployment, the live wireless signature is compared to the gallery of enrolled signatures; the closest match "
     "wins.")
equation("eq_match.png")
runs_para([("What it says.  ", True, False),
           ("The predicted identity k-star is the enrolled person whose camera signature is closest (smallest "
            "distance d) to the live wireless signature z_rf.", False, False)])
variable_table([
    ("k*", "the chosen identity (the best match)."),
    ("z_rf", "the live wireless signature to be identified."),
    ("z_cam^k", "the enrolled signature of candidate person k."),
    ("d( , )", "a distance (for example Euclidean or cosine) between signatures."),
])
para("For an open-world setting, one also verifies rather than just ranks: accept the match only if the distance is "
     "below a threshold, otherwise declare \u201cunknown\u201d.")
equation("eq_verify.png")
variable_table([
    ("tau", "the acceptance threshold; below it, the two are judged the same person."),
])

h2("3.3  Training the signatures: metric learning")
para("Encoders are trained so that same-person pairs are pulled together and different-person pairs pushed apart. The "
     "classic objective is the triplet loss.")
equation("eq_triplet.png")
runs_para([("What it says.  ", True, False),
           ("For an anchor a, a positive p (same person) and a negative n (different person), the loss pushes the "
            "anchor-positive distance to be smaller than the anchor-negative distance by at least a margin m. The "
            "plus sign means only positive violations are penalised.", False, False)])
variable_table([
    ("z_a, z_p, z_n", "signatures of the anchor, a same-person sample, and a different-person sample."),
    ("m", "the margin: how much closer the positive must be than the negative."),
    ("[ ]_+", "the hinge: keep the value only when positive, else zero."),
])

h2("3.4  Linking the two modalities: cross-modal alignment")
para("The key trick that lets a camera teach the wireless sensor is to force paired camera and wireless signatures of "
     "the same moment to land at the same place in the signature space.")
equation("eq_align.png")
runs_para([("What it says.  ", True, False),
           ("Train the camera encoder g and the wireless encoder f together so that, for each synchronised pair, the "
            "wireless signature and the camera signature are as close as possible. After training, a wireless "
            "measurement alone lands where the matching person's camera signature would.", False, False)])
variable_table([
    ("f_theta(CSI_i)", "the wireless signature for sample i."),
    ("g_phi(video_i)", "the camera signature for the same moment i."),
    ("theta, phi", "the weights of the wireless and camera encoders."),
])

h2("3.5  Following people across sensors: association")
para("When several people and several sensors are present, deciding which wireless track belongs to which camera "
     "track is an assignment problem, solved by minimising a total matching cost.")
equation("eq_assoc.png")
runs_para([("What it says.  ", True, False),
           ("Choose the assignment A (who matches whom) that minimises the total cost, where each pair's cost is the "
            "negative log-probability that the camera track i and wireless track j are the same person.", False, False)])
variable_table([
    ("A", "the assignment matrix (1 if camera track i is matched to wireless track j)."),
    ("C_ij", "the cost of matching i to j; lower means more likely the same person."),
    ("p(match | ...)", "the probability that the two tracks are the same person, from signatures and timing."),
])
sci_box("These five equations cover the whole catalogue. Different ideas differ mainly in the sensor (Wi-Fi CSI, "
        "FMCW radar, BLE), what the signature captures (gait, body shape, trajectory), and whether identity is "
        "device-free or tied to a phone. The learning machinery (embeddings, metric learning, cross-modal alignment, "
        "association) is shared.")
figure("fig_embedding.png",
       "Why it works: camera signatures of each enrolled person form clusters in a shared space. A live wireless "
       "measurement is encoded into the same space and assigned to the nearest cluster \u2014 here, Person 2.", width=5.6)

# ===========================================================================
# 4. THE IDEAS
# ===========================================================================
h1("4.  The Ideas")

# ---- Idea 1 ----
idea("Vision-Enrolled, Wireless-Tracked Identity (flagship)")
scorecard("Medium-High", "Low-Medium", "Medium", "High", "Medium", "Headline research contribution")
field("Concept.", "Enrol each known person once with a camera (face, gait, body shape), build a per-person signature, "
      "then re-identify and track that person via Wi-Fi or radar in rooms with no camera. This is the user's core "
      "idea in full.")
field("How it works.", "Camera and wireless encoders are aligned into a shared signature space (Section 3.4). At "
      "deployment the wireless signature is matched to the enrolled gallery (Section 3.2). XModal-ID proved the "
      "cross-modal match works even through a wall (matching a person behind a wall to candidate video, with ranked "
      "top-1/2/3 accuracy of 75 / 90 / 97 percent among eight candidates). WhoFi shows Wi-Fi-only re-identification "
      "at about 95.5 percent on a benchmark, and RF-ReID shows radar body-shape signatures robust to clothing.")
field("Execution steps.", "")
numbered("install one camera at an enrolment point and a wireless sensor (Wi-Fi CSI node or radar) covering each "
         "room to be monitored.", "")
numbered("record synchronised camera video and wireless measurements while each resident walks naturally through the "
         "enrolment area and rooms.", "")
numbered("use the camera to auto-label identity (and gait/pose) for each moment; pair each label with the "
         "simultaneous wireless measurement.", "")
numbered("train the wireless encoder with cross-modal alignment and a triplet or contrastive loss so wireless "
         "signatures match the camera gallery.", "")
numbered("deploy wireless-only matching in the camera-free rooms; accept a match only above a confidence threshold, "
         "else mark unknown.", "")
numbered("evaluate with rank-1 accuracy and mean average precision on held-out sessions and, ideally, a different "
         "day or room arrangement.", "")
field("Requirements.", "One camera; a CSI-capable Wi-Fi setup (ESP32-C6/S3 nodes, or an Intel 5300 / Atheros CSI "
      "rig, or an FMCW radar for higher fidelity); a way to time-synchronise camera and wireless; a GPU for "
      "training; a small closed set of people; and consented enrolment.")
field("Challenges and mitigations.", "Environment dependence (mitigate with per-home enrolment and multiple links); "
      "multi-person separation (add sensors, or enrol distinctive gaits); spoofing and robustness (verify with a "
      "threshold, reject unknowns); privacy (keep signatures on-device, enrol only consenting residents).")
field("Closest prior work.", "XModal-ID (UCSB, MobiCom 2019); WhoFi (2025); RF-ReID (MIT, CVPR 2020).")

# ---- Idea 2 ----
idea("Commodity-Wi-Fi Gait Re-ID for a Household (de-risked core)")
scorecard("Low-Medium", "Low", "Low-Medium", "Medium", "Medium-High", "Fast, reliable first result")
field("Concept.", "A lighter version of Idea 1: identify among a small set of known housemates (2 to 5) by the way "
      "they walk, using only commodity Wi-Fi CSI. The camera is used solely to label the enrolment walks.")
field("How it works.", "Each person's gait produces a distinctive Doppler / CSI spectrogram; a compact classifier or "
      "embedding recognises which housemate is walking. WiWho reported 80 to 92 percent among 2 to 6 people; WifiU "
      "and WhoFi extend this with spectrograms and transformers.")
field("Execution steps.", "")
numbered("place one CSI node per room (or a few along a corridor).", "")
numbered("have each housemate walk a set path several times; label each walk (by camera or by hand).", "")
numbered("extract gait features (CSI amplitude spectrograms) and train a small classifier or embedding.", "")
numbered("deploy per-room recognition among the enrolled housemates; report confidence.", "")
field("Requirements.", "ESP32 / CSI hardware, one camera or manual labels for enrolment, a laptop for training. No "
      "GPU strictly required for small sets.")
field("Challenges and mitigations.", "Works only for a small known set and mainly while walking (accept graceful "
      "\u201cunknown / no-ID\u201d when uncertain); recalibrate after big furniture changes.")
field("Closest prior work.", "WiWho (IPSN 2016); WifiU (2016); WhoFi (2025).")

# ---- Idea 3 ----
idea("Cross-Modal Distillation for Pose and Activity in Blind Rooms")
scorecard("Medium", "Low-Medium", "Medium", "Medium", "Medium", "Complement to Idea 1 / 2")
field("Concept.", "Rather than identity, teach the wireless sensor to output pose or activity (walking, sitting, "
      "falling) in rooms with no camera, by distilling knowledge from a camera during a short overlap.")
field("How it works.", "The camera generates pose/activity labels; a wireless model is trained to predict them "
      "(teacher-student). RF-Pose recovered through-wall skeletons this way; DensePose From WiFi mapped Wi-Fi to a "
      "dense body surface; Person-in-WiFi produced keypoints.")
field("Execution steps.", "")
numbered("overlap a camera and the wireless sensor on the same scene during a collection phase.", "")
numbered("run an off-the-shelf pose/activity model on the video to create labels.", "")
numbered("train the wireless model to reproduce those labels from wireless input.", "")
numbered("remove the camera; the wireless model reports pose/activity in that room.", "")
field("Requirements.", "A camera for enrolment; ideally a multi-antenna Wi-Fi or radar receiver for spatial "
      "resolution; a GPU for training.")
field("Challenges and mitigations.", "Rich pose needs more antennas / access points; start with coarse activity "
      "classes and expand.")
field("Closest prior work.", "RF-Pose (MIT, CVPR 2018); DensePose From WiFi (CMU, 2023); Person-in-WiFi (ICCV 2019).")

# ---- Idea 4 ----
idea("Vision-Wireless Track Handoff Across a Sensor Network")
scorecard("Medium", "Low-Medium", "Low-Medium", "Medium", "Medium-High", "Corridors, large spaces")
field("Concept.", "Keep cameras only at a few nodes; when a tracked person walks out of camera view, associate their "
      "track with a wireless signature (Wi-Fi angle-of-arrival, or their phone's BLE) so tracking continues in the "
      "camera-free stretch and resumes at the next camera.")
field("How it works.", "Vision detects and tracks people; the wireless side estimates each person's trajectory or "
      "device signature; an assignment step (Section 3.5) matches the two, handing the identity across the gap. "
      "Vi-Fi associated pedestrians across vision and phone wireless/IMU; EyeFi associated Wi-Fi transmitters with "
      "camera detections using angle-of-arrival.")
field("Execution steps.", "")
numbered("deploy cameras at entry/exit nodes and wireless sensing along the route.", "")
numbered("detect and track people in each camera; estimate wireless trajectories or device IDs.", "")
numbered("solve the association between camera tracks and wireless tracks by minimising the matching cost.", "")
numbered("carry the identity label across camera-free segments; re-confirm at the next camera.", "")
field("Requirements.", "A few cameras; angle-of-arrival-capable Wi-Fi (multi-antenna) or BLE scanning; optionally the "
      "occupants' phones (with consent).")
field("Challenges and mitigations.", "MAC-address randomisation weakens device IDs (prefer device-free signatures or "
      "consented apps); crowded scenes raise association errors (use motion continuity as extra cost).")
field("Closest prior work.", "Vi-Fi (Rutgers, 2022); EyeFi (2020).")

# ---- Idea 5 ----
idea("mmWave / FMCW Radar Biometric Signature (high fidelity)")
scorecard("High", "Medium-High", "Medium", "High", "High", "Best accuracy, if budget allows")
field("Concept.", "Use a millimetre-wave FMCW radar instead of (or with) Wi-Fi to build a higher-fidelity body-shape "
      "and gait signature that is robust to clothing and lighting, for accurate long-term re-identification.")
field("How it works.", "The radar produces range-Doppler or point-cloud data; a network extracts a skeleton and a "
      "persistent body-shape embedding. RF-ReID (MIT) did exactly this with a multi-task recurrent architecture "
      "(skeleton prediction plus environment discrimination to reduce bias), evaluated on campus and home datasets.")
field("Execution steps.", "")
numbered("mount an FMCW radar per zone; enrol residents as in Idea 1.", "")
numbered("train the body-shape / gait embedding with metric learning and skeleton supervision.", "")
numbered("deploy radar-only re-identification and tracking; fuse with a gate camera for enrolment refresh.", "")
field("Requirements.", "An mmWave radar module (for example a TI IWR-series board); more compute; higher cost per "
      "node than Wi-Fi.")
field("Challenges and mitigations.", "Radar cost and coverage per room; combine with cheaper Wi-Fi nodes for wide "
      "areas and reserve radar for key zones.")
field("Closest prior work.", "RF-ReID (MIT, CVPR 2020); mmWave gait-recognition literature.")

# ---- Idea 6 ----
idea("Device-Tied Identity Plus Device-Free Sensing (pragmatic baseline)")
scorecard("Low", "Low", "Low", "Low-Medium", "Medium", "Quick baseline / fallback")
field("Concept.", "Attach identity to the person's phone (its BLE or Wi-Fi presence) and use device-free CSI for "
      "motion; the phone provides the name, the channel provides the movement.")
field("How it works.", "A probabilistic association links the times/places a phone is seen with the device-free motion "
      "track, assigning the phone's identity to that track. This is the identity source used in parts of Vi-Fi and "
      "EyeFi.")
field("Execution steps.", "")
numbered("with consent, register each resident's phone identifier.", "")
numbered("scan for that identifier and correlate its presence with device-free CSI motion per room.", "")
numbered("assign the phone's identity to the co-located motion track.", "")
field("Requirements.", "Minimal extra hardware; a BLE / Wi-Fi scanner; explicit consent and an app if MAC "
      "randomisation must be defeated legitimately.")
field("Challenges and mitigations.", "Only identifies phone-carriers, and modern MAC randomisation breaks passive "
      "IDs (use a consented companion app or a wearable beacon); treat as a helper, not the core.")
field("Closest prior work.", "EyeFi (2020); Vi-Fi (2022).")

# ---- Idea 7 ----
idea("Anchor-Camera Privacy-Zone Monitoring (deployable demonstrator)")
scorecard("Low-Medium", "Low", "Low", "Medium", "Medium-High", "A working end-to-end demo")
field("Concept.", "A product-leaning integration: one camera in a common area does identity and alarm verification; "
      "cheap Wi-Fi nodes cover private rooms for presence, activity and fall; identity is assigned at the camera and "
      "carried by room-to-room handoff.")
field("How it works.", "Combine the handoff of Ideas 1 or 2 with the camera-reduction architecture from the companion "
      "study: identity is stamped at the gate camera, then propagated by tracking which room the person moves to.")
field("Execution steps.", "")
numbered("deploy a gate camera plus one ESP32 CSI node per room (the ESPectre platform works out of the box for "
         "presence and motion).", "")
numbered("assign identity at the gate; update the current room by following motion events between rooms.", "")
numbered("raise identity-aware events (\u201cknown resident in bathroom, no motion 60 s -> alert\u201d).", "")
field("Requirements.", "One camera; a few ESP32 nodes; a hub (Home Assistant or a small server). Low cost, quick to "
      "stand up.")
field("Challenges and mitigations.", "Handoff can lose identity when two people cross (fall back to "
      "\u201cunknown/occupied\u201d and re-confirm at the next camera pass).")
field("Closest prior work.", "Integration of ESPectre (open source) with XModal-ID-style handoff.")

# ---- Idea 8 ----
idea("Privacy-Preserving, Consent-Based Re-ID (the responsible-design angle)")
scorecard("Medium", "Low", "Low-Medium", "High", "Medium", "The distinctive thesis framing")
field("Concept.", "Make privacy the contribution: a system that only ever recognises people who explicitly enrolled, "
      "keeps signatures on-device, is revocable, and actively resists covert tracking of non-consenting people.")
field("How it works.", "On-device (or federated) embeddings; matching restricted to the enrolled gallery with strong "
      "unknown-rejection; no raw video or CSI leaves the hub; an audit log of when identity was used. This directly "
      "answers the surveillance concerns raised by WhoFi and by the 2026 Karlsruhe result that ordinary Wi-Fi can "
      "identify people with about 99.5 percent accuracy.")
field("Execution steps.", "")
numbered("design a consented enrolment protocol with per-person revocation.", "")
numbered("store signatures locally; match only against enrolled identities; reject everything else as unknown.", "")
numbered("add an audit trail and a red-team evaluation of covert-tracking resistance.", "")
field("Requirements.", "Edge compute; careful system and protocol design; a threat model.")
field("Challenges and mitigations.", "Balancing utility and privacy; document the trade-offs and evaluate both "
      "recognition accuracy and privacy leakage.")
field("Closest prior work.", "WhoFi privacy discussion (2025); Karlsruhe Institute of Technology Wi-Fi identification "
      "(2026); biometric-data regulation (for example GDPR special-category data).")

# ===========================================================================
# 5. DECISION MATRIX
# ===========================================================================
h1("5.  Decision Matrix")
para("The eight ideas side by side. Use this to shortlist with your professor.")
make_table(
    ["#", "Idea", "Effort", "Cost", "Novelty", "Robustness", "Recommended role"],
    [
        ["1", "Vision-enrolled wireless identity", "Med-High", "Low-Med", "High", "Medium", "Flagship / headline"],
        ["2", "Wi-Fi gait Re-ID (household)", "Low-Med", "Low", "Medium", "Med-High", "De-risked core to start"],
        ["3", "Cross-modal pose/activity", "Medium", "Low-Med", "Medium", "Medium", "Complementary capability"],
        ["4", "Vision-wireless track handoff", "Medium", "Low-Med", "Medium", "Med-High", "Corridors / networks"],
        ["5", "mmWave radar biometric", "High", "Med-High", "High", "High", "If budget allows accuracy"],
        ["6", "Device-tied + device-free", "Low", "Low", "Low-Med", "Medium", "Quick baseline / fallback"],
        ["7", "Anchor-camera privacy zones", "Low-Med", "Low", "Medium", "Med-High", "Deployable demonstrator"],
        ["8", "Privacy-preserving Re-ID", "Medium", "Low", "High", "Medium", "Ethics-forward framing"],
    ],
    widths=[0.35, 2.5, 0.95, 0.7, 0.85, 1.0, 1.75], fs=8.8,
)

# ===========================================================================
# 6. RECOMMENDATION
# ===========================================================================
h1("6.  A Recommended Path to Finalise")
para("A pragmatic plan that de-risks early and still reaches an ambitious, publishable result:")
numbered("build Idea 2 (Wi-Fi gait Re-ID for the household) to get a working, measurable result quickly and to "
         "learn the data pipeline and hardware.", "Start (weeks): ")
numbered("grow it into Idea 1 (vision-enrolled, wireless-tracked identity) by adding the camera-supervised "
         "cross-modal alignment \u2014 this is the headline contribution.", "Extend: ")
numbered("add Idea 3 (pose/activity) so the camera-free rooms report not just who but what, and optionally Idea 7 "
         "to package it as a working demonstrator.", "Enrich: ")
numbered("frame the whole project with Idea 8 (consent-based, on-device, privacy-preserving), which turns the "
         "biggest risk of this technology into the project's distinctive strength.", "Frame: ")
numbered("if a millimetre-wave radar is available, use Idea 5 for a high-accuracy comparison point.", "Optional: ")
note_box("For finalising with your professor: Idea 2 is the safest single choice; Idea 1 is the most impressive; "
         "Idea 8 is the most defensible and novel framing. The three combine naturally into one coherent project: "
         "a consented, camera-enrolled, wireless-tracked identity system for a small household.")

# ===========================================================================
# 7. ETHICS, PRIVACY AND LEGAL
# ===========================================================================
h1("7.  Ethics, Privacy and Legal Considerations")
para("This class of system identifies and tracks people through walls, so it must be handled responsibly. This is "
     "not optional boilerplate \u2014 it should shape the design.")
bullet("recognise only people who explicitly opted in; reject everyone else as \u201cunknown\u201d rather than profiling "
       "them.", "Consent and closed set: ")
bullet("keep biometric signatures and raw data on the local hub; never stream them out; make enrolment revocable.", "Data locality: ")
bullet("under regimes such as the GDPR, biometric identifiers are special-category data with strict handling "
       "requirements; document lawful basis and retention.", "Legal status: ")
bullet("researchers have warned that ordinary Wi-Fi can be turned into covert person tracking (WhoFi, 2025; "
       "Karlsruhe Institute of Technology, 2026, about 99.5 percent identification). Design against misuse and "
       "evaluate the privacy leakage explicitly.", "Surveillance risk: ")
ethics_box("Recommended stance for the project: build the smallest system that serves the goal (a consenting "
           "household), keep all identity processing on-device, publish the privacy analysis alongside the accuracy "
           "numbers, and avoid any capability aimed at identifying non-consenting people.")

# ===========================================================================
# 8. REFERENCES
# ===========================================================================
h1("8.  Key References")
refs = [
    "B. Korany, C. R. Karanam, H. Cai, Y. Mostofi. XModal-ID: Using WiFi for Through-Wall Person Identification from Candidate Video Footage. ACM MobiCom, 2019. (Ranked top-1/2/3 accuracy 75 / 90 / 97 percent among 8 candidates.)",
    "L. Fan, T. Li, R. Fang, R. Hristov, Y. Yuan, D. Katabi. Learning Longterm Representations for Person Re-Identification Using Radio Signals (RF-ReID). CVPR, 2020. arXiv:2004.01091.",
    "WhoFi: Deep Person Re-Identification via Wi-Fi Channel Signal Encoding. arXiv:2507.12869, 2025. (Transformer on CSI; about 95.5 percent on NTU-Fi.)",
    "H. Cai et al. / Rutgers WINLAB. Vi-Fi: Associating Moving Subjects across Vision and Wireless Sensors, 2022.",
    "Z. Fang et al. EyeFi: Fast Human Identification through Vision and WiFi-based Trajectory Matching, 2020.",
    "Y. Zeng, P. H. Pathak, P. Mohapatra. WiWho: WiFi-Based Person Identification in Smart Spaces. IPSN, 2016.",
    "W. Wang et al. WifiU: Human Gait Recognition Using WiFi Signals, 2016.",
    "M. Zhao et al. Through-Wall Human Pose Estimation Using Radio Signals (RF-Pose). CVPR, 2018.",
    "J. Geng, D. Huang, F. De la Torre. DensePose From WiFi. arXiv:2301.00250, 2023.",
    "F. Wang et al. Person-in-WiFi: Fine-grained Person Perception using WiFi. ICCV, 2019.",
    "Y. Zheng et al. Widar3.0: Zero-Effort Cross-Domain Gesture Recognition with WiFi.",
    "Karlsruhe Institute of Technology. Person identification via WiFi beamforming feedback (2026), reported about 99.5 percent accuracy \u2014 a privacy caution.",
    "F. Pace. ESPectre \u2014 WiFi CSI motion detection for Home Assistant (open source). github.com/francescopace/espectre.",
]
for rf in refs:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(rf); r.font.size = Pt(9.5)

closing = doc.add_paragraph(); closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("This catalogue is a decision aid. Accuracy figures are quoted from the cited papers and their "
                    "public materials; scorecard ratings are relative guidance to help shortlist a direction with "
                    "your professor. Figures and rendered equations were generated for this document.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE_LT

doc.save(OUTPUT)
print("wrote", OUTPUT)
