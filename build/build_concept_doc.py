"""Build the concept Word document:
"Fewer Cameras, More Coverage - Using Wireless Perception to Reduce Camera
Count in a Multi-Room Home".

Covers: whether this has been done (literature and products), how it can be done,
and the requirements - mathematically, realistically, and practically.
No emojis are used anywhere.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "assets_wireless")
OUTPUT = os.path.join(os.path.dirname(HERE), "Wireless_Perception_Reducing_Cameras.docx")

INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
PLUM = RGBColor(0x7A, 0x64, 0x9A)
RUST = RGBColor(0xA8, 0x5A, 0x4A)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_AMBER = "FBF2E7"
SHADE_GREY = "F2F3F5"
HEADER_FILL = "DCE6EF"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(size)); e.set(qn("w:space"), "8"); e.set(qn("w:color"), color)
    pbdr.append(e); pPr.append(pbdr)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def para(text, italic=False, size=11, align=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def runs_para(parts):
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        r = p.add_run(text); r.bold = bold; r.italic = italic
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def layman_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREEN); _border(p, color="AFC8B2")
    r = p.add_run("In plain terms:  "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def sci_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_BLUE); _border(p, color="9DB6CC")
    r = p.add_run("Technical detail:  "); r.bold = True; r.font.color.rgb = SLATE
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def caution_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_AMBER); _border(p, color="D9B98A")
    r = p.add_run("Reality check:  "); r.bold = True; r.font.color.rgb = RUST
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def note_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREY); _border(p, color="C9CDD3")
    r = p.add_run(text); r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


_eqn = {"n": 0}


def equation(img, label=True):
    _eqn["n"] += 1
    from PIL import Image
    path = os.path.join(ASSETS, img)
    with Image.open(path) as im:
        w_in = im.width / 200.0; h_in = im.height / 200.0
    max_w = 5.9
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if w_in > max_w:
        run.add_picture(path, width=Inches(max_w))
    else:
        run.add_picture(path, height=Inches(max(min(h_in, 0.6), 0.26)))
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    if label:
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"( {_eqn['n']} )"); r.font.size = Pt(9); r.font.color.rgb = SLATE_LT
        cap.paragraph_format.space_after = Pt(8)
    return _eqn["n"]


_fig = {"n": 0}


def figure(img, caption, width=6.2):
    _fig["n"] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(width))
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {_fig['n']}.  "); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE
    r2 = cap.add_run(caption); r2.font.size = Pt(9.5); r2.font.color.rgb = INK
    cap.paragraph_format.space_after = Pt(12)


def variable_table(rows):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(["Symbol", "What it means"]):
        hdr[i].text = ""; p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL); tcPr.append(shd)
    for ri, (sym, mean) in enumerate(rows):
        cells = t.add_row().cells
        cells[0].text = ""; p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(sym); r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = SLATE
        cells[1].text = ""; p1 = cells[1].paragraphs[0]
        r1 = p1.add_run(mean); r1.font.size = Pt(9.5); r1.font.color.rgb = INK
        if ri % 2 == 1:
            for c in cells:
                tcPr = c._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA"); tcPr.append(shd)
    t.columns[0].width = Inches(1.3); t.columns[1].width = Inches(4.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def make_table(headers, rows, widths=None, fs=9.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""; p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL); tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; p = cells[i].paragraphs[0]
            r = p.add_run(val); r.font.size = Pt(fs); r.font.color.rgb = INK
            if i == 0:
                r.bold = True; r.font.color.rgb = SLATE
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA"); tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ===========================================================================
# TITLE
# ===========================================================================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Fewer Cameras, More Coverage"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(60); title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Using Wireless Perception to Reduce the Number of Cameras in a Multi-Room Home")
r.font.size = Pt(15); r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Can Wi-Fi sensing let one camera cover a three-room apartment that would normally need three or four? "
                 "What has been done, how to do it, and what it truly requires.")
r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = INK

line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  Feasibility Study and Design Note  \u2014"); r.font.size = Pt(11); r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(24)

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Multi-Modal Perception (MPL) programme  \u00b7  Prepared with the ESPectre Wi-Fi-sensing platform as a reference building block")
r.font.size = Pt(10); r.font.color.rgb = SLATE_LT
doc.add_page_break()

# ===========================================================================
# CONTENTS
# ===========================================================================
h1("Contents")
toc_p = doc.add_paragraph(); run = toc_p.add_run()
fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
t_run = OxmlElement("w:t"); t_run.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
for e in (fldChar, instr, fldChar2, t_run, fldChar3): run._r.append(e)
doc.add_page_break()

# ===========================================================================
# 1. THE QUESTION AND THE SHORT ANSWER
# ===========================================================================
h1("1.  The Question and the Short Answer")
para("A three-room apartment is hard to cover with cameras. A camera only sees a cone of space in front of it, "
     "and walls stop that cone at the edge of each room. To watch the living room, the bedroom, and the bathroom, "
     "you typically need one camera per room \u2014 three or four in total once hallways and blind spots are counted. "
     "Cameras are also unwelcome in bedrooms and bathrooms for obvious privacy reasons.")
runs_para([("The question of this study:  ", True, False),
           ("can we keep just one camera and use wireless perception (Wi-Fi sensing) to cover the other rooms, "
            "and still get results that are, for practical purposes, similar to a three-or-four-camera setup?", False, False)])
para("The honest short answer has two parts:")
bullet("if \u201csimilar results\u201d means seeing faces, reading text, or capturing fine visual detail in every "
       "room, then no \u2014 wireless perception cannot replace a camera for that, at any price.", "For appearance and identity: ")
bullet("if \u201csimilar results\u201d means knowing whether a room is occupied, how many people are in it, roughly where "
       "they are, what activity is happening, and whether someone has fallen, then yes \u2014 one camera in the room "
       "that needs visual detail, plus one cheap Wi-Fi sensor per remaining room, can match the task-level outcome "
       "of three or four cameras.", "For presence, motion, count, activity and safety: ")
layman_box("Cameras and Wi-Fi answer different questions. A camera answers \u201cwho is this and what exactly are they "
           "doing?\u201d Wi-Fi answers \u201cis someone here, roughly where, and are they moving or in trouble?\u201d For most "
           "home-monitoring goals you only need the second answer in most rooms, and Wi-Fi already passes through "
           "walls to give it. So you keep one camera where detail matters and let Wi-Fi handle the rest.")
figure("fig_apartment.png",
       "The core idea. Left: vision alone needs a camera in every room because walls block each field of view, and "
       "cameras are intrusive in private rooms. Right: one camera is kept where visual detail matters, while a cheap "
       "Wi-Fi link (router to ESP32) senses each of the other rooms straight through the walls.", width=6.6)

# ===========================================================================
# 2. HAS THIS BEEN DONE?
# ===========================================================================
h1("2.  Has This Been Done Before?")
para("The specific product framing \u2014 \u201creplace three home cameras with one camera plus Wi-Fi\u201d \u2014 is not a "
     "single published paper or off-the-shelf product. However, every building block needed to do it is well "
     "established in peer-reviewed research and in shipping commercial systems. In other words, the exact packaging "
     "is novel, but nothing in it is unproven. The relevant prior work falls into five groups.")

h2("2.1  Wi-Fi sensing exists and is being standardised")
para("Using ordinary Wi-Fi signals to sense the environment is a recognised field. The IEEE approved the 802.11bf "
     "\u201cWLAN Sensing\u201d task group in 2020 to standardise exactly this, exposing low-level channel measurements "
     "(Channel State Information, or CSI) for sensing. The classic enabling tool was the Linux 802.11n CSI Tool "
     "(Halperin et al., 2011), which first made CSI available from commodity network cards. Wi-Fi sensing is "
     "categorised into detection (presence, intrusion, falls), localisation (where motion occurs), recognition "
     "(gesture, gait, activity), and estimation (people counting, breathing and heart rate).")

h2("2.2  Through-wall sensing with radio")
para("A defining advantage of radio over light is that it passes through walls. The MIT group demonstrated a series "
     "of systems \u2014 RF-Capture, RF-Pose, and RF-Pose3D \u2014 that recover a coarse human skeleton through walls using "
     "radio reflections, trained with a camera as the teacher. This is direct evidence that radio can report human "
     "posture in rooms a camera cannot see.")

h2("2.3  Human pose and dense body from Wi-Fi alone")
para("Several works reconstruct human pose from commodity Wi-Fi. Person-in-WiFi (ICCV 2019) and WiSPPN estimate body "
     "keypoints from CSI; DensePose From WiFi (Carnegie Mellon, 2023) maps Wi-Fi amplitude and phase to a dense body "
     "surface across 24 regions, reporting performance comparable to early image-based methods using Wi-Fi as the "
     "only input. The authors explicitly frame this as a low-cost, privacy-preserving alternative to cameras.")

h2("2.4  Localisation, activity, and cross-domain robustness")
para("The Widar line of work (Widar, Widar2.0, Widar3.0) performs device-free localisation, tracking, and "
     "cross-environment gesture recognition from Wi-Fi. A large body of work covers Wi-Fi human activity recognition, "
     "fall detection, and vital-sign monitoring (breathing and heart rate). These are the exact capabilities we would "
     "ask Wi-Fi to provide in the camera-free rooms.")

h2("2.5  Camera-supervised (cross-modal) learning and commercial products")
para("A recurring pattern makes deployment realistic: use a camera to automatically label Wi-Fi data during a short "
     "period when both see the same scene, then let the Wi-Fi model run alone afterwards. RF-Pose and DensePose From "
     "WiFi both use vision as the teacher in this way. On the commercial side, Wi-Fi motion sensing already ships: "
     "Cognitive Systems (WiFi Motion, now embedded in many routers), Origin Wireless (health and presence), and the "
     "open-source ESPectre project used as the reference platform for this study. A 2026 Karlsruhe Institute of "
     "Technology result even identified individuals from ordinary Wi-Fi with very high accuracy, underlining both the "
     "capability and the privacy stakes.")
note_box("Bottom line: something very close to this has been done in the research literature (camera-taught Wi-Fi that "
         "covers what the camera cannot see), and the Wi-Fi half is already a commercial reality. What is not yet a "
         "packaged product is the explicit \u201cone-camera-plus-Wi-Fi replaces N cameras\u201d home system \u2014 which is "
         "precisely the gap this study targets.")
make_table(
    ["Prior work / product", "What it demonstrates", "Relevance here"],
    [
        ["IEEE 802.11bf; Halperin CSI Tool", "Standardised Wi-Fi sensing; CSI from commodity hardware", "The measurement substrate we build on"],
        ["RF-Pose / RF-Pose3D / RF-Capture (MIT)", "Through-wall human pose from radio, camera-supervised", "Radio reports posture where cameras cannot see"],
        ["Person-in-WiFi; DensePose From WiFi (CMU)", "Body keypoints / dense surface from Wi-Fi only", "Rich human sensing without a camera"],
        ["Widar / Widar3.0", "Localisation, tracking, cross-domain gestures", "Coarse location and activity per room"],
        ["Cognitive Systems; Origin Wireless", "Shipping Wi-Fi motion and presence products", "Commercial proof the Wi-Fi half works"],
        ["ESPectre (open source)", "ESP32 CSI motion detection into Home Assistant", "A ready, low-cost per-room sensor node"],
    ],
    widths=[2.1, 2.6, 1.9],
)

# ===========================================================================
# 3. WHAT "SIMILAR RESULTS" REALLY MEANS
# ===========================================================================
h1("3.  What \u201cSimilar Results\u201d Really Means")
para("The whole feasibility question hinges on defining the goal precisely. Cameras and Wi-Fi are not "
     "interchangeable; they are complementary, with almost opposite strengths and weaknesses. Deciding how many "
     "cameras can be removed is really deciding which questions each room must answer.")
figure("fig_modalities.png",
       "Capability comparison. A camera in view can do almost everything (bar reliable breathing and privacy in "
       "sensitive rooms). Wi-Fi sensing is strong for presence, counting, coarse location, activity and falls, weak "
       "for pose without richer hardware, and essentially cannot do identity or appearance.", width=6.4)
para("The design rule that follows is simple:")
runs_para([("Keep a camera only where a question truly needs pixels ", True, False),
           ("(a face at the front door, reading a label, verifying an alarm). ", False, False),
           ("Use Wi-Fi everywhere the question is about presence, count, location, activity, or safety.", True, False)])
make_table(
    ["Monitoring goal", "Needs a camera?", "Wi-Fi sufficient?"],
    [
        ["Is anyone in the room?", "No", "Yes (mature)"],
        ["How many people?", "Helpful", "Yes, approximately"],
        ["Roughly where in the room?", "Helpful", "Yes, coarse"],
        ["What activity (walk, sit, fall)?", "Helpful", "Yes, with a trained model"],
        ["Has someone fallen / not moved?", "Helpful", "Yes (active research, some products)"],
        ["Breathing / heart rate at rest", "No (hard for camera)", "Yes, near the link"],
        ["Who is this person (identity)?", "Yes", "No (except constrained lab settings)"],
        ["Read text / see fine detail / colour", "Yes", "No"],
    ],
    widths=[2.7, 1.9, 2.0],
)
caution_box("\u201cSimilar results\u201d is achievable at the level of events and decisions (occupied, two people, someone "
            "fell in the bedroom) but not at the level of imagery. If the requirement is a recorded video of every "
            "room, Wi-Fi cannot deliver it and no reduction in cameras is possible.")

# ===========================================================================
# 4. HOW WE CAN DO IT
# ===========================================================================
h1("4.  How We Can Do It")
para("Two approaches are worth separating: a simple one that is deployable today, and an advanced one that recovers "
     "richer information. They can be combined \u2014 start simple, add the advanced layer later.")

h2("4.1  Approach A \u2014 complementary coverage (deployable now)")
para("Give every room a sensing modality and merge their outputs. The camera covers the one room where visual detail "
     "is wanted. Each other room gets one Wi-Fi link: the home router as transmitter and a small ESP32 sensor as "
     "receiver, positioned so the link crosses the room. Each Wi-Fi node reports presence, motion, and (with a model) "
     "activity. A hub (for example a Raspberry Pi running Home Assistant) collects everything and raises events.")
figure("fig_fusion_arch.png",
       "Reference architecture. One camera in the key room and one ESP32 Wi-Fi node per remaining room feed a hub. "
       "Each stream is turned into features and per-room decisions, then merged by late fusion into events such as "
       "presence, count, activity and fall.", width=6.5)

h2("4.2  Approach B \u2014 cross-modal learning (camera teaches Wi-Fi)")
para("To get more than presence from the Wi-Fi rooms \u2014 for example coarse pose or reliable activity classes \u2014 use "
     "the camera as a teacher. During an enrolment period, place the camera so it temporarily overlaps a Wi-Fi room "
     "(or move it room to room). While both observe the same scene, the camera automatically generates labels "
     "(person present, count, pose, activity) that are paired with the simultaneous Wi-Fi CSI. A model is trained to "
     "predict the camera-derived labels from CSI alone. Afterwards the camera is returned to its permanent room and "
     "the Wi-Fi model runs unaided in the others.")
layman_box("Think of the camera as a tutor that sits next to the Wi-Fi sensor for a while, pointing and saying "
           "\u201cthat pattern means one person walking, that one means someone sat down.\u201d Once the Wi-Fi sensor has "
           "learned the association, the tutor can leave and the sensor keeps making those calls on its own.")
sci_box("This is knowledge distillation across modalities (a teacher-student scheme). The camera provides "
        "pseudo-labels y_cam; a network f_theta maps CSI to those labels; training minimises a supervised loss over "
        "the overlap window. It is exactly the mechanism behind RF-Pose and DensePose From WiFi, and it removes the "
        "need for hand-labelling, which is the main practical barrier to Wi-Fi sensing.")

# ===========================================================================
# 5. THE MATHEMATICS
# ===========================================================================
h1("5.  The Mathematics, With Every Variable Explained")
para("This section gives the equations that justify the approach. For each, we state what it says, why it matters, "
     "and what each symbol means.")

h2("5.1  What the receiver measures: the wireless channel")
para("A receiver does not see an image; it measures how the environment transformed the transmitted signal. In "
     "vector form for one snapshot:")
equation("eq_channel.png")
runs_para([("What it says.  ", True, False),
           ("The received signal y is the channel matrix H acting on the transmitted signal x, plus noise n. All the "
            "information about the room is inside H.", False, False)])
variable_table([
    ("y", "the received signal (what the ESP32 records)."),
    ("H", "the channel: how the room scaled, delayed and echoed the signal. This is the sensing quantity."),
    ("x", "the known transmitted signal (Wi-Fi packets from the router)."),
    ("n", "receiver noise."),
])

h2("5.2  The channel as a sum over propagation paths (CSI)")
para("Wi-Fi measures the channel separately on each OFDM subcarrier (frequency bin). That per-frequency channel is "
     "the Channel State Information (CSI), and it is a sum over all the paths the signal travelled:")
equation("eq_csi.png")
runs_para([("What it says.  ", True, False),
           ("The channel at frequency f_k is the sum of P paths; each path p contributes an amplitude a_p and a phase "
            "that depends on its delay tau_p. Reflections off walls, furniture and people are these paths.", False, False)])
runs_para([("Why it matters.  ", True, False),
           ("A person is a moving reflector. When they move, their path's delay and amplitude change, so the CSI "
            "changes. Reading those changes is the whole basis of Wi-Fi sensing.", False, False)])
variable_table([
    ("H(f_k, t)", "the CSI on subcarrier k at time t."),
    ("P", "the number of propagation paths (direct plus reflections)."),
    ("a_p(t)", "the amplitude (strength) of path p, possibly changing with time."),
    ("f_k", "the frequency of subcarrier k."),
    ("tau_p(t)", "the time delay of path p (its length divided by the speed of light)."),
])
figure("fig_csi_tensor.png",
       "CSI is the wireless equivalent of a pixel grid: a value of amplitude and phase for every antenna, every "
       "subcarrier, and every instant. Motion writes patterns into this grid that a model can learn to read.", width=5.8)

h2("5.3  Why motion is visible: static plus dynamic paths")
para("Splitting the channel into a fixed part and a moving part makes the sensing signal explicit:")
equation("eq_dyn.png")
runs_para([("What it says.  ", True, False),
           ("The channel is a static component from the unchanging room (walls, furniture) plus a sum of dynamic "
            "components from moving reflectors m. Each moving reflector adds a term whose phase changes as its "
            "distance d_m(t) changes.", False, False)])
runs_para([("How it works.  ", True, False),
           ("A person walking changes d_m(t) continuously, which rotates the phase and produces a Doppler shift. "
            "Subtracting or filtering out the static part leaves exactly the human-motion signal.", False, False)])
variable_table([
    ("H_static", "the unchanging channel from the room's fixed geometry."),
    ("alpha_m(t)", "the complex reflection coefficient of moving object m."),
    ("d_m(t)", "the changing path length via moving object m."),
    ("lambda", "the wavelength (about 12.5 cm at 2.4 GHz)."),
])

h2("5.4  The simplest detector: turbulence and its variance")
para("A device as small as an ESP32 does not image anything; it computes a single number per packet that measures "
     "how disturbed the channel is, then watches how much that number fluctuates. This is exactly what the ESPectre "
     "platform does.")
equation("eq_turbulence.png")
runs_para([("What it says.  ", True, False),
           ("The turbulence tau_t is the spread (standard deviation) of the amplitudes across the selected "
            "subcarriers at time t. The moving variance V_t is how much that turbulence varies over a window of W "
            "recent packets. A still room gives low variance; a moving person gives high variance.", False, False)])
variable_table([
    ("tau_t", "spatial turbulence: standard deviation of subcarrier amplitudes in one packet."),
    ("|H_1|..|H_K|", "the amplitudes of the K selected subcarriers."),
    ("V_t", "moving variance of turbulence over the recent window (the detection score)."),
    ("W", "the window length in packets."),
])

h2("5.5  What sets the resolution: bandwidth, aperture, and Doppler")
para("Three physical limits explain why cheap Wi-Fi gives coarse information and why richer output needs more "
     "hardware.")
equation("eq_range.png")
runs_para([("Range and angle resolution.  ", True, False),
           ("How finely you can separate reflectors in distance is set by bandwidth B; how finely in angle is set by "
            "the antenna aperture D. A single small antenna and 20-40 MHz of Wi-Fi bandwidth give metre-scale range "
            "and very coarse angle \u2014 enough for presence and rough location, not for imaging.", False, False)])
variable_table([
    ("Delta R", "range resolution (smallest separable distance difference)."),
    ("c", "speed of light."),
    ("B", "signal bandwidth (tens of MHz for Wi-Fi)."),
    ("Delta theta", "angular resolution."),
    ("D", "antenna array aperture (size); larger means finer angle."),
])
equation("eq_doppler.png")
runs_para([("Motion and sampling.  ", True, False),
           ("A reflector moving at speed v produces a Doppler shift f_D. To capture activity or breathing without "
            "aliasing, the packet sampling rate f_s must be at least twice the highest Doppler present (the Nyquist "
            "rule). This is why ESPectre streams around 100 packets per second.", False, False)])
variable_table([
    ("f_D", "Doppler frequency shift caused by motion."),
    ("v", "speed of the moving body part."),
    ("beta", "angle between the motion and the signal path."),
    ("f_s", "packet sampling rate (packets per second)."),
    ("f_D,max", "the largest Doppler shift to be captured."),
])
equation("eq_dof.png")
runs_para([("How much a link can resolve.  ", True, False),
           ("The number of independent things a link can distinguish grows with the transmit and receive apertures "
            "and shrinks with wavelength and distance. Small commodity radios have a low degree-of-freedom count, "
            "which is the fundamental reason a single Wi-Fi link reports coarse, not detailed, scenes.", False, False)])
variable_table([
    ("N_dof", "number of independent spatial degrees of freedom (resolvable elements)."),
    ("A_T, A_R", "transmit and receive antenna aperture areas."),
    ("lambda", "wavelength."),
    ("d", "link distance."),
])

h2("5.6  Combining camera and Wi-Fi: fusion")
para("When two sensors observe overlapping evidence, the principled way to combine them is a posterior over the "
     "scene state:")
equation("eq_fusion.png")
runs_para([("What it says.  ", True, False),
           ("The probability of a scene state s given both the camera evidence and the radio evidence is proportional "
            "to each sensor's likelihood times a prior. Whichever sensor is more confident for a given question "
            "dominates the answer automatically.", False, False)])
variable_table([
    ("s", "the scene state we want (occupied, count, activity, location)."),
    ("z_cam", "the camera's evidence."),
    ("z_rf", "the radio (Wi-Fi) evidence."),
    ("p(z|s)", "the likelihood of each sensor's evidence given the state."),
    ("p(s)", "the prior probability of the state."),
])

h2("5.7  Training the Wi-Fi model from the camera")
para("The cross-modal learning of Approach B is one optimisation:")
equation("eq_crossmodal.png")
runs_para([("What it says.  ", True, False),
           ("Choose the model parameters theta that minimise, over the overlap period O, the loss between the model's "
            "prediction from CSI and the camera-derived label. After training, the model needs only CSI.", False, False)])
variable_table([
    ("theta*", "the trained model parameters."),
    ("f_theta(CSI_t)", "the Wi-Fi model's prediction at time t."),
    ("y_cam_t", "the label the camera produced automatically at time t."),
    ("O", "the enrolment window when camera and Wi-Fi overlap."),
    ("L", "the loss measuring prediction error."),
])

h2("5.8  The camera-count argument")
para("Finally, the counting logic behind the headline claim:")
equation("eq_cameras.png")
runs_para([("What it says.  ", True, False),
           ("Vision alone needs, per room r, enough cameras to tile that room's area with camera fields of view \u2014 "
            "summed over all rooms, this is several cameras. Because radio passes through walls, the same coverage "
            "for event-level tasks becomes one camera (kept where detail is needed) plus one Wi-Fi link per room.", False, False)])
variable_table([
    ("C_min (vision)", "minimum cameras for vision-only coverage."),
    ("R", "number of rooms."),
    ("A_r", "area of room r."),
    ("A_FoV", "area one camera can usefully cover."),
    ("1 + R", "one camera plus one Wi-Fi link per room (the proposed system)."),
])

# ===========================================================================
# 6. REQUIREMENTS
# ===========================================================================
h1("6.  Requirements: Mathematical, Realistic, Practical")

h2("6.1  Mathematical requirements (what must be true to work)")
bullet("only what the physics allows can be recovered. Presence, motion, coarse location, activity and "
       "breathing are identifiable from CSI; identity and appearance are not, from a single commodity link.", "Identifiability: ")
bullet("anything beyond presence needs diversity \u2014 multiple subcarriers (always available), and ideally "
       "multiple antennas or multiple links, to raise the degree-of-freedom count for location and pose.", "Spatial diversity: ")
bullet("the packet rate must satisfy Nyquist for the fastest motion of interest (about 100 packets per "
       "second is ample for walking and breathing).", "Adequate sampling: ")
bullet("the learning approaches need paired examples; the camera supplies them automatically during "
       "enrolment, but the coverage of that enrolment set bounds what the model can later recognise.", "Labelled data: ")

h2("6.2  Realistic requirements (what to expect in the wild)")
caution_box("Wi-Fi sensing models are environment-specific. A model trained in one apartment often degrades in "
            "another because the multipath is different (this is called domain shift). Expect to calibrate or "
            "re-train per home, and to re-calibrate after major furniture changes.")
bullet("presence and motion are robust and mature; activity and pose are noisier and improve with more "
       "links and more training data.", "Graceful capability ladder: ")
bullet("distinguishing several people in one room from a single link is hard; more links or a brief camera "
       "assist help.", "Multi-person is harder: ")
bullet("coarse, non-identifying sensing is a feature, not a bug, for bedrooms and bathrooms where cameras "
       "are unacceptable.", "Privacy is an advantage: ")

h2("6.3  Practical requirements (what to buy and do)")
make_table(
    ["Item", "Role", "Approximate cost"],
    [
        ["1 x camera (existing IP or USB)", "Detail and identity in the key room; teacher for training", "already owned / low"],
        ["Home Wi-Fi router (2.4 GHz)", "Transmitter for all Wi-Fi links", "already owned"],
        ["3 x ESP32-C6 or ESP32-S3 (CSI-capable)", "One receiver per remaining room", "about 10 euro each"],
        ["Raspberry Pi (or existing PC / NAS)", "Hub: Home Assistant + ESPHome + fusion", "low / already owned"],
        ["USB cables, small antennas", "Flashing and better reception", "minimal"],
    ],
    widths=[2.5, 2.6, 1.4],
)
para("Practical placement and process follow directly from the physics:")
numbered("place each ESP32 so the router-to-ESP32 line crosses the room to be monitored, ideally 3 to 8 metres from "
         "the router, at table height, away from large metal objects.", "Geometry: ")
numbered("keep the room still for the brief auto-calibration at power-on (for the ESPectre motion mode).", "Calibrate: ")
numbered("for richer output, run the short camera-supervised enrolment described in Section 4.2, then remove the "
         "camera from that room.", "Enrol (optional): ")
numbered("collect the per-room decisions in Home Assistant and define the fusion rules and alerts (occupied, count, "
         "fall, no-motion timeout).", "Fuse: ")

# ===========================================================================
# 7. A CONCRETE PLAN FOR THE THREE-ROOM APARTMENT
# ===========================================================================
h1("7.  A Concrete Plan for the Three-Room Apartment")
para("The following is a end-to-end recipe that reduces cameras from three or four to one while preserving "
     "event-level coverage in every room.")
figure("fig_geometry.png",
       "Each Wi-Fi link senses one room. Placing the router-to-ESP32 line across the space means anyone crossing the "
       "link's Fresnel zones perturbs the signal, which the node reports as motion or (with a model) activity.", width=5.9)
make_table(
    ["Room", "Sensor kept", "What it delivers"],
    [
        ["Living room (entry)", "The one camera + an ESP32", "Identity/verification at the door; presence and activity"],
        ["Bedroom", "ESP32 Wi-Fi link only", "Presence, in/out of bed motion, breathing near the link, fall alert"],
        ["Bathroom", "ESP32 Wi-Fi link only", "Presence, no-motion / fall alert (no camera, by design)"],
    ],
    widths=[1.9, 2.3, 2.4],
)
para("Deployment steps:")
numbered("flash three CSI-capable ESP32 boards with the ESPectre firmware and add them to Home Assistant (see the "
         "companion ESPectre guide). Position one per room per the geometry rule.", "Stand up the Wi-Fi layer: ")
numbered("verify each node reports clean presence and motion, tuning the detection threshold per room.", "Baseline: ")
numbered("run a short camera-supervised enrolment to add activity and fall classes to the Wi-Fi rooms, if desired.", "Add intelligence: ")
numbered("keep the single camera in the living room for identity and for verifying alarms raised by Wi-Fi elsewhere.", "Keep one camera: ")
numbered("write fusion automations: for example, \u201cWi-Fi says fall in bathroom AND no motion for 60 s -> alert\u201d, "
         "or \u201ccamera confirms unknown person at door -> notify\u201d.", "Define events: ")
para("Evaluation metrics to confirm \u201csimilar results\u201d:")
bullet("presence and motion: recall and false-positive rate per room (target recall above 95 percent, false "
       "positives below 5 percent, which ESPectre already reports for motion).", None)
bullet("localisation: which-room accuracy (near 100 percent with one link per room) and coarse in-room position error.", None)
bullet("safety: fall-detection rate and time-to-alert; no-motion timeout reliability.", None)

# ===========================================================================
# 8. LIMITATIONS AND HONEST EXPECTATIONS
# ===========================================================================
h1("8.  Limitations and Honest Expectations")
bullet("Wi-Fi will not give faces, appearance, colour, or readable detail. Keep the camera "
       "for those.", "Not a camera replacement for imagery. ")
bullet("models may need per-home calibration and can drift after big layout changes.", "Environment dependence. ")
bullet("separating several people in one room from a single link is unreliable; add links or a brief "
       "camera assist.", "Multi-person scenes. ")
bullet("coarse pose and dense body from Wi-Fi (the DensePose-style results) need multiple antennas or "
       "access points and non-trivial training; plan for it as a phase two, not day one.", "Rich pose needs more. ")
bullet("even coarse sensing reveals occupancy patterns; treat the data as sensitive, keep it on the local "
       "hub, and be transparent with occupants.", "Privacy is real. ")

# ===========================================================================
# 9. RELATION TO THE OTHER WORK IN THIS PROGRAMME
# ===========================================================================
h1("9.  How This Connects to the Rest of the Programme")
para("Two nearby references make this plan concrete and forward-looking. ESPectre (documented in the companion "
     "guide) is a ready, roughly ten-euro, open-source Wi-Fi-sensing node \u2014 it is the exact hardware and software "
     "for the per-room Wi-Fi layer in Section 7, providing presence and motion out of the box and an experimental "
     "on-device ML detector. The MetaAI paper studied earlier in this programme points to the frontier: pushing "
     "computation into the wireless channel itself with programmable metasurfaces, which is a way to make Wi-Fi "
     "sensing richer and cheaper in the longer term.")
note_box("Recommended path: build Approach A with ESPectre nodes now to get reliable multi-room presence, motion and "
         "occupancy with a single camera; then layer Approach B (camera-supervised activity and fall detection) once "
         "the basic system is trusted. This realistically takes the apartment from three or four cameras to one, with "
         "similar results at the level of the events that actually matter.")

# ===========================================================================
# 10. REFERENCES
# ===========================================================================
h1("10.  Key References and Further Reading")
refs = [
    "IEEE 802.11bf Task Group \u2014 WLAN Sensing standardisation (2020- ).",
    "D. Halperin, W. Hu, A. Sheth, D. Wetherall. Tool release: gathering 802.11n traces with channel state information. ACM SIGCOMM CCR, 2011.",
    "F. Adib, D. Katabi. See Through Walls with WiFi. ACM SIGCOMM, 2013.",
    "M. Zhao et al. Through-Wall Human Pose Estimation Using Radio Signals (RF-Pose). CVPR, 2018.",
    "M. Zhao et al. RF-Pose3D / RF-based 3D skeletons. MIT CSAIL.",
    "F. Wang et al. Person-in-WiFi: Fine-grained Person Perception using WiFi. ICCV, 2019.",
    "J. Geng, D. Huang, F. De la Torre. DensePose From WiFi. arXiv:2301.00250, 2023.",
    "Y. Zheng et al. Widar3.0: Zero-Effort Cross-Domain Gesture Recognition with Wi-Fi. IEEE TPAMI / MobiSys.",
    "Wireless Broadband Alliance. Wi-Fi Sensing white papers (2019- ).",
    "Karlsruhe Institute of Technology. Person identification from commodity Wi-Fi (2026) \u2014 reported near-perfect accuracy; a privacy caution.",
    "C. Feng et al. Enabling Over-the-Air AI for Edge Computing via Metasurface-Driven Physical Neural Networks (MetaAI). ACM SIGCOMM, 2025.",
    "F. Pace. ESPectre \u2014 Wi-Fi CSI motion detection for Home Assistant (open source). github.com/francescopace/espectre.",
]
for rf in refs:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(rf); r.font.size = Pt(9.5)

closing = doc.add_paragraph(); closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("This note is a feasibility and design study. Figures and rendered equations were generated for "
                    "this document; performance figures cited for Wi-Fi motion detection are drawn from the ESPectre "
                    "project's published metrics, and capability claims are grounded in the referenced literature.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE_LT

doc.save(OUTPUT)
print("wrote", OUTPUT)
