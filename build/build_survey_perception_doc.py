"""Build Document 1 of the literature survey:
"Wireless Perception: A Beginner's Survey of Wi-Fi and RF Sensing Papers".

One page per paper, plain-language + technical, honest limitations, real
citations only. Companion to Document 2 (the MetaAI / computing-in-the-channel
survey). No emojis.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
OUTPUT = os.path.join(os.path.dirname(HERE), "Survey_1_Wireless_Perception_Papers.docx")

INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
PLUM = RGBColor(0x7A, 0x64, 0x9A)
RUST = RGBColor(0xA8, 0x5A, 0x4A)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_AMBER = "FBF2E7"
SHADE_GREY = "F2F3F5"
HEADER_FILL = "DCE6EF"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = INK
pf = normal.paragraph_format; pf.space_after = Pt(8); pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(size)); e.set(qn("w:space"), "8"); e.set(qn("w:color"), color)
    pbdr.append(e); pPr.append(pbdr)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)


def para(text, italic=False, size=11, space_after=8):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def field(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + "  "); r.bold = True; r.font.color.rgb = SLATE
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def layman_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREEN); _border(p, color="AFC8B2")
    r = p.add_run("In plain terms:  "); r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def note_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_GREY); _border(p, color="C9CDD3")
    r = p.add_run(text); r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


def caution_box(text):
    p = doc.add_paragraph(); _shade(p, SHADE_AMBER); _border(p, color="D9B98A")
    r = p.add_run("Honest caveat:  "); r.bold = True; r.font.color.rgb = RUST
    p.add_run(text)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    return p


def make_table(headers, rows, widths=None, fs=9.5, first_bold=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""; p = hdr[i].paragraphs[0]
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), HEADER_FILL); tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; p = cells[i].paragraphs[0]
            r = p.add_run(val); r.font.size = Pt(fs); r.font.color.rgb = INK
            if i == 0 and first_bold:
                r.bold = True; r.font.color.rgb = SLATE
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F6F8FA"); tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


_paper = {"n": 0}


def paper(title):
    _paper["n"] += 1
    doc.add_heading(f"Paper {_paper['n']}.  {title}", level=2)


# ===========================================================================
# TITLE PAGE
# ===========================================================================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Wireless Perception"); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(60); title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Beginner's Survey of Wi-Fi and RF Sensing Papers"); r.font.size = Pt(16); r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("What has been done in seeing people and their actions through radio signals \u2014 the "
                 "landmark works, the datasets, the methods, and how well they generalise. One page per "
                 "paper, explained from scratch, with real citations only.")
r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = INK

line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  Document 1 of 2  \u00b7  Companion to \u201cComputing in the Wireless Channel\u201d  \u2014")
r.font.size = Pt(11); r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(24)

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Multi-Modal Perception (MPL) programme  \u00b7  Prepared as a reading guide")
r.font.size = Pt(10); r.font.color.rgb = SLATE_LT
doc.add_page_break()

# ===========================================================================
# HOW TO READ
# ===========================================================================
h1("How to Read This Document")
para("This is the first of two survey documents. It covers wireless perception: using radio signals "
     "(mostly Wi-Fi, some radar) to sense people \u2014 where they are, how they move, their pose, their "
     "activity, and even their identity. The second document covers the separate idea of computing "
     "inside the wireless channel (the MetaAI paper and its relatives).")
para("Every paper below is laid out the same way so you can compare them and read at your own depth:")
bullet("the one-line summary of what the paper is about.", "What it is: ")
bullet("a from-scratch explanation with no jargon assumed.", "In plain terms: ")
bullet("the method, the data used, and the headline result.", "How it works / Data / Result: ")
bullet("why the paper mattered and what it unlocked.", "Why it matters: ")
bullet("the honest limitations \u2014 what it cannot do.", "Honest caveat: ")
bullet("the exact reference so you can find it.", "Citation: ")
note_box("A note on honesty. Where a number is well established I state it. Where I am not certain of an "
         "exact figure, I describe the result qualitatively rather than invent a number. Nothing here is "
         "fabricated; every paper is real and the citation is given so you can verify it yourself.")

# ===========================================================================
# BACKGROUND
# ===========================================================================
h1("1.  The Idea Behind Wireless Perception (Read This First)")
para("Before the papers, three background facts make all of them easy to understand.")
bullet("radio waves (Wi-Fi, radar) travel through a room, bounce off walls, furniture and human "
       "bodies, and arrive at a receiver slightly changed. A moving body changes those reflections in a "
       "characteristic way.", "Radio reflects off people. ")
bullet("Wi-Fi chips can report Channel State Information (CSI): a fine-grained measurement of how the "
       "signal's amplitude and phase changed across many frequencies. CSI is the raw material of Wi-Fi "
       "sensing. (Older work used coarse signal strength, RSSI, which carries far less detail.)",
       "Wi-Fi already measures the channel. ")
bullet("if a moving body changes the radio in a repeatable way, a neural network can be trained to read "
       "those changes and output something meaningful \u2014 presence, a gesture, a pose, an identity.",
       "Machine learning reads the pattern. ")
layman_box("Think of the room as a still pond and your Wi-Fi router as someone tapping the water. When a "
           "person moves, they make ripples. The receiver watches the ripples arrive and a neural network "
           "learns to say \u201cthat ripple pattern means a person waving\u201d or \u201cthat one means "
           "someone fell.\u201d No camera needed, and it works in the dark and through walls.")
para("The papers below are grouped from easier tasks (presence, gesture) to harder ones (pose, identity). "
     "A recurring trick appears again and again: use a camera during training only, to automatically "
     "label what the radio is seeing, then throw the camera away and let the radio work alone. This is "
     "called cross-modal supervision, and it is the bridge to your MPL \u201creduce the cameras\u201d idea.")

# ===========================================================================
# 2. THE PAPERS
# ===========================================================================
h1("2.  The Papers, One Per Page")

# ---- Paper 1: Widar3.0 ----
paper("Widar3.0 \u2014 Cross-Domain Gesture Recognition with Wi-Fi (the generalisation landmark)")
field("What it is.", "A system that recognises hand gestures from commodity Wi-Fi, and \u2014 crucially \u2014 "
      "keeps working in a new room or a new orientation it never saw during training. This \u201czero-effort "
      "cross-domain\u201d property is the reason the paper is famous.")
layman_box("Most Wi-Fi sensing models are like a student who memorised one exam room and gets lost in "
           "any other room. Widar3.0's trick is to convert the raw signal into a quantity that describes "
           "the motion itself \u2014 how fast each part of the body moves \u2014 rather than how the signal "
           "looked in that one room. Motion is the same everywhere, so the model transfers.")
field("How it works.", "It computes a Body-coordinate Velocity Profile (BVP): a domain-independent "
      "representation of how the body's parts move, derived from the Doppler shifts measured across "
      "several Wi-Fi links. A neural network then classifies the BVP into a gesture. Because the BVP is "
      "tied to the body, not the environment, the model generalises across locations and orientations.")
field("Data and result.", "The authors released the Widar3.0 dataset (a widely reused public benchmark: "
      "multiple people, gestures, rooms and orientations). It reports strong cross-domain gesture "
      "accuracy where earlier methods collapsed when the environment changed.")
field("Why it matters.", "It named and tackled the central problem of Wi-Fi sensing \u2014 generalisation. "
      "Any project you propose will be judged on whether it works outside the exact room it was trained "
      "in, and BVP-style domain-independent features are one of the main answers.")
caution_box("Gestures are a relatively easy task with clear motion. The domain-independence is strong for "
            "gestures but does not automatically extend to fine tasks like identity. It still needs "
            "several synchronised Wi-Fi links, not a single link.")
field("Citation.", "Y. Zheng, Y. Zhang, K. Qian, G. Zhang, Y. Liu, C. Wu, Z. Yang. \u201cZero-Effort "
      "Cross-Domain Gesture Recognition with Wi-Fi.\u201d ACM MobiSys 2019.")

# ---- Paper 2: RF-Pose ----
paper("RF-Pose \u2014 Through-Wall Human Pose Estimation Using Radio Signals (MIT)")
field("What it is.", "A system that draws a 2D stick-figure skeleton of a person \u2014 even when the person "
      "is behind a wall \u2014 using only radio signals, no camera at inference time.")
layman_box("The team put a camera and a radio device side by side. The camera watched people and "
           "automatically produced skeletons. The radio watched the same people through the reflections. "
           "The radio network was trained to output the same skeleton the camera saw. After training, the "
           "camera is removed and the radio alone can draw the skeleton \u2014 including through a wall, where "
           "the camera never could.")
field("How it works.", "A custom FMCW radio (sweeping across a frequency band, with antenna arrays) "
      "produces horizontal and vertical heatmaps of reflections. A teacher-student setup uses a "
      "vision-based pose estimator as the teacher to label the radio data; a convolutional network learns "
      "to predict keypoints from the radio heatmaps. This cross-modal supervision is the key idea.")
field("Data and result.", "Collected in-house with synchronised camera and radio. It demonstrated "
      "believable skeletons through walls and in the dark, conditions where cameras fail outright.")
field("Why it matters.", "It is the flagship proof that radio can do a rich vision task (pose) that "
      "everyone assumed needed a camera, and it popularised the camera-teaches-radio training recipe that "
      "underlies your reduce-the-cameras direction.")
caution_box("It used a specialised FMCW radio, not a commodity Wi-Fi chip, so the hardware is more capable "
            "(and more expensive) than an ESP32. Results are strongest for one or a few people in "
            "controlled settings.")
field("Citation.", "M. Zhao, T. Li, M. Abu Alsheikh, Y. Tian, H. Zhao, A. Torralba, D. Katabi. "
      "\u201cThrough-Wall Human Pose Estimation Using Radio Signals.\u201d IEEE/CVF CVPR 2018.")

# ---- Paper 3: Person-in-WiFi ----
paper("Person-in-WiFi \u2014 Fine-Grained Person Perception Using Commodity Wi-Fi")
field("What it is.", "Body segmentation (the person's outline) and skeleton keypoints from ordinary Wi-Fi, "
      "using just a few commodity antennas rather than a special radar.")
layman_box("RF-Pose used a custom radio. This paper asked: can we do something similar with the cheap "
           "Wi-Fi hardware people already own? It showed that even a handful of standard Wi-Fi antennas "
           "carry enough information to outline a body and place joints, again by learning from a camera "
           "during training.")
field("How it works.", "One transmitter and three receiver antennas capture CSI. A camera provides the "
      "training labels (segmentation masks and keypoints). A deep network maps the 1D Wi-Fi signals to "
      "these 2D body representations. It reduces the hardware requirement from radar to commodity Wi-Fi.")
field("Data and result.", "In-house synchronised Wi-Fi and camera data. It produced coarse but real body "
      "masks and joint locations from Wi-Fi alone \u2014 a strong signal that commodity gear suffices for "
      "meaningful perception.")
field("Why it matters.", "It is the bridge between \u201crich radar sensing\u201d and \u201ccheap Wi-Fi you can "
      "actually deploy.\u201d It is the closest ancestor of DensePose-from-WiFi and directly relevant to a "
      "low-cost ESP32-style build.")
caution_box("Accuracy is coarser than camera or radar; it works best for a small number of people in the "
            "training environment, and generalisation to new rooms is limited.")
field("Citation.", "F. Wang, S. Zhou, S. Panev, J. Han, D. Huang. \u201cPerson-in-WiFi: Fine-grained Person "
      "Perception using WiFi.\u201d IEEE/CVF ICCV 2019.")

# ---- Paper 4: DensePose From WiFi ----
paper("DensePose From WiFi \u2014 Dense Body Surface From Wi-Fi (CMU)")
field("What it is.", "The most detailed Wi-Fi body sensing to date: it maps Wi-Fi signals to a dense "
      "surface of the human body (24 regions with UV coordinates), the same output as the camera-based "
      "DensePose, but from radio.")
layman_box("Instead of just a stick figure, this paints the whole body surface \u2014 which patch of the "
           "signal corresponds to which patch of skin \u2014 using only Wi-Fi. It is the clearest evidence "
           "that a lot of the visual information about a body survives in ordinary Wi-Fi signals.")
field("How it works.", "It uses phase and amplitude from a 3x3 arrangement of commodity Wi-Fi antennas, "
      "borrows the DensePose network design from computer vision, and is supervised by an image-based "
      "DensePose model. The Wi-Fi network learns to output the dense UV body map for multiple people.")
field("Data and result.", "In-house synchronised Wi-Fi and video. The authors report performance in the "
      "same ballpark as some image-based methods on their data, for the dense-pose task, from Wi-Fi only.")
field("Why it matters.", "It sets the current high-water mark for how much body detail Wi-Fi can recover, "
      "and it is strongly motivated by privacy and low cost \u2014 exactly your MPL framing (get camera-like "
      "understanding without keeping cameras everywhere).")
caution_box("It is a research prototype, environment-specific, and evaluated on the authors' own data; it "
            "is not a plug-and-play product, and reproducing it needs careful multi-antenna CSI capture.")
field("Citation.", "J. Geng, D. Huang, F. De la Torre. \u201cDensePose From WiFi.\u201d arXiv:2301.00250, 2023 "
      "(Carnegie Mellon University).")

# ---- Paper 5: XModal-ID ----
paper("XModal-ID \u2014 Through-Wall Person Identification by Matching Wi-Fi to Video (UCSB)")
field("What it is.", "A method to decide whether the person sensed behind a wall by Wi-Fi is the same "
      "person shown in a piece of video footage \u2014 a cross-modal identity match between radio and video.")
layman_box("Imagine security footage of a suspect from one place, and a Wi-Fi signal of someone walking "
           "in another room. XModal-ID asks: are these the same person? It converts the video into the "
           "Wi-Fi signal that such a walk would produce, then compares it to the real measured Wi-Fi. If "
           "they match, it is likely the same person \u2014 all without a camera in the second room.")
field("How it works.", "From a video it generates a simulated Wi-Fi time-frequency signature of that "
      "person's gait, then compares it with the actually measured Wi-Fi signature using their similarity. "
      "This avoids needing paired Wi-Fi-plus-video training data for every target.")
field("Data and result.", "Evaluated on volunteers with through-wall Wi-Fi. It reported ranked "
      "identification accuracy that improves as you allow the correct person to be within the top few "
      "candidates (top-1 lower, top-2 and top-3 higher) among a small candidate set \u2014 promising, but for "
      "modest numbers of people.")
field("Why it matters.", "It is the canonical demonstration that a camera-derived identity can be handed "
      "over to Wi-Fi through a wall. It is the direct scientific backbone of your camera-enrol, "
      "wireless-track idea.")
caution_box("It works for a small candidate set, mainly while walking, and is environment-dependent. It is "
            "identity verification/ranking, not open-world recognition of anyone.")
field("Citation.", "B. Korany, C. R. Karanam, H. Cai, Y. Mostofi. \u201cXModal-ID: Using WiFi for Through-Wall "
      "Person Identification from Candidate Video Footage.\u201d ACM MobiCom 2019.")

# ---- Paper 6: RF-ReID ----
paper("RF-ReID \u2014 Long-Term Person Re-Identification Using Radio Signals (MIT)")
field("What it is.", "Recognising the same person over long periods from radio, using body shape and "
      "movement rather than clothing \u2014 so it still works when the person changes outfits.")
layman_box("Cameras often re-identify people by their clothes, which fails the next day when they change. "
           "Radio does not see clothes; it senses the shape and gait of the body underneath. This paper "
           "builds a per-person radio signature that stays stable across days and clothing changes.")
field("How it works.", "An FMCW radar produces reflections; a multi-task network predicts a skeleton and "
      "learns a person embedding, with a component that discourages the model from latching onto the "
      "environment (so the signature depends on the person, not the room). Same-person samples are pulled "
      "together and different-person samples pushed apart (metric learning).")
field("Data and result.", "Evaluated on in-house campus and home datasets over time. It demonstrated "
      "re-identification that survives clothing changes, which camera systems struggle with.")
field("Why it matters.", "It shows radio identity can be robust and long-term \u2014 the most convincing "
      "evidence for the \u201ctrack a known person across camera-free rooms\u201d goal, and a template for the "
      "signature-matching machinery you would reuse.")
caution_box("It uses a specialised FMCW radar (higher fidelity and cost than Wi-Fi/ESP32), and works for a "
            "known enrolled set of people, not arbitrary strangers.")
field("Citation.", "L. Fan, T. Li, R. Fang, R. Hristov, Y. Yuan, D. Katabi. \u201cLearning Longterm "
      "Representations for Person Re-Identification Using Radio Signals.\u201d IEEE/CVF CVPR 2020.")

# ---- Paper 7: WhoFi ----
paper("WhoFi \u2014 Person Re-Identification via Wi-Fi Channel Encoding (2025, very recent)")
field("What it is.", "A recent pipeline that re-identifies people from Wi-Fi CSI alone using a modern "
      "Transformer encoder, framed explicitly as a privacy-relevant alternative to camera surveillance.")
layman_box("This is the 2025 update of the identity-from-Wi-Fi idea. It takes the channel measurements, "
           "feeds them through a Transformer (the same family of models behind modern AI), and produces a "
           "biometric signature per person that can be matched later. It shows the field is still active "
           "and moving toward stronger, more general models.")
field("How it works.", "Biometric features are extracted from CSI and passed through a modular deep "
      "network with a Transformer-based encoder, trained with an in-batch negative loss (a contrastive "
      "objective that separates different people's signatures).")
field("Data and result.", "Evaluated on the public NTU-Fi dataset, it reports results competitive with "
      "state-of-the-art Wi-Fi identification methods. (It is an arXiv preprint, so treat the exact numbers "
      "as not yet peer-reviewed.)")
field("Why it matters.", "It is the freshest reference point for identity-from-Wi-Fi, uses a public "
      "dataset you can also use, and gives you the surveillance-versus-privacy framing that makes a strong "
      "thesis motivation.")
caution_box("It is a 2025 preprint on a specific public dataset; real-world, cross-environment performance "
            "is unproven, and identity-from-Wi-Fi raises serious privacy questions you should address "
            "head-on rather than ignore.")
field("Citation.", "D. Avola, E. Emam, D. Montagnini, D. Pannone, A. Ranaldi. \u201cWhoFi: Deep Person "
      "Re-Identification via Wi-Fi Channel Signal Encoding.\u201d arXiv:2507.12869, 2025.")

# ---- Paper 8: Vi-Fi ----
paper("Vi-Fi \u2014 Associating Moving People Across Vision and Wireless Sensors (Rutgers)")
field("What it is.", "A method to match people seen by a camera with their phones' wireless signals, so a "
      "camera track and a wireless track can be linked to the same person.")
layman_box("A camera sees several people walking; several phones broadcast wireless and motion data. "
           "Vi-Fi figures out which phone belongs to which person on screen, by correlating how the "
           "on-screen person moves with what the phone's sensors report. This is how you hand a camera "
           "identity to a wireless signal in corridors or at building scale.")
field("How it works.", "It fuses visual tracks with each phone's wireless ranging (FTM) and inertial "
      "(IMU) data, learning to associate the trajectories that move together. The output is a "
      "correspondence between camera detections and wireless device identities.")
field("Data and result.", "The authors released a multimodal Vi-Fi dataset (synchronised vision, "
      "wireless and IMU across many subjects and scenes) and demonstrated reliable association in "
      "multi-person outdoor and indoor settings.")
field("Why it matters.", "It is the cleanest example of vision-to-wireless handoff at scale, and its "
      "public dataset is directly usable if you pursue the corridor/network version of your idea.")
caution_box("It ties identity to a phone the person carries (device-tied), and modern MAC-address "
            "randomisation weakens passive phone identification, so it usually needs a consented app or "
            "ranging feature.")
field("Citation.", "H. Liu, A. Alali, M. Ibrahim, B. B. Cao, N. Meegan, H. Li, M. Gruteser, et al. "
      "\u201cVi-Fi: Associating Moving Subjects across Vision and Wireless Sensors.\u201d ACM/IEEE IPSN 2022.")

# ---- Paper 9: SenseFi ----
paper("SenseFi \u2014 A Library and Benchmark for Deep Wi-Fi Sensing (the practical toolkit)")
field("What it is.", "Not a single new sensing trick, but a benchmark and open-source code library that "
      "compares many deep-learning models across several public Wi-Fi sensing datasets and tasks.")
layman_box("When you start, the hardest questions are \u201cwhich model should I use, on which dataset, and "
           "how do I even load Wi-Fi data?\u201d SenseFi answers exactly that: it gives you ready code, "
           "standard datasets, and a fair comparison of models, so you do not reinvent the plumbing.")
field("How it works.", "It implements common architectures (MLPs, CNNs, RNNs, and Transformers) and "
      "evaluates them on public CSI datasets for tasks such as activity and gesture recognition and "
      "person identification, including transfer- and few-shot-learning settings.")
field("Data and result.", "It aggregates and standardises public datasets (for example UT-HAR and the "
      "NTU-Fi datasets) and reports which model families work best where, providing baselines you can "
      "reproduce.")
field("Why it matters.", "It is the fastest way to get hands-on, get baselines, and understand "
      "generalisation empirically. Start any Wi-Fi sensing project by running SenseFi first.")
caution_box("Benchmarks on curated public datasets tend to be optimistic; real deployments in your own "
            "room will be harder. Treat SenseFi numbers as a ceiling, not a promise.")
field("Citation.", "J. Yang, X. Chen, H. Zou, D. Wang, Q. Xu, L. Xie. \u201cSenseFi: A library and benchmark "
      "on deep-learning-empowered WiFi human sensing.\u201d Patterns (Cell Press), 2023.")

# ===========================================================================
# 3. DATASETS AND TOOLS
# ===========================================================================
h1("3.  Datasets and Capture Tools You Can Actually Use")
para("Everything below is public or commodity, so you can start without special equipment.")
make_table(
    ["Resource", "Type", "What it gives you"],
    [
        ["Widar3.0", "Dataset", "Cross-domain gesture CSI across rooms/orientations; the standard generalisation test."],
        ["NTU-Fi", "Dataset", "CSI for human identification and activity; used by WhoFi and SenseFi."],
        ["UT-HAR", "Dataset", "Early public CSI activity-recognition set; a common baseline."],
        ["Vi-Fi dataset", "Dataset", "Synchronised vision + phone wireless (FTM) + IMU for cross-modal association."],
        ["SenseFi", "Code/benchmark", "Ready models and loaders for the datasets above; your starting point."],
        ["Intel 5300 CSI Tool", "Capture tool", "Classic laptop/PC CSI capture (Halperin et al.); lots of prior work uses it."],
        ["Atheros CSI Tool", "Capture tool", "Alternative commodity CSI capture on Atheros chips."],
        ["Nexmon CSI", "Capture tool", "CSI extraction on Broadcom chips (some phones and the Raspberry Pi)."],
        ["ESP32 CSI Toolkit", "Capture tool", "Low-cost CSI from ESP32 boards \u2014 the cheapest entry, matches your current work."],
    ],
    widths=[1.4, 1.1, 3.8],
)
note_box("Method families you will meet across these papers: convolutional networks on CSI spectrograms; "
         "recurrent/Transformer encoders for time series; teacher-student cross-modal supervision "
         "(camera teaches radio); and metric/contrastive learning for identity. Generalisation is the "
         "recurring weakness \u2014 domain-independent features (like Widar3.0's BVP), data from many "
         "environments, and domain-adaptation are the main cures.")

# ===========================================================================
# 4. HOW THEY CONNECT TO YOUR PROJECT
# ===========================================================================
h1("4.  How These Papers Connect to Your MPL Direction")
para("Read as a group, the papers give you a ready-made ladder from easy to hard, all supporting the "
     "\u201cuse a few cameras to teach cheap wireless sensors, then reduce the cameras\u201d idea:")
bullet("start with presence, motion and gesture (Widar3.0, SenseFi) on ESP32 CSI \u2014 cheap and reliable.",
       "Rung 1: ")
bullet("add activity and coarse pose via camera-teaches-radio supervision (RF-Pose, Person-in-WiFi, "
       "DensePose-from-WiFi).", "Rung 2: ")
bullet("add identity of a small known group via gait/body-shape signatures (RF-ReID, WhoFi, XModal-ID).",
       "Rung 3: ")
bullet("extend to corridors and buildings with vision-to-wireless handoff (Vi-Fi).", "Rung 4: ")
para("Document 2 then takes the separate, more futuristic idea \u2014 making the wireless channel itself do "
     "the neural-network computation (MetaAI and its relatives) \u2014 and assesses how, and whether, it can "
     "be combined with the sensing above.")

closing = doc.add_paragraph(); closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("End of Document 1. Every paper listed is real; citations are provided so each claim "
                    "can be checked at source. Where an exact metric was uncertain it was described "
                    "qualitatively rather than invented.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = SLATE_LT

doc.save(OUTPUT)
print("wrote", OUTPUT)
