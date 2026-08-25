"""Build the Word document explaining the MetaAI paper.

Paper: "Enabling Over-the-Air AI for Edge Computing via Metasurface-Driven
Physical Neural Networks" (MetaAI), SIGCOMM '25, Feng et al.

The document explains the paper at two levels:
  (1) a plain-language, from-scratch account for a complete newcomer, and
  (2) a scientific account for a researcher, in which every equation is
      explained in terms of what it says, why it matters, how it works, and
      what each variable means.

No emojis are used anywhere in the document.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "assets_metaai")
OUTPUT = os.path.join(os.path.dirname(HERE), "MetaAI_OverTheAir_Edge_AI.docx")

# Light, presentable palette
INK = RGBColor(0x33, 0x37, 0x3D)
SLATE = RGBColor(0x4A, 0x6B, 0x8A)
SLATE_LT = RGBColor(0x6E, 0x8C, 0xA8)
ACCENT = RGBColor(0x7A, 0x9A, 0x7E)
PLUM = RGBColor(0x7A, 0x64, 0x9A)
SHADE_BLUE = "EAF1F7"
SHADE_GREEN = "EDF4ED"
SHADE_PLUM = "F1EDF6"
SHADE_GREY = "F2F3F5"
HEADER_FILL = "DCE6EF"

doc = Document()

# ---- base styles -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, SLATE), ("Heading 2", 13.5, SLATE), ("Heading 3", 12, SLATE_LT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="9DB6CC", size=18, where="left"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{where}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), "8")
    e.set(qn("w:color"), color)
    pbdr.append(e)
    pPr.append(pbdr)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def para(text, italic=False, size=11, align=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def runs_para(parts):
    """parts: list of (text, bold, italic)."""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    r2 = p.add_run(text)
    return p


def layman_box(text):
    p = doc.add_paragraph()
    _shade(p, SHADE_GREEN)
    _border(p, color="AFC8B2")
    r = p.add_run("In plain terms:  ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    return p


def sci_box(text):
    p = doc.add_paragraph()
    _shade(p, SHADE_BLUE)
    _border(p, color="9DB6CC")
    r = p.add_run("For the researcher:  ")
    r.bold = True
    r.font.color.rgb = SLATE
    r2 = p.add_run(text)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    return p


def note_box(text):
    p = doc.add_paragraph()
    _shade(p, SHADE_GREY)
    _border(p, color="C9CDD3")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)
    return p


_eqn = {"n": 0}


def equation(img, label=True):
    _eqn["n"] += 1
    from PIL import Image
    path = os.path.join(ASSETS, img)
    with Image.open(path) as im:
        w_in = im.width / 200.0
        h_in = im.height / 200.0
    max_w = 5.9
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if w_in > max_w:
        run.add_picture(path, width=Inches(max_w))
    else:
        target_h = max(min(h_in, 0.6), 0.28)
        run.add_picture(path, height=Inches(target_h))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    if label:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"( {_eqn['n']} )")
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE_LT
        cap.paragraph_format.space_after = Pt(8)
    return _eqn["n"]


_fig = {"n": 0}


def figure(img, caption, width=6.2):
    _fig["n"] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(width))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {_fig['n']}.  ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = SLATE
    r2 = cap.add_run(caption)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = INK
    cap.paragraph_format.space_after = Pt(12)


def variable_table(rows):
    """rows: list of (symbol, meaning)."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(["Symbol", "What it means"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for ri, (sym, mean) in enumerate(rows):
        cells = t.add_row().cells
        cells[0].text = ""
        p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(sym)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = SLATE
        cells[1].text = ""
        p1 = cells[1].paragraphs[0]
        r1 = p1.add_run(mean)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = INK
        if ri % 2 == 1:
            for c in cells:
                tcPr = c._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F6F8FA")
                tcPr.append(shd)
    t.columns[0].width = Inches(1.3)
    t.columns[1].width = Inches(4.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = INK
        _shade(p, HEADER_FILL)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = INK
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F6F8FA")
                tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ===========================================================================
# TITLE PAGE
# ===========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Computing in Thin Air")
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = SLATE
title.paragraph_format.space_before = Pt(70)
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("How MetaAI Turns the Wireless Channel into a Neural Network")
r.font.size = Pt(16)
r.font.color.rgb = SLATE_LT
sub.paragraph_format.space_after = Pt(18)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("A guided explanation of \u201cEnabling Over-the-Air AI for Edge Computing via "
                 "Metasurface-Driven Physical Neural Networks\u201d (SIGCOMM 2025)")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = INK

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("\u2014  Explained for the Newcomer and for the Researcher  \u2014")
r.font.size = Pt(11)
r.font.color.rgb = SLATE_LT
line.paragraph_format.space_before = Pt(24)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Original authors: Chao Feng, Shuo Liang, Chenghui Li, Gaogeng Zhao, Beier Jing, "
                 "Yaxiong Xie, Xiaojiang Chen  \u00b7  Northwest University and University at Buffalo (SUNY)")
r.font.size = Pt(10)
r.font.color.rgb = SLATE_LT

doc.add_page_break()

# ===========================================================================
# TABLE OF CONTENTS
# ===========================================================================
h1("Contents")
toc_p = doc.add_paragraph()
run = toc_p.add_run()
fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
t_run = OxmlElement("w:t"); t_run.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
run._r.append(fldChar); run._r.append(instr); run._r.append(fldChar2); run._r.append(t_run); run._r.append(fldChar3)

note_box("How to read this document. Every major idea appears twice. The green boxes marked "
         "\u201cIn plain terms\u201d explain the concept as if to someone who has never studied wireless "
         "or machine learning. The blue boxes marked \u201cFor the researcher\u201d add the precise, "
         "technical version. Each equation is followed by an explanation of what it says, why it is "
         "used, how it works, and a table defining every symbol in it.")
doc.add_page_break()

# ===========================================================================
# 1. EXECUTIVE SUMMARY
# ===========================================================================
h1("1.  Executive Summary")
para("Modern \u201csmart\u201d devices \u2014 doorbells, sensors, cameras, tags on shelves \u2014 usually cannot "
     "run artificial intelligence (AI) themselves. They are too small, too cheap, and too battery-limited. "
     "So they send their raw data over a wireless link to a powerful server, and the server does the "
     "thinking. In this normal way of working, sending the data and computing on the data are two "
     "separate jobs, each costing time and energy.")
para("The paper introduces MetaAI, which asks a bold question: what if the wireless signal did the "
     "computation on its way to the server, so that sending and computing became a single job? MetaAI "
     "achieves exactly this by placing a programmable \u201csmart surface\u201d (a metasurface) in the room. "
     "As the device\u2019s radio signal bounces off this surface on its way to the receiver, the surface "
     "reshapes the signal in a carefully chosen way. That reshaping is mathematically identical to the "
     "multiplications and additions that a neural network performs. The receiver therefore gets the "
     "answer (for example, \u201cthis is a cat\u201d) rather than the raw data.")
layman_box("A normal setup is like posting a photo to an expert and waiting for a reply. MetaAI is like "
           "installing a magic mirror in the hallway: as your message passes the mirror it is turned into "
           "the expert\u2019s verdict, so the person at the other end receives the verdict directly. The "
           "clever part is that the \u201cmirror\u201d (a metasurface) can be reprogrammed to hold whatever the "
           "AI model has learned.")
sci_box("MetaAI is an over-the-air computing paradigm that embeds a linear neural network into the "
        "wireless propagation channel. A reconfigurable metasurface synthesises a time-varying channel "
        "response H(t) whose successive values equal the trained weights of a complex-valued fully "
        "connected layer. The transmitter streams the input vector one symbol at a time; the receiver "
        "accumulates the products, computing a matrix-vector product physically during transmission. A "
        "single 16 x 16 (256-element), 2-bit metasurface prototype at 2.4/5 GHz and 3.5 GHz reaches "
        "82.8 percent average classification accuracy (up to 89.8 percent) across six datasets, with "
        "multi-sensor late fusion adding up to 27 percent, while remaining compatible with commodity "
        "wireless links.")
para("This document explains the whole idea from the ground up. It covers why the problem matters, the "
     "two obstacles the authors had to overcome, the mathematics that makes it possible (with every "
     "equation and variable explained), the engineering that makes it work in a real room, and the "
     "measured results, strengths, and limitations.")

# ===========================================================================
# 2. THE BIG IDEA IN PLAIN LANGUAGE
# ===========================================================================
h1("2.  The Big Idea, in Plain Language")
para("Start with three everyday facts and join them together.")
bullet("radio signals are waves. When your phone or a sensor sends data, it sends invisible radio "
       "waves through the air, exactly like ripples spreading on water.", "Fact 1 \u2014 ")
bullet("waves change as they travel. When a wave bounces off walls, furniture, or a person, it is "
       "bent, delayed, and weakened. Engineers call this whole effect \u201cthe channel\u201d. The channel "
       "is really just a rule that says: whatever wave you send in, this is the different wave that "
       "comes out.", "Fact 2 \u2014 ")
bullet("a neural network is mostly multiply-and-add. The heart of a simple neural network is: take "
       "each input number, multiply it by a learned weight, and add the results up. That sum is the "
       "answer (or the score for each possible answer).", "Fact 3 \u2014 ")
para("Now the key realisation: \u201cmultiply the wave by something and add the results up\u201d is precisely "
     "what a channel already does to a radio signal. If we could choose the exact amount by which the "
     "channel multiplies the signal at each moment, we could make the channel carry out the neural "
     "network\u2019s multiply-and-add for us. A metasurface is the device that lets us choose it.")
figure("fig_paradigms.png",
       "Three ways to run AI on data from a small device. (a) The usual way: send raw data to a server, "
       "then compute there \u2014 two separate costs. (b) A physical neural network still needs the data "
       "sent first; the wave only powers a fixed structure. (c) MetaAI: the metasurface computes while "
       "the signal travels, so the server receives the result directly.", width=6.6)
layman_box("A metasurface is a flat panel covered with hundreds of tiny controllable cells. Each cell "
           "can nudge the radio wave that hits it \u2014 speeding it up or slowing it down a little. By "
           "setting all the cells together, we control the overall reshaping of the wave. MetaAI sets "
           "them so the reshaping equals the AI model\u2019s calculation.")

# ===========================================================================
# 3. WHY THIS MATTERS AND WHAT CAME BEFORE
# ===========================================================================
h1("3.  Why It Matters, and What Came Before")

h2("3.1  The problem MetaAI attacks")
para("Edge AI is the practice of running AI close to where data is produced, on or near small "
     "\u201cInternet of Things\u201d (IoT) devices. It faces a hard trade-off. Running AI on the device "
     "gives fast answers but is often impossible: the device lacks the processing power and the battery. "
     "The common alternative \u2014 the \u201ctransmit then compute\u201d model \u2014 sends the raw data to an edge "
     "server, which logs it and computes on it. In that model, communication and computation are two "
     "separate, sequential costs in both energy and time. MetaAI\u2019s goal is to merge them.")

h2("3.2  Physical neural networks: fast, but still separate")
para("A promising earlier idea is the Physical Neural Network (PNN). A PNN uses a physical effect \u2014 "
     "usually the way waves spread and bend (diffraction) \u2014 to perform many multiplications and "
     "additions at once, at the speed of light, using almost no electronic energy. In an optical PNN, "
     "light passes through patterned glass plates; in a radio PNN, a signal passes through stacked "
     "metasurfaces. The wave itself does the arithmetic.")
para("But traditional PNNs share a weakness. They are co-processors: the input data must first be fully "
     "delivered to the structure, and a light or radio source is used merely to \u201cpower\u201d the "
     "computation, not to carry information. So communication and computation remain separate. Worse, "
     "because all inputs strike a cell at the same instant, a single cell cannot give each input its own "
     "weight, so many stacked layers are needed just to approximate one simple layer.")
sci_box("A PNN implements linear operations through two physical mechanisms: multiplication as a wave "
        "traverses a transmissive metasurface (each meta-atom modulates phase or amplitude) and addition "
        "through free-space superposition. Because superimposed inputs at a single meta-atom are "
        "modulated jointly, a one-layer PNN cannot realise an arbitrary matrix; stacking L layers adds "
        "degrees of freedom that asymptotically approach the target linear map (the paper proves this in "
        "its Appendix A.1).")

h2("3.3  Over-the-air computing: addition was easy, multiplication was hard")
para("Another line of work, over-the-air computing, exploited the fact that when several radio signals "
     "arrive together they naturally add up. That gives addition for free. The trouble was "
     "multiplication: earlier systems performed it with complex, expensive processing inside the "
     "transmitter, which small commodity devices cannot do. A pioneering system called AirNN used "
     "metasurfaces to compute a convolution in the air, but it needed a complicated multi-antenna relay "
     "steering signals to several separate metasurfaces, making it a specialised apparatus rather than a "
     "normal communication link.")
runs_para([("The gap MetaAI fills.  ", True, False),
           ("The authors ask whether a simple, single-metasurface design can implement a complete, "
            "end-to-end neural network that still works over an ordinary wireless link with an ordinary "
            "IoT transmitter. MetaAI answers yes.", False, False)])

# ===========================================================================
# 4. HOW METAAI WORKS: THE TWO CHALLENGES
# ===========================================================================
h1("4.  How MetaAI Works: Two Challenges and Two Insights")
para("Turning the channel into a neural network requires solving two problems. The paper frames the "
     "whole design around them.")

h2("4.1  Challenge 1: sequential transmission versus parallel computation")
para("Wireless systems send data one small piece (one \u201csymbol\u201d) at a time, in sequence. A neural "
     "network, by contrast, normally wants all its inputs at once, in parallel. These two habits seem "
     "incompatible.")
para("The insight: because the network is linear, its parallel calculation can be broken into a "
     "sequence without changing the answer. A sum such as \u201cweight one times input one, plus weight two "
     "times input two, and so on\u201d can be built up one term at a time. So we multiply the first input by "
     "its weight now, the second input by its weight a moment later, and keep a running total. The final "
     "total is identical to computing everything at once.")
figure("fig_sequential.png",
       "The linearity trick. A parallel weighted sum (left) is mathematically identical to computing one "
       "product per time slot and accumulating them (right). This lets a naturally sequential wireless "
       "link carry out a neural-network layer.", width=6.5)
layman_box("Adding a column of numbers gives the same total whether you add them all in one glance or "
           "one at a time down the page. MetaAI adds them one at a time \u2014 which is exactly how a radio "
           "link likes to send them.")

h2("4.2  Challenge 2: doing the multiplication in the wave itself")
para("The second problem is where the actual multiply-and-add happens. Here MetaAI uses the basic "
     "behaviour of a wireless channel: as a signal travels, it is linearly transformed by the channel. "
     "If we call the transmitted signal x(t) and the channel H(t), the received signal is their product. "
     "By programming the metasurface to make the channel take a specific value at each instant, we make "
     "the channel multiply each input by exactly the weight the neural network wants.")
equation("eq_channel.png")
para("This tiny equation is the seed of the whole system. It states that the signal coming out of the "
     "channel is the channel response multiplied by the signal going in.")
variable_table([
    ("x(t)", "the signal the device transmits at time t (one input value of the network, encoded as a radio symbol)."),
    ("H(t)", "the wireless channel response at time t \u2014 how much the channel scales and phase-shifts the signal. The metasurface makes this controllable."),
    ("y(t)", "the signal that arrives at the receiver: the input already multiplied by the chosen weight."),
])
sci_box("Signals here are complex-valued: a radio symbol carries both amplitude and phase, so it is a "
        "complex number, and H(t) is a complex gain. Complex multiplication scales magnitude and rotates "
        "phase simultaneously, which is precisely a complex weight in a complex-valued neural network. "
        "This is why MetaAI trains a complex-valued network rather than a real one.")

# ===========================================================================
# 5. THE MATHEMATICS, EXPLAINED
# ===========================================================================
h1("5.  The Mathematics, Explained Step by Step")
para("This section walks through the equations that define MetaAI. For each one we give the plain "
     "meaning, the reason it appears, how it works, and a table defining every symbol.")

h2("5.1  A linear neural network is a matrix times a vector")
para("The simplest classifier MetaAI uses is a single fully connected linear layer. It takes the input "
     "as a list of numbers (a vector) and produces one score per possible category by multiplying with "
     "a table of learned weights (a matrix).")
equation("eq_lnn.png")
runs_para([("What it says.  ", True, False),
           ("The output vector Y equals the weight matrix W multiplied by the input vector X. Each output "
            "score is a weighted sum of all the inputs.", False, False)])
runs_para([("Why it is used.  ", True, False),
           ("Classification means turning raw input into a score for each class; a matrix-vector product "
            "is the most basic way to do that, and it is entirely linear, which is what makes it "
            "implementable in the wireless channel.", False, False)])
runs_para([("How it works.  ", True, False),
           ("Row r of W holds the weights for class r. Multiplying that row by the input and summing "
            "gives the score for class r. Doing this for every row yields the full output vector.", False, False)])
variable_table([
    ("Y", "the output vector; entry y_r is the score (class probability) for category r."),
    ("W", "the weight matrix, size R by U; entry w_(r,i) links input i to output r. These are the numbers learned during training."),
    ("X", "the input vector of length U (the data sample, encoded as complex symbols)."),
    ("R", "the number of output categories (for example, 10 for the ten digits 0-9)."),
    ("U", "the number of input values (the length of the data sample)."),
])
note_box("A useful property: because every operation is linear, several stacked linear layers can always "
         "be collapsed into one equivalent layer. That is why MetaAI needs only a single layer to "
         "represent any linear transformation \u2014 and why a single metasurface can suffice.")

h2("5.2  Turning the parallel sum into a sequence")
para("Focusing on one output score shows how the parallel sum becomes a step-by-step process.")
equation("eq_decomp.png")
runs_para([("What it says.  ", True, False),
           ("The score for the first class, y_1, is the sum over all inputs of each input x_i multiplied "
            "by its weight w_(1,i).", False, False)])
runs_para([("Why it is used.  ", True, False),
           ("Written as a sum of separate products, the calculation no longer needs all inputs at once. "
            "Each product can be formed at a different moment and added to a running total \u2014 the exact "
            "match to sequential wireless transmission described in Section 4.1.", False, False)])
runs_para([("How it works.  ", True, False),
           ("At time slot i the device sends x_i, the channel multiplies it by w_(1,i), and the receiver "
            "adds the result to what it already has. After U slots the running total equals y_1.", False, False)])
variable_table([
    ("y_1", "the score for the first category (the same idea applies to every category)."),
    ("x_i", "the i-th input value, transmitted during time slot i."),
    ("w_(1,i)", "the weight connecting input i to output 1, applied by the channel at time slot i."),
    ("U", "the number of inputs, and therefore the number of time slots used."),
])

h2("5.3  The core MetaAI computation")
para("Combining the two insights gives the equation at the heart of the paper: the receiver forms each "
     "output by accumulating, over time, the transmitted symbols after the metasurface has multiplied "
     "each one by the desired weight.")
equation("eq_core.png")
runs_para([("What it says.  ", True, False),
           ("Output y_r is the magnitude of the sum, over all input indices i, of the channel response "
            "H_r(t_i) multiplied by the input x_i.", False, False)])
runs_para([("Why it is used.  ", True, False),
           ("This is the neural-network layer of Section 5.1, but realised physically: the multiplication "
            "happens over the air inside the channel, and only the running addition is done in software "
            "at the receiver. The absolute-value bars turn the complex result into a real score and "
            "provide the network\u2019s nonlinearity.", False, False)])
runs_para([("How it works.  ", True, False),
           ("To compute the score for class r, the metasurface is set so that at time t_i its channel "
            "response equals the weight w_(r,i). The device transmits x_i, the channel multiplies, and "
            "the receiver accumulates. The magnitude of the final sum is the class score.", False, False)])
variable_table([
    ("y_r", "the final score for category r."),
    ("H_r(t_i)", "the channel response used to compute output r at time t_i; the metasurface makes this equal to the trained weight w_(r,i)."),
    ("x_i", "the i-th transmitted input symbol."),
    ("| . |", "the magnitude (absolute value) of the complex sum, giving a real, non-negative score and acting as the nonlinearity."),
    ("U", "the number of inputs accumulated."),
])
sci_box("Multiplication is performed physically (over the air); accumulation is performed digitally at "
        "the receiver. Because each input occupies its own time slot, MetaAI can assign an independent "
        "weight to each input, so a single-layer, single-metasurface design achieves what a multi-layer "
        "stacked PNN needs many surfaces to approximate.")

h2("5.4  How the metasurface creates the weight")
para("Next we need to know what channel a metasurface actually produces, so we can set it to the value "
     "we want. The metasurface is an array of small elements called meta-atoms; each reflects the signal "
     "with a controllable phase shift. The total channel through the metasurface is the sum of the "
     "contributions of all the meta-atoms.")
equation("eq_mts.png")
runs_para([("What it says.  ", True, False),
           ("The metasurface channel H_mts is an overall amplitude factor times the sum, over all "
            "meta-atoms, of two phase rotations: the phase the wave picks up travelling to and from that "
            "atom, and the phase the atom itself deliberately adds.", False, False)])
runs_para([("Why it is used.  ", True, False),
           ("It connects the physical control knob we actually have \u2014 the phase of each meta-atom \u2014 "
            "to the channel value the neural network needs. Only by modelling this can we solve for the "
            "right settings.", False, False)])
runs_para([("How it works.  ", True, False),
           ("Each meta-atom sends back a small wave. Those waves add up at the receiver. By changing each "
            "atom\u2019s deliberate phase, we steer how the contributions combine, and thus tune the total "
            "channel to a chosen complex value.", False, False)])
variable_table([
    ("H_mts", "the total channel through the metasurface path (the physical realisation of a weight)."),
    ("M", "the number of meta-atoms on the metasurface (256 in the prototype)."),
    ("phi_m", "the phase shift the m-th meta-atom deliberately adds \u2014 the only quantity we can actively control."),
    ("phi^p_m", "the phase the wave naturally acquires along its propagation path to and from atom m."),
    ("alpha_p", "an overall amplitude factor common to all atoms (far-field assumption); it scales every output equally and so does not change the classification."),
])
para("The natural path phase depends on distance, through the wavelength of the signal:")
equation("eq_wavenumber.png")
runs_para([("What it says and how it works.  ", True, False),
           ("The propagation phase for atom m is the wavenumber k_0 times the total travelled distance "
            "(transmitter to atom, plus atom to receiver). The wavenumber k_0 = 2 pi / lambda converts a "
            "distance into a phase angle: every wavelength travelled adds a full turn of 2 pi.", False, False)])
variable_table([
    ("k_0", "the wavenumber, equal to 2 pi divided by the wavelength lambda; it turns a distance into a phase."),
    ("lambda", "the wavelength of the radio signal (about 6 cm at 5 GHz)."),
    ("d_(Tx,m)", "distance from the transmitter to the m-th meta-atom (fixed and known)."),
    ("d_(m,Rx)", "distance from the m-th meta-atom to the receiver."),
])
para("The receiver distance is not known in advance, but under the far-field assumption the reflected "
     "rays are nearly parallel, which simplifies it to a neat, regular pattern:")
equation("eq_farfield.png")
runs_para([("What it says.  ", True, False),
           ("The distance from atom m to the receiver equals the distance from the first atom, minus a "
            "steady step that grows with the atom index m and depends on the viewing angle theta.", False, False)])
runs_para([("Why and how it matters.  ", True, False),
           ("The leading term (the distance to the first atom) is common to every atom, so, like the "
            "amplitude factor, it scales all outputs equally and cannot change which class wins. That "
            "means the receiver\u2019s exact position is not needed \u2014 only the angle theta, which is found "
            "with standard beam scanning. A hard three-dimensional positioning problem collapses into a "
            "one-angle estimation.", False, False)])
variable_table([
    ("d_(m,Rx)", "distance from atom m to the receiver."),
    ("d_(1,Rx)", "distance from the first atom to the receiver (a common factor, so it does not affect the result)."),
    ("d_s", "the fixed spacing between neighbouring meta-atoms."),
    ("theta", "the angle between the receiver direction and the metasurface plane; the only unknown that matters, obtained by beam scanning."),
    ("m", "the meta-atom index, running from 1 to M."),
])

h2("5.5  Solving for the metasurface settings")
para("With a model of what the metasurface produces, and a target weight from training, we choose the "
     "meta-atom phases that make the two match as closely as possible.")
equation("eq_goal.png", label=True)
para("The goal, stated simply, is that the metasurface channel should equal the desired (trained) "
     "weight. Because real hardware cannot hit the target exactly, we minimise the difference instead:")
equation("eq_config.png")
runs_para([("What it says.  ", True, False),
           ("Choose the set of meta-atom phases Phi that makes the metasurface channel H_mts as close as "
            "possible to the desired weight H_des, by minimising the magnitude of their difference.", False, False)])
runs_para([("Why it is used.  ", True, False),
           ("This is the bridge from the trained digital model to the physical device. Training produces "
            "ideal continuous weights; this optimisation finds the best achievable hardware setting for "
            "each of them.", False, False)])
runs_para([("How it works.  ", True, False),
           ("Each meta-atom offers only a few discrete phase states (four, in the 2-bit prototype). The "
            "optimiser searches these discrete choices across all atoms to bring the realised channel "
            "closest to the target. Using more meta-atoms gives denser coverage of the complex plane and "
            "a better fit \u2014 with diminishing returns past 256 atoms.", False, False)])
variable_table([
    ("Phi", "the list of chosen phase shifts for all meta-atoms, [phi_1, ..., phi_M]; the solution we solve for."),
    ("phi_m", "the phase of the m-th atom, restricted to the discrete states the hardware supports."),
    ("H_mts", "the channel the chosen phases actually produce (from Equation 5)."),
    ("H_des", "the desired weight obtained from training the digital network."),
    ("arg min", "the operation that returns the phases minimising the following quantity."),
])
figure("fig_metasurface.png",
       "The metasurface is a grid of programmable meta-atoms (left). In the 2-bit prototype each atom "
       "can pick one of four phase states (right): 0, pi/2, pi, or 3pi/2. Setting all 256 atoms together "
       "realises the desired complex weight as closely as the discrete states allow.", width=5.8)

# ===========================================================================
# 6. MAKING IT WORK IN A REAL ROOM
# ===========================================================================
h1("6.  Making It Work in a Real Room")
para("A clean equation assumes an ideal world. A real deployment adds reflections, timing errors, and "
     "noise. MetaAI includes a specific remedy for each, and also ways to go faster and to combine "
     "several sensors.")

h2("6.1  From trained model to configured surface")
para("The workflow ties the mathematics together into a repeatable procedure.")
figure("fig_workflow.png",
       "The MetaAI workflow: train a complex-valued linear network in software, read out its desired "
       "weights, solve for the meta-atom phases that reproduce them, load those phases onto the "
       "metasurface, and then let the signal compute in the air at run time.", width=6.6)
para("A subtle but important choice: MetaAI trains with fully continuous weights and only afterwards "
     "snaps them to the metasurface\u2019s discrete states. The paper shows this continuous-then-discretise "
     "strategy clearly beats training a network that is forced to use discrete weights from the start.")

h2("6.2  Cancelling the echoes (multipath)")
para("In any room the signal also arrives by other paths \u2014 reflections off walls and furniture. These "
     "extra copies, called multipath, would corrupt the computation. MetaAI removes them with an elegant "
     "trick based on how digital symbols are designed.")
para("Communication symbols are deliberately built to average to zero over their period (this keeps the "
     "transmission balanced and helps timing recovery). The environmental echoes inherit this "
     "zero-average property. MetaAI, however, deliberately varies the metasurface weights within a "
     "single symbol period, which breaks the zero-average property for the metasurface path only. So "
     "when the receiver sums samples across the symbol, the echoes cancel to zero while the wanted "
     "metasurface signal survives.")
figure("fig_multipath.png",
       "Multipath cancellation. (a) A communication symbol averages to zero over its period. (b) The "
       "environmental echoes keep that zero average and therefore cancel when summed. (c) Because the "
       "metasurface changes its weight within the symbol, its contribution does not average to zero and "
       "is preserved.", width=6.6)
layman_box("Imagine everyone in a choir singing notes that cancel out to silence over one bar \u2014 that "
           "is the background echo. MetaAI makes its own singer change notes mid-bar so that this one "
           "voice does not cancel and can still be heard clearly above the silence.")
para("If instead the fixed room response is known, it can simply be subtracted inside the optimisation, "
     "though that needs a static environment:")
equation("eq_multipath.png")
runs_para([("What it says.  ", True, False),
           ("Find the phases that make the metasurface channel match the desired weight after removing "
            "the known environmental response H_e.", False, False)])
variable_table([
    ("H_e", "the environmental (multipath) channel from all the fixed reflectors in the room."),
    ("H_des", "the desired weight from training."),
    ("H_mts", "the metasurface channel produced by the chosen phases."),
    ("Phi", "the meta-atom phases being solved for."),
])

h2("6.3  Getting the timing right (synchronisation)")
para("The transmitter and the metasurface do not share a clock, so the stream of inputs can drift out "
     "of step with the stream of weights. Even a few microseconds of drift can collapse accuracy \u2014 a "
     "4-microsecond error dropped recognition to about 25 percent in the authors\u2019 tests. MetaAI fixes "
     "this with a two-part scheme called Coarse-Grained Detection and Fine-Grained Adjustment (CDFA).")
bullet("a low-power energy detector on the metasurface senses when the "
       "incident signal arrives and triggers the weight sequence to start. This gets the timing "
       "roughly right but not perfectly.", "Coarse-grained detection \u2014 ")
bullet("the remaining error is absorbed during training. The training "
       "data is deliberately shifted by random amounts (drawn from a Gamma distribution that matches "
       "the real timing errors), so the learned model becomes robust to drift. No extra hardware or "
       "expensive shared clock is needed.", "Fine-grained adjustment \u2014 ")
sci_box("The fine-grained stage injects synchronisation error into training by cyclically shifting the "
        "input vector by a random number of positions drawn from a Gamma distribution matched to "
        "measured timing statistics. This data-augmentation approach lifts tolerance so that accuracy "
        "holds until roughly 4 microseconds of drift, versus a sharp collapse without it. In the "
        "reported ablation, no-sync accuracy is about 19 percent, coarse detection alone about 56 "
        "percent, and full CDFA about 89 percent.")

h2("6.4  Coping with noise")
para("Two kinds of noise disturb the result: hardware noise (small differences between meta-atoms) and "
     "environmental noise. MetaAI models both and, crucially, trains the network to expect them.")
equation("eq_noise.png")
runs_para([("What it says.  ", True, False),
           ("The computed output includes, at each step, a hardware-noise term N_d added to the "
            "metasurface weight, and an overall environmental-noise term N_e added to the sum.", False, False)])
para("A short rearrangement shows the hardware noise can be treated as if the input signal itself had "
     "been slightly corrupted before transmission:")
equation("eq_noise2.png")
runs_para([("Why this matters.  ", True, False),
           ("Written this way, both noise sources look like a signal that was pre-disturbed by noise. "
            "That means MetaAI can simulate them simply by lowering the signal-to-noise ratio during "
            "training. A network trained on noisier data copes better with real noise at run time \u2014 an "
            "80th-percentile accuracy improvement from about 80.5 to 87.9 percent in the experiments.", False, False)])
variable_table([
    ("N_d", "hardware noise, mainly phase differences among the meta-atoms."),
    ("N_e", "environmental noise added to the final combined result."),
    ("N_d (hat)", "the equivalent input-referred noise, equal to x_i divided by the weight, times N_d."),
    ("H_mts(t_i)", "the metasurface weight applied at step i."),
    ("x_i", "the i-th input symbol."),
])

h2("6.5  Going faster: parallelism")
para("Computing one class score needs one full pass of the input. For R classes that is R sequential "
     "passes, which is slow. MetaAI recovers speed with two parallelism schemes that compute several "
     "class scores at once.")
bullet("different subcarriers (frequency slots in an OFDM signal) carry "
       "different weight patterns, so several classes are computed in parallel across frequency.", "Subcarrier-based \u2014 ")
bullet("several receive antennas each act as an independent output "
       "neuron, computing a different class score at the same time.", "Antenna-based \u2014 ")
para("Because a single metasurface applies one physical phase per atom at any instant, the different "
     "parallel paths cannot be weighted fully independently, so the authors solve one joint optimisation "
     "across all of them:")
equation("eq_parallel.png")
runs_para([("What it says and how it works.  ", True, False),
           ("This is a training loss (a cross-entropy) summed over the K parallel channels. It rewards "
            "phase settings phi_(i,m) that make the correct class score large across every subcarrier or "
            "antenna at once. Minimising it yields one set of metasurface phases that serves all parallel "
            "outputs. Parallelism trades a little accuracy for a large cut in latency.", False, False)])
variable_table([
    ("loss", "the training objective to minimise (cross-entropy over the parallel outputs)."),
    ("K", "the number of parallel channels (subcarriers or antennas), equal to the number of classes."),
    ("y_k", "the ground-truth label: 1 for the correct class, 0 otherwise."),
    ("x_(i,k)", "input i as seen on parallel channel k."),
    ("phi_(i,m)", "the phase assigned to atom m for input i (the quantities being optimised)."),
    ("phi^p_(m,k)", "the propagation phase for atom m on parallel channel k."),
])

h2("6.6  Combining several sensors (multi-sensor fusion)")
para("Real deployments often use several sensors \u2014 cameras from different angles, or a camera plus a "
     "microphone plus a motion sensor. MetaAI supports this naturally. Because the weights for different "
     "sensors are independent, each sensor\u2019s data is processed in its own time window through the same "
     "shared metasurface, and the results are combined.")
figure("fig_multisensor.png",
       "Multi-sensor fusion improves accuracy. Combining several camera views, several antennas, or "
       "different modalities (accelerometer plus gyroscope) all raise recognition accuracy over a single "
       "sensor, using one shared metasurface via time-division.", width=6.0)
para("Each sensor produces its own partial score, and the final score is the combination across all "
     "sensors:")
equation("eq_sensor.png")
equation("eq_fusion.png")
runs_para([("What they say.  ", True, False),
           ("The first expression is the score for class r computed from sensor s alone (the same "
            "accumulate-over-time rule as before, using that sensor\u2019s weights and data). The second "
            "combines the per-sensor scores across all N_s sensors into the final score.", False, False)])
runs_para([("Why and how.  ", True, False),
           ("Handling sensors in separate time windows lets one metasurface serve them all, avoiding the "
            "separate hardware that traditional PNNs would need per sensor. Late fusion \u2014 summing the "
            "per-sensor evidence \u2014 lets complementary sensors reinforce the correct answer, which is why "
            "accuracy rises by up to 27 percent.", False, False)])
variable_table([
    ("y_r^s", "the score for category r from sensor s."),
    ("H_r^s(t_i^s)", "the metasurface weight for sensor s, output r, at its time step i."),
    ("x_i^s", "the i-th input from sensor s."),
    ("U_s", "the number of inputs from sensor s."),
    ("N_s", "the total number of sensors being fused."),
    ("y_r^multi", "the final fused score for category r, combining all sensors."),
])

# ===========================================================================
# 7. THE HARDWARE
# ===========================================================================
h1("7.  The Hardware in Brief")
para("MetaAI is not only theory; the authors built and measured it. The key facts are collected below.")
make_table(
    ["Component", "What the authors used"],
    [
        ["Metasurfaces", "Two prototypes: one dual-band (2.4 and 5 GHz) and one single-band (3.5 GHz)."],
        ["Meta-atom grid", "16 x 16 = 256 meta-atoms per surface."],
        ["Per-atom control", "Two PIN diodes give 4 phase states (0, pi/2, pi, 3pi/2): a 2-bit shifter."],
        ["Controller", "An STM32 microcontroller drives the atoms via shift registers, in parallel groups."],
        ["Radios", "USRP X310 software-defined radios as transmitter and receiver."],
        ["Default settings", "Carrier 5.25 GHz, 256-QAM modulation, 1 M symbols/second, MTS switching up to 2.56 MHz."],
        ["Environments", "Corridor, laboratory, and office; also non-line-of-sight and cross-room tests."],
    ],
    widths=[1.9, 4.4],
)
para("Why 256 meta-atoms? More atoms give a denser set of achievable weights and better accuracy, but "
     "the benefit levels off. The authors measure both accuracy and a \u201cweight distribution density\u201d "
     "metric and find both saturate around 256 atoms, making it the sweet spot between performance and "
     "cost.")
figure("fig_metaatoms.png",
       "Choosing the metasurface size. Recognition accuracy (blue) and the weight distribution density "
       "metric (green) both rise sharply and then saturate near 256 meta-atoms, which the authors adopt "
       "as the best trade-off between accuracy and hardware cost.", width=5.4)

# ===========================================================================
# 8. RESULTS
# ===========================================================================
h1("8.  What the Experiments Show")
para("MetaAI was tested on six datasets spanning handwritten digits, clothing, fruit, animal faces, "
     "human faces, and hand gestures. Even with just a single linear layer, the physical prototype "
     "classifies well.")
figure("fig_accuracy.png",
       "Prototype accuracy across six datasets. MetaAI (blue) stays close to a full digital ResNet-18 "
       "reference (grey) and clearly beats a network constrained to discrete weights from the start "
       "(orange), confirming the continuous-then-discretise strategy.", width=6.4)
para("The headline numbers reported by the authors:")
bullet("82.8 percent average accuracy across the six datasets, peaking at 89.8 percent on the digits.")
bullet("the prototype trails its own digital simulation by no more than about 7 percent, showing the "
       "physical implementation is faithful.")
bullet("it beats a discrete-from-the-start baseline (DiscreteNN) on every dataset, validating training "
       "with continuous weights and mapping to discrete states afterwards.")
bullet("multi-sensor fusion adds up to 27.06 percent (across modalities such as accelerometer plus "
       "gyroscope) and up to 25 percent (across several sensors of the same type).")
para("It also stays robust under stress. Each remedy from Section 6 was verified separately: the CDFA "
     "timing scheme lifts accuracy from about 19 percent to about 89 percent; multipath cancellation "
     "keeps accuracy above about 82 percent in three different rooms; the parallelism schemes add speed "
     "with only a small accuracy cost; and the noise-aware training lifts the 80th-percentile accuracy "
     "from about 80.5 to 87.9 percent. MetaAI further works across 2.4, 3.5, and 5 GHz bands, several "
     "modulation schemes, a range of distances and angles, non-line-of-sight corners, and across rooms, "
     "and a real-time face-recognition case study with IoT cameras reached about 78.5 percent.")

# ===========================================================================
# 9. STRENGTHS, LIMITATIONS, FUTURE
# ===========================================================================
h1("9.  Strengths, Limitations, and Where It Goes Next")

h2("9.1  Why it is compelling")
figure("fig_benefits.png",
       "The main practical benefits: small devices are spared from running AI, the server sees results "
       "rather than raw data (privacy by design), one metasurface is shared by many sensors, and the "
       "whole thing runs over standard commodity wireless links.", width=6.4)
bullet("energy and hardware. The IoT device only transmits; it never runs the AI, which can extend "
       "battery life and lower cost.", "Saves ")
bullet("by design. The server receives the inference result, not the raw images or signals, so "
       "sensitive data need never leave the device in raw form.", "Private ")
bullet("and shareable. A single reconfigurable metasurface implements a whole network and can serve "
       "many sensors by time-division.", "Simple ")
bullet("with standards. It works with ordinary transmitters and links \u2014 no exotic pre-coding or "
       "custom radios.", "Compatible ")

h2("9.2  Honest limitations")
bullet("only. MetaAI currently implements linear networks. Deeper, nonlinear architectures (such as "
       "Transformers) would need nonlinear physical components, which is left as future work.", "Linear ")
bullet("size versus latency. Because inputs are sent sequentially, a larger model means more "
       "transmissions and higher latency; model size is bounded by the latency budget.", "Model ")
bullet("limits precision. Accuracy is capped by the metasurface resolution \u2014 the number of atoms and "
       "their bit-depth (2-bit here).", "Hardware ")
bullet("When the transmitter or receiver moves, the propagation paths change and the phase-to-weight "
       "mapping must be recomputed; supporting fast motion is a race between the target\u2019s speed and "
       "this recalibration.", "Mobility.  ")

h2("9.3  The takeaway")
para("MetaAI reframes the wireless channel from a passive pipe into an active part of the computer. By "
     "exploiting the shared linearity of wireless propagation and neural networks, and by programming a "
     "single metasurface to become the network\u2019s weights, it unifies communication and computation into "
     "one step. For always-on, low-power, privacy-sensitive edge sensing, that is a genuinely new and "
     "practical direction \u2014 and a template that future nonlinear and larger-scale designs can build on.")

# ===========================================================================
# 10. GLOSSARY
# ===========================================================================
h1("10.  Glossary of Key Terms")
make_table(
    ["Term", "Plain meaning"],
    [
        ["Edge AI", "Running AI close to where data is produced, on or near small devices, rather than in a distant cloud."],
        ["IoT device", "A small, cheap, often battery-powered internet-connected sensor or gadget."],
        ["Wireless channel", "The rule describing how a transmitted radio wave is changed (scaled, delayed, echoed) before it is received."],
        ["Metasurface / meta-atom", "A flat panel of many tiny controllable cells (meta-atoms); each cell nudges the phase of the reflected wave."],
        ["Physical Neural Network (PNN)", "A device that performs neural-network arithmetic using a physical wave effect instead of digital circuits."],
        ["Over-the-air computing", "Doing computation on signals as they travel through the air, using the physics of wave combination."],
        ["Linear neural network", "A network whose layer is a plain weighted sum (matrix times vector), with no nonlinear bending in the middle."],
        ["Complex-valued weight", "A weight that both scales and rotates the signal, matching the amplitude-and-phase nature of radio symbols."],
        ["Multipath", "Extra copies of a signal that arrive via reflections off walls and objects."],
        ["Synchronisation (CDFA)", "Keeping the input stream and the weight stream in step, via coarse energy detection plus training-time error injection."],
        ["Subcarrier", "One of many closely spaced frequency slots in an OFDM signal, usable as a parallel computing lane."],
        ["Late fusion", "Combining several sensors by computing each one's result first and then merging the results."],
    ],
    widths=[2.2, 4.1],
)

closing = doc.add_paragraph()
closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("Prepared as an explanatory companion to the SIGCOMM 2025 paper by Feng et al. "
                    "All figures and rendered equations were generated for this explainer and are "
                    "faithful restatements of the concepts and equations in the original work.")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = SLATE_LT

doc.save(OUTPUT)
print("wrote", OUTPUT)
