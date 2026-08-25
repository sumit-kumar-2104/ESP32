# -*- coding: utf-8 -*-
"""
Short companion doc for the 22 July meeting.
Plain talking-points write-up of the two-ESP32 WiFi sensing trial.
Factual, beginner-friendly, no invented numbers.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MEET = r"c:\Users\fs161326\Downloads\SNU Research\22 july meet"

INK   = RGBColor(0x1F, 0x2A, 0x37)
BLUE  = RGBColor(0x2E, 0x5A, 0x88)
SLATE = RGBColor(0x54, 0x63, 0x72)
ACCENT= RGBColor(0xC0, 0x4A, 0x3A)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
st.font.color.rgb = INK

for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)


def shade(par, hexcolor):
    pPr = par._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    pPr.append(sh)


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = INK
    p.space_after = Pt(2)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(11.5); r.font.color.rgb = SLATE; r.italic = True
    p.space_after = Pt(14)


def h(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13.5); r.font.color.rgb = BLUE
    p.space_before = Pt(12); p.space_after = Pt(6)


def body(text, after=8):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(11); r.font.color.rgb = INK
    p.space_after = Pt(after)
    return p


def bullet(text, after=4):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(11); r.font.color.rgb = INK
    p.space_after = Pt(after)
    return p


def callout(head, text, fill="EFF3F7", edge=BLUE):
    p = doc.add_paragraph(); shade(p, fill)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1); p.paragraph_format.right_indent = Inches(0.1)
    r = p.add_run(head + "  "); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = edge
    r2 = p.add_run(text); r2.font.size = Pt(11); r2.font.color.rgb = INK


# ---- header ----
title("WiFi motion & zone sensing with two ESP32 boards")
subtitle("Trial notes for the 22 July meeting \u00b7 MPL Research")

body("Short version: I set up two ESP32 boards to sense movement from the WiFi "
     "signal, ran a person around a zoned room with a camera watching for ground "
     "truth, and recorded everything together. The boards clearly react to motion. "
     "One board is over-sensitive and needs tuning before we can tell zones apart "
     "reliably. It is a working proof of concept, not yet a trained location model.")

# ---- what I built ----
h("What I set up")
bullet("Two ESP32-WROOM-32 boards running the ESPectre firmware. Each reads WiFi "
       "channel data and outputs a Movement Score from 0 to 10, plus a simple "
       "\u201cmotion detected\u201d flag, viewable on a live web dashboard.")
bullet("One board keeps WiFi traffic flowing (pings the router ~100 times per second) "
       "so there is always a signal to measure.")
bullet("A webcam with automatic person detection labels which floor zone the person "
       "is standing in. The room was divided into a 4 \u00d7 3 grid of zones.")
bullet("A collection script saves the zone label, both boards\u2019 movement scores, and "
       "the person\u2019s position into one CSV, all time-aligned.")

# ---- the run ----
h("This run")
bullet("Mode: free walking around the room for about 3.5 minutes.")
bullet("Roughly 4,660 samples recorded into a single combined CSV.")
bullet("Labelling was fully automatic from the camera \u2014 no manual tagging.")

# ---- results ----
h("What the data shows")
body("The time-series plot shows both boards rising and falling as the person moves, "
     "and settling when the room is still \u2014 so the signal is genuinely tracking "
     "activity rather than random noise.", after=6)
bullet("ESP32-2 gives a nicely graded response that changes with the person\u2019s position "
       "\u2014 exactly what we want for telling zones apart.")
bullet("ESP32-1 is very sensitive and mostly sits near the top of the scale (close to 10) "
       "almost everywhere.")
bullet("Most of the walking this run stayed along the middle row of zones, so coverage "
       "was uneven.")

# ---- honest read ----
h("Honest read \u2014 what works and what doesn\u2019t")
callout("Works:", "The whole pipeline runs end to end (sense \u2192 camera-label \u2192 save), "
        "and both boards genuinely detect human motion.", fill="EAF3EA", edge=RGBColor(0x2E,0x6B,0x3A))
callout("Needs fixing:", "ESP32-1 saturates near 10, so right now it can\u2019t distinguish "
        "one zone from another \u2014 it behaves like a coarse presence detector. It needs "
        "threshold, placement, and calibration tuning.", fill="FBEEEA", edge=ACCENT)
callout("Also:", "Because the walk covered mostly one row, the per-zone averages aren\u2019t "
        "a fair comparison yet. We need balanced coverage of every zone.", fill="FBF4EA", edge=RGBColor(0xB5,0x7A,0x1F))

# ---- next steps ----
h("Next steps")
bullet("Re-tune ESP32-1 (threshold / spacing / calibration) so both boards give a "
       "comparable graded response.")
bullet("Do a guided walk that spends equal time in every zone to build fair per-zone "
       "signal fingerprints.")
bullet("Train a simple classifier that predicts the zone from the two movement scores "
       "(and later the raw channel data), using the camera labels as ground truth.")
bullet("Tie this back to the wider goal: doing perception from the wireless signal "
       "itself and reducing reliance on cameras.")

# ---- talking points ----
h("If asked in the meeting")
callout("Why two boards?", "Two viewpoints of the same room give more information than "
        "one \u2014 the plan is to combine them to localise, not just detect presence.")
callout("Why a camera if the goal is camera-free?", "Only for ground-truth labels while "
        "training. Once a model is trained, the camera can be removed.")
callout("Is this novel?", "The trial itself is a standard WiFi-sensing setup; the value "
        "is a working, camera-labelled data pipeline we control end to end, which we can "
        "now build proper experiments on.")

out = os.path.join(MEET, "ESP32_WiFi_Sensing_22July_notes.docx")
doc.save(out)
print("wrote", out)
